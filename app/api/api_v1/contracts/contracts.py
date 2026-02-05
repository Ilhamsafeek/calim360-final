# =====================================================
# FILE: app/api/api_v1/contracts/contracts.py
# COMPLETE VERSION - MERGE CONFLICTS RESOLVED
# =====================================================

from fastapi import APIRouter, Depends, HTTPException, status, Query, File, UploadFile, Form, Request, Body
from sqlalchemy.orm import Session
from sqlalchemy import text, func, and_, or_
from typing import Optional, Dict, List, Any, Union
from datetime import datetime, timedelta
from decimal import Decimal
import logging
import json
import os
import re
from pathlib import Path
import asyncio
import sys
from app.core.config import settings
from fastapi.responses import StreamingResponse
import hashlib
import uuid
import traceback
from app.utils.datetime_helpers import format_datetime_to_iso
from app.services.audit_service import log_contract_action
from weasyprint import HTML
import pypandoc
import tempfile
from app.services.workflow_email_service import WorkflowEmailService
import shutil
import secrets

import hashlib
import uuid
import json
from pathlib import Path


import pdfplumber
import docx
from sqlalchemy.sql import text

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from html.parser import HTMLParser
from bs4 import BeautifulSoup


from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.models.user import User

from pydantic import BaseModel, Field

from app.models.contract import Contract, ContractVersion

from app.services.claude_service import ClaudeService, claude_service
from app.services.blockchain_service import blockchain_service


from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak, HRFlowable
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import base64
from bs4 import BeautifulSoup



logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/contracts", tags=["contracts"])

# =====================================================
# PYDANTIC MODELS
# =====================================================

class AIContractGenerationRequest(BaseModel):
    """Request for AI contract generation"""
    contract_title: str
    contract_type: str
    profile_type: str
    parties: Optional[Dict[str, Any]] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    contract_value: Optional[float] = None
    currency: str = "QAR"
    selected_clauses: Optional[Dict[str, bool]] = None
    clause_descriptions: Optional[List[str]] = None
    jurisdiction: str = "Qatar"
    language: str = "en"
    project_id: Optional[int] = None
    metadata: Optional[Dict[str, Any]] = None
    tags: Optional[str] = None


class UpdateMetadataRequest(BaseModel):
    """Request to update contract AI generation metadata"""
    ai_generation_params: Dict[str, Any]



class ContractHTMLParser(HTMLParser):
    """
    Enhanced HTML parser to convert contract HTML to Word document
    Handles images (including base64 signature images)
    """
    
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.bold = False
        self.italic = False
        self.underline = False
        self.heading_level = 0
        self.in_list = False
        self.list_type = None
        self.list_counter = 0
        
    def handle_starttag(self, tag, attrs):
        """Handle HTML opening tags"""
        tag = tag.lower()
        attrs_dict = dict(attrs)
        
        # Extract style attributes
        style_dict = {}
        if 'style' in attrs_dict:
            styles = attrs_dict['style'].split(';')
            for style in styles:
                if ':' in style:
                    key, value = style.split(':', 1)
                    style_dict[key.strip().lower()] = value.strip()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = int(tag[1])
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.space_after = Pt(12)
            self.current_paragraph.paragraph_format.space_before = Pt(6)
            
        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.space_after = Pt(10)
            self.current_paragraph.paragraph_format.line_spacing = 1.15
            
            # Apply text alignment
            if 'text-align' in style_dict:
                align = style_dict['text-align']
                if align == 'center':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == 'justify':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
        elif tag == 'br':
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()
                self.current_run.add_break()
                
        elif tag == 'strong' or tag == 'b':
            self.bold = True
            
        elif tag == 'em' or tag == 'i':
            self.italic = True
            
        elif tag == 'u':
            self.underline = True
            
        elif tag == 'ul':
            self.in_list = True
            self.list_type = 'ul'
            
        elif tag == 'ol':
            self.in_list = True
            self.list_type = 'ol'
            self.list_counter = 0
            
        elif tag == 'li':
            self.current_paragraph = self.doc.add_paragraph()
            if self.list_type == 'ul':
                self.current_paragraph.style = 'List Bullet'
            else:
                self.list_counter += 1
                self.current_paragraph.style = 'List Number'
            self.current_paragraph.paragraph_format.space_after = Pt(6)
            
        elif tag == 'img':
            # Handle images (including base64 signatures)
            src = attrs_dict.get('src', '')
            if src:
                try:
                    if src.startswith('data:image'):
                        # Base64 image
                        header, base64_data = src.split(',', 1)
                        image_data = base64.b64decode(base64_data)
                        image_stream = BytesIO(image_data)
                        
                        # Create paragraph if needed
                        if not self.current_paragraph:
                            self.current_paragraph = self.doc.add_paragraph()
                        
                        # Add image
                        run = self.current_paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(2.0))
                        
                except Exception as e:
                    # If image fails, add placeholder text
                    if not self.current_paragraph:
                        self.current_paragraph = self.doc.add_paragraph()
                    run = self.current_paragraph.add_run('[Signature Image]')
                    run.font.italic = True
                    
        elif tag == 'div':
            # Divs might have special styling
            if 'text-align' in style_dict:
                self.current_paragraph = self.doc.add_paragraph()
                align = style_dict['text-align']
                if align == 'center':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == 'right':
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
    def handle_endtag(self, tag):
        """Handle HTML closing tags"""
        tag = tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = 0
            self.current_paragraph = None
            
        elif tag == 'p':
            self.current_paragraph = None
            
        elif tag == 'strong' or tag == 'b':
            self.bold = False
            
        elif tag == 'em' or tag == 'i':
            self.italic = False
            
        elif tag == 'u':
            self.underline = False
            
        elif tag == 'ul' or tag == 'ol':
            self.in_list = False
            self.list_type = None
            self.list_counter = 0
            
        elif tag == 'li':
            self.current_paragraph = None
            
    def handle_data(self, data):
        """Handle text content"""
        if not data.strip():
            return
            
        # Create paragraph if none exists
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.space_after = Pt(10)
            
        # Add text run
        run = self.current_paragraph.add_run(data)
        
        # Apply formatting
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        if self.heading_level > 0:
            run.bold = True
            if self.heading_level == 1:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif self.heading_level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(30, 41, 59)
            elif self.heading_level == 3:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(51, 65, 85)
            else:
                run.font.size = Pt(12)
        else:
            if self.bold:
                run.bold = True
            if self.italic:
                run.italic = True
            if self.underline:
                run.underline = True


# =====================================================
# EXISTING ENDPOINT: GET MY CONTRACTS
# =====================================================

@router.get("/my-contracts")
async def get_my_contracts(
    status_filter: Optional[str] = Query(None),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all contracts accessible to the current user"""
    try:
        logger.info(f"Fetching contracts for user: {current_user.email}")
        
        where_conditions = ["(c.created_by = :user_id OR c.company_id = :company_id)"]
        params = {
            "user_id": str(current_user.id),
            "company_id": str(current_user.company_id) if current_user.company_id else None,
            "limit": limit,
            "offset": offset
        }
        
        if status_filter:
            where_conditions.append("c.status = :status_filter")
            params["status_filter"] = status_filter
        
        where_clause = " AND ".join(where_conditions)
        
        query_sql = text(f"""
        SELECT 
            c.id,
            c.contract_number,
            c.contract_title,
            c.contract_type,
            c.status,
            c.start_date,
            c.end_date,
            c.contract_value,
            c.currency,
            c.created_at,
            c.updated_at, 
            c.party_b_name as counterparty_name,
            u.first_name,
            u.last_name
        FROM contracts c
        LEFT JOIN users u ON c.created_by = u.id
        WHERE {where_clause}
        AND contract_type <> 'risk_analysis'
        ORDER BY c.created_at DESC
        LIMIT :limit OFFSET :offset
        """)
        
        result = db.execute(query_sql, params)
        rows = result.fetchall()
        
        contracts = []
        for row in rows:
            contracts.append({
                "id": row[0],
                "contract_number": row[1],
                "contract_title": row[2],
                "contract_type": row[3],
                "status": row[4],
                "start_date": str(row[5]) if row[5] else None,
                "end_date": str(row[6]) if row[6] else None,
                "contract_value": float(row[7]) if row[7] else 0,
                "currency": row[8],
                "created_at": str(row[9]) if row[9] else None,
                "updated_at": str(row[10]) if row[10] else None,
                "counterparty_name": row[11],
                "created_by_name": f"{row[12]} {row[13]}" if row[12] and row[13] else "Unknown"
            })
        
        return {
            "success": True,
            "contracts": contracts,
            "total": len(contracts)
        }
        
    except Exception as e:
        logger.error(f"Error fetching contracts: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to fetch contracts: {str(e)}"
        )

# =====================================================
# TEMPLATES LIST ENDPOINT
# =====================================================

@router.get("/templates/list")
async def list_contract_templates(
    category: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get list of available contract templates"""
    try:
        logger.info(f"📋 Fetching templates for category: {category}")
        
        query = """
            SELECT id, template_name, template_type, template_category,
                   description, is_active, language
            FROM contract_templates
            WHERE is_active = 1
        """
        
        params = {}
        
        if category:
            # Include generic templates marked as 'all' in addition to the requested category
            query += " AND (template_category = :category OR template_category = 'all')"
            params["category"] = category
        
        query += " ORDER BY language DESC, template_name"
        
        result = db.execute(text(query), params)
        rows = result.fetchall()
        
        templates = []
        for row in rows:
            templates.append({
                "id": row[0],
                "name": row[1],
                "type": row[2],
                "category": row[3],
                "description": row[4],
                "is_active": row[5],
                "language": row[6]
            })
        
        logger.info(f" Found {len(templates)} templates")
        
        return {
            "success": True,
            "templates": templates,
            "count": len(templates)
        }
        
    except Exception as e:
        logger.error(f" Error fetching templates: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to fetch templates: {str(e)}"
        )

# =====================================================
# CREATE CONTRACT FROM TEMPLATE
# =====================================================


@router.post("/", status_code=status.HTTP_201_CREATED)
async def create_contract_from_template(
    request: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Create a new contract from template with ACTUAL template content"""
    try:
        # RBAC: Check if user has permission to create contracts
        user_role = getattr(current_user, 'user_role', '').lower() if current_user.user_role else ''
        if user_role in ['viewer', 'guest', 'readonly']:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="You don't have permission to create contracts"
            )
        
        logger.info(f"📋 Creating contract for user: {current_user.email}")
        
        template_id = request.get("template_id")
        
        #  CRITICAL: Load template content FIRST
        template_content = None
        template_content_ar = None
        template_name = None
        template_type = "general"
        
        if template_id:
            logger.info(f" Loading template ID: {template_id}")
            
            # Fetch template with content
            template_query = text("""
                SELECT id, template_name, template_type, template_content, template_content_ar
                FROM contract_templates 
                WHERE id = :template_id AND is_active = 1
            """)
            template_result = db.execute(template_query, {"template_id": template_id}).fetchone()
            
            if not template_result:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Template {template_id} not found or inactive"
                )
            
            template_name = template_result[1]
            template_type = template_result[2] or "general"
            template_content = template_result[3]
            template_content_ar = template_result[4]
            
            logger.info(f" Template loaded: {template_name}")
            logger.info(f"📊 Content length: {len(template_content) if template_content else 0} chars")
            
            # If template has NO content, use meaningful default
            if not template_content or template_content.strip() == "":
                logger.warning(f" Template has no content! Using default structure")
                template_content = f"""
                <div class="contract-document">
                    <h1>{template_name}</h1>
                    <p><strong>Type:</strong> {template_type}</p>
                    <p><em>Template content needs to be added in database.</em></p>
                </div>
                """
        else:
            # No template - blank contract
            logger.info(" Creating blank contract")
            template_content = """
            <div class="contract-document">
                <h1>New Contract</h1>
                <p>Start editing your contract...</p>
            </div>
            """
        
        # Generate contract number
        contract_number = generate_contract_number(db, current_user.company_id)
        
        # Prepare contract data
        contract_data = {
            "company_id": str(current_user.company_id) if current_user.company_id else None,
            "project_id": request.get("project_id"),
            "contract_number": contract_number,
            "contract_title": request.get("contract_title") or f"New Contract - {contract_number}",
            "contract_type": template_type,
            "profile_type": request.get("profile_type", "contractor"),
            "template_id": template_id,
            "status": "draft",
            "workflow_status": "drafting",
            "created_by": current_user.id,
            "created_at": datetime.utcnow(),  #  FIXED
            "updated_at": datetime.utcnow(),  #  FIXED
            "single_tag": request.get("tags") if request.get("tags") else None,  #  Convert array to comma-separated
            
        }
        
        # Insert contract
        insert_query = text("""
            INSERT INTO contracts (
    company_id, project_id, contract_number, contract_title,
    contract_type, profile_type, template_id, status,
    workflow_status, created_by, created_at, updated_at, single_tag
) VALUES (
    :company_id, :project_id, :contract_number, :contract_title,
    :contract_type, :profile_type, :template_id, :status,
    :workflow_status, :created_by, :created_at, :updated_at, :single_tag
)
        """)
        
        result = db.execute(insert_query, contract_data)
        contract_id = result.lastrowid
        
        logger.info(f" Contract created with ID: {contract_id}")
        
        #  CRITICAL: Create contract_versions entry with ACTUAL template content
        version_data = {
            "contract_id": contract_id,
            "version_number": 1,
            "version_type": "draft",
            "contract_content": template_content,  #  ACTUAL CONTENT HERE
            "contract_content_ar": template_content_ar,
            "change_summary": f"Initial contract creation from template: {template_name}" if template_name else "Initial contract creation",
            "is_major_version": False,
            "created_by": current_user.id,
            "created_at": datetime.utcnow()
        }
        
        version_insert = text("""
            INSERT INTO contract_versions (
                contract_id, version_number, version_type, 
                contract_content, contract_content_ar, change_summary,
                is_major_version, created_by, created_at
            ) VALUES (
                :contract_id, :version_number, :version_type,
                :contract_content, :contract_content_ar, :change_summary,
                :is_major_version, :created_by, :created_at
            )
        """)
        
        db.execute(version_insert, version_data)
        db.commit()
        
        logger.info(f" Contract version created with content length: {len(template_content)}")

        log_contract_action(
            db=db,
            action_type="contract_created",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "contract_number": contract_number,
                "contract_title": contract_data["contract_title"],
                "template_id": template_id,
                "template_name": template_name,
                "creation_method": "template"
            },
            ip_address=None  # request.client.host not available in this context
        )
        
        # Return complete contract data
        return {
            "id": contract_id,
            "contract_number": contract_number,
            "contract_title": contract_data["contract_title"],
            "status": "draft",
            "template_content_loaded": bool(template_content and len(template_content) > 100),
            "message": "Contract created successfully with template content"
        }
        
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        logger.error(f" Error creating contract: {str(e)}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to create contract: {str(e)}"
        )


# =====================================================
# UPLOAD CONTRACT - FIXED VERSION
# =====================================================

# @router.post("/upload")
# async def upload_contract(
#     file: UploadFile = File(...),
#     profile_type: str = Form(...),
#     metadata: str = Form(...),
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     """Upload an existing contract document"""
#     try:
#         logger.info(f"Uploading contract file: {file.filename}")
        
#         # Validate file type
#         allowed_types = ["application/pdf", "application/msword", 
#                         "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
#         if file.content_type not in allowed_types:
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="Invalid file type. Only PDF, DOC, and DOCX are allowed."
#             )
        
#         # Parse metadata
#         metadata_dict = json.loads(metadata)
        
#         # Generate contract number
#         contract_number = generate_contract_number(db, current_user.company_id)
        
#         # Save file
#         file_path = save_uploaded_file(file, contract_number, current_user.company_id)
        
#         # Prepare contract data - NO file_url in contracts table
#         contract_data = {
#             "company_id": current_user.company_id,
#             "project_id": metadata_dict.get("project_id"),
#             "contract_number": contract_number,
#             "contract_title": metadata_dict.get("contract_title"),
#             "profile_type": profile_type,
#             "status": "draft",
#             "created_by": current_user.id,
#             "created_at": datetime.utcnow(),
#             "updated_at": datetime.utcnow()
#         }
        
#         # Insert into contracts table (without file_url)
#         insert_query = text("""
#             INSERT INTO contracts (
#                 company_id, project_id, contract_number, contract_title,
#                 profile_type, status, created_by, created_at, updated_at
#             ) VALUES (
#                 :company_id, :project_id, :contract_number, :contract_title,
#                 :profile_type, :status, :created_by, :created_at, :updated_at
#             )
#         """)
        
#         result = db.execute(insert_query, contract_data)
#         db.commit()
        
#         contract_id = result.lastrowid
        
#         # Store the file in contract_versions table
#         version_query = text("""
#             INSERT INTO contract_versions (
#                 contract_id, version_number, file_url, version_type, created_by, created_at
#             ) VALUES (
#                 :contract_id, :version_number, :file_url, :version_type, :created_by, :created_at
#             )
#         """)
        
#         db.execute(version_query, {
#             "contract_id": contract_id,
#             "version_number": 1,
#             "file_url": file_path,
#             "version_type": "uploaded",
#             "change_summary": f"Uploaded from {file.filename} - Text extracted and ready for editing",
#             "created_by": current_user.id,
#             "created_at": datetime.utcnow()
#         })
        
#         db.commit()
        
#         logger.info(f" Upload complete: {contract_number}")
        
#         return {
#             "success": True,
#             "id": contract_id,
#             "contract_number": contract_number,
#             "file_path": relative_path,
#             "content_extracted": True,
#             "content_length": len(extracted_html),
#             "message": "Contract uploaded and text extracted successfully. Ready for editing."
#         }
        
#     except HTTPException:
#         raise
#     except Exception as e:
#         db.rollback()
#         logger.error(f" Upload error: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Failed to upload contract: {str(e)}"
#         )

# =====================================================
# COMPLETE FIX: Add party_b_lead_id and party_b_id to INSERT
# File: app/api/api_v1/contracts/contracts.py
# Function: generate_contract_with_ai()
# =====================================================


def is_internal_user(user) -> bool:
    """
    Check if user is internal (staff/admin) vs external (client/contractor/consultant)
    
    Based on user_type field in database:
    - Internal: 'internal', 'super_admin', 'admin' (LP Global Strategic Page staff)
    - External: 'client', 'contractor', 'consultant', 'sub_contractor' (customer companies)
    
    Internal users can edit without blockchain recording (for draft contracts),
    but edits to finalized contracts WILL trigger tampering detection.
    """
    user_type = getattr(user, 'user_type', '').lower().strip()
    
    # Define internal user types (system staff/administrators)
    internal_types = [
        'internal',      # System internal staff
        'super_admin',   # Super administrators
        'admin',         # Company administrators who are internal
    ]
    
    is_internal = user_type in internal_types
    
    logger.info(f"🔍 User type check: '{user_type}' → {'INTERNAL' if is_internal else 'EXTERNAL'}")
    
    return is_internal



@router.post("/ai-generate")
async def generate_contract_with_ai(
    request_data: AIContractGenerationRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Create contract metadata for AI generation (content will be streamed separately)"""
    try:
        logger.info("🤖 Creating contract for AI generation")
        
        # Extract metadata
        contract_title = request_data.contract_title
        contract_type = request_data.contract_type
        profile_type = request_data.profile_type
        
        # Generate contract number
        contract_number = generate_contract_number(db, current_user.company_id)
        
        # Save generation params as JSON
        generation_params_json = json.dumps({
            "contract_type": request_data.contract_type,
            "profile_type": request_data.profile_type,
            "parties": request_data.parties,
            "selected_clauses": request_data.selected_clauses,
            "clause_descriptions": request_data.clause_descriptions,
            "jurisdiction": request_data.jurisdiction,
            "language": request_data.language,
            "contract_value": request_data.contract_value,
            "currency": request_data.currency,
            "start_date": request_data.start_date,
            "end_date": request_data.end_date,
            "metadata": request_data.metadata
        })
        
        # Create contract record
        contract_data = {
            "company_id": str(current_user.company_id) if current_user.company_id else None,
            "project_id": request_data.project_id,
            "contract_number": contract_number,
            "contract_title": contract_title,
            "contract_type": contract_type,
            "profile_type": profile_type,
            "contract_value": request_data.contract_value,
            "currency": request_data.currency,
            "language": request_data.language,
            "status": "draft",
            "current_version": 1,
            "is_ai_generated": 1,
            "ai_generation_params": generation_params_json,
            "party_b_lead_id": None, 
            "party_b_id": None,      
            "created_by": str(current_user.id),
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow(),
            "single_tag": request_data.tags if request_data.tags else None,
        }
        
        #  FIXED: Include party_b_lead_id and party_b_id in INSERT
        result = db.execute(text("""
            INSERT INTO contracts (company_id, project_id, contract_number, contract_title,
                         contract_type, profile_type, contract_value, currency, language,
                         status, current_version, is_ai_generated, ai_generation_params,
                         party_b_lead_id, party_b_id, single_tag,
                         created_by, created_at, updated_at)
    VALUES (:company_id, :project_id, :contract_number, :contract_title,
            :contract_type, :profile_type, :contract_value, :currency, :language,
            :status, :current_version, :is_ai_generated, :ai_generation_params,
            :party_b_lead_id, :party_b_id, :single_tag,
            :created_by, :created_at, :updated_at)
        """), contract_data)
        
        # Use lastrowid for MySQL compatibility
        contract_id = result.lastrowid
        db.commit()

        log_contract_action(
            db=db,
            action_type="contract_created",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "contract_number": contract_number,
                "contract_title": contract_title,
            },
            ip_address=None
        )
        
        logger.info(f" Contract created: {contract_number} (ID: {contract_id})")
        
        return {
            "id": contract_id,
            "contract_id": contract_id,
            "contract_number": contract_number,
            "message": "Contract created. AI generation can now be streamed."
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error creating contract: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to create contract: {str(e)}"
        )

# =====================================================
# STATISTICS ENDPOINT - FIXED VERSION
# =====================================================

@router.get("/statistics")
async def get_contract_statistics(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get dashboard statistics - INCLUDES MY PENDING APPROVALS
    SCR_010 - Contract Dashboard
    """
    from app.models.contract import Contract
    from app.models.project import Project
    from app.models.obligation import Obligation
    
    company_id = current_user.company_id
    user_id = current_user.id
    
    # Total contracts (including AI-generated)
    total_contracts = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.is_deleted == False,
        Contract.contract_type != 'risk_analysis'
    ).scalar() or 0
    
    # Active contracts
    active_contracts = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.status.in_(['active', 'signed', 'executed']),
        Contract.is_deleted == False,
        Contract.contract_type != 'risk_analysis'
    ).scalar() or 0
    
    # Pending review
    pending_review = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.status.in_(['pending_review', 'review', 'pending_approval']),
        Contract.is_deleted == False,
        Contract.contract_type != 'risk_analysis'
    ).scalar() or 0
    
    # Expiring soon (within 30 days)
    thirty_days = datetime.now() + timedelta(days=30)
    today = datetime.now()
    
    expiring_soon = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.status.in_(['active', 'signed', 'executed']),
        Contract.end_date.isnot(None),
        Contract.end_date <= thirty_days,
        Contract.end_date >= today,
        Contract.is_deleted == False,
        Contract.contract_type != 'risk_analysis'
    ).scalar() or 0
    
    # Completed contracts
    completed_contracts = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        or_(Contract.status == 'completed', Contract.status == 'expired'),
        Contract.is_deleted == False,
        Contract.contract_type != 'risk_analysis'
    ).scalar() or 0
    
    # Active projects
    active_projects = db.query(func.count(Project.id)).filter(
        Project.company_id == company_id,
        Project.status == 'active'
    ).scalar() or 0
    
    # Due obligations (within 7 days)
    seven_days = datetime.now() + timedelta(days=7)
    due_obligations = db.query(func.count(Obligation.id)).filter(
        Obligation.due_date <= seven_days,
        Obligation.due_date >= today,
        Obligation.status.in_(['PENDING', 'IN_PROGRESS'])
    ).scalar() or 0
    
    # 🆕 MY PENDING APPROVALS - Contracts waiting for current user's approval
    # This checks both workflow_stages and approval_requests tables
    my_pending_approvals = db.execute(text("""
        SELECT COUNT(DISTINCT c.id) as count
        FROM contracts c
        WHERE c.company_id = :company_id
        AND c.contract_type <> 'risk_analysis'
        AND c.is_deleted = 0
        AND c.status IN ('pending_approval', 'pending_review', 'review', 'approval','counterparty_internal_review')
        AND (
            -- Check workflow_stages for pending approvals assigned to this user
            EXISTS (
                SELECT 1 
                FROM workflow_instances wi
                JOIN workflow_stages ws ON wi.id = ws.workflow_instance_id
                WHERE wi.contract_id = c.id
                AND wi.status IN ('active', 'in_progress', 'pending')
                AND ws.approver_user_id = :user_id
                AND ws.status = 'pending'
            )
            OR
            -- Check approval_requests for pending approvals assigned to this user
            EXISTS (
                SELECT 1
                FROM approval_requests ar
                WHERE ar.contract_id = c.id
                AND ar.approver_id = :user_id
                AND ar.responded_at IS NULL
                AND (ar.status = 'pending' OR ar.approval_status = 'pending')
            )
        )
    """), {
        "company_id": company_id,
        "user_id": user_id
    }).scalar() or 0
    
    # Module counts - FIXED to include AI-generated contracts
    
    # Drafting count - includes ALL draft contracts and AI-generated ones
    drafting_count = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.is_deleted == False,
        or_(
            # Explicit workflow statuses for drafting
            Contract.workflow_status.in_(['draft', 'internal_review', 'clause_analysis']),
            # NULL workflow_status with drafting statuses
            and_(
                Contract.workflow_status.is_(None),
                Contract.status.in_(['draft', 'pending_review', 'in_progress'])
            ),
            # Explicitly include any contract with status='draft' (AI-generated fall here)
            Contract.status == 'draft'
        )
    ).scalar() or 0
    
    # Negotiation count
    negotiation_count = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.is_deleted == False,
        or_(
            Contract.workflow_status.in_(['external_review', 'negotiation', 'approval']),
            and_(
                Contract.workflow_status.is_(None),
                Contract.status.in_(['negotiation', 'pending_approval'])
            )
        )
    ).scalar() or 0
    
    # Operations count
    operations_count = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.is_deleted == False,
        Contract.status.in_(['active', 'expired', 'terminated', 'completed', 'executed', 'signed'])
    ).scalar() or 0
    
    # Additional stat: Count of AI-generated contracts
    ai_generated_count = db.query(func.count(Contract.id)).filter(
        Contract.company_id == company_id,
        Contract.is_deleted == False,
        Contract.id.in_(
            db.query(ContractVersion.contract_id).filter(
                ContractVersion.version_type == 'ai_generated'
            )
        )
    ).scalar() or 0
    
    logger.info(f"📊 Statistics - Total: {total_contracts}, Drafting: {drafting_count}, AI-Generated: {ai_generated_count}, My Pending Approvals: {my_pending_approvals}")
    
    return {
        "total_contracts": total_contracts,
        "active_contracts": active_contracts,
        "pending_review": pending_review,
        "expiring_soon": expiring_soon,
        "completed_contracts": completed_contracts,
        "active_projects": active_projects,
        "due_obligations": due_obligations,
        "my_pending_approvals": my_pending_approvals,  # 🆕 NEW: Contracts awaiting my approval
        "drafting_count": drafting_count,
        "negotiation_count": negotiation_count,
        "operations_count": operations_count,
        "ai_generated_count": ai_generated_count
    }


# =====================================================
# GET CONTRACTS LIST
# =====================================================
@router.get("")
async def get_contracts(
    module: str = Query("drafting"),
    status: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    search: Optional[str] = Query(None),
    type: Optional[str] = Query(None),
    profile: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get contracts list with filters - SCR_010 - Gets counterparty company name"""
    company_id = current_user.company_id
    
    # Base query
    query = db.query(Contract).filter(
        or_(
            Contract.company_id == company_id,
            Contract.party_b_id == company_id
        ),
        Contract.is_deleted == 0,
        Contract.contract_type != 'risk_analysis'
    )
    
    # Module filter
    if module == "drafting":
        query = query.filter(
            # Displays all the contracts with any status
            # Contract.status.in_(['draft', 'pending_review', 'in_progress', 'review','review_completed','counterparty_internal_review','negotiation_completed','negotiation'])

        )
    elif module == "negotiation":
        query = query.filter(
            or_(
                # Contract.workflow_status.in_(['external_review', 'negotiation', 'approval']),
                # and_(
                #     Contract.workflow_status.is_(None),
                    Contract.status.notin_(['draft', 'executed'])
                ) 
            # )
        )
    elif module == "operations":
        query = query.filter(
            Contract.status.in_(['executed'])
        )
    
    # Status filter for sub-tabs
    if status and status != "all":
        if status == "my":
            query = query.filter(Contract.created_by == current_user.id)
        elif status == "draft":
            query = query.filter(Contract.status == 'draft')
        elif status == "review":
            query = query.filter(
                or_(
                    Contract.status == 'review',
                    Contract.status == 'pending_review',
                    Contract.workflow_status == 'review'
                )
            )
        elif status == "expiring":
            thirty_days = datetime.now() + timedelta(days=30)
            query = query.filter(
                Contract.end_date.isnot(None),
                Contract.end_date <= thirty_days,
                Contract.end_date >= datetime.now()
            )
        else:
            query = query.filter(Contract.status == status)
    
    # Search filter
    if search:
        search_term = f"%{search}%"
        query = query.filter(
            or_(
                Contract.contract_title.ilike(search_term),
                Contract.contract_number.ilike(search_term)
            )
        )
    
    # Type filter
    if type:
        query = query.filter(Contract.contract_type == type)
    
    # Profile filter
    if profile:
        query = query.filter(Contract.profile_type == profile)
    
    # Get total count before pagination
    total = query.count()
    
    # Apply pagination
    offset = (page - 1) * limit
    contracts = query.order_by(Contract.created_at.desc()).offset(offset).limit(limit).all()
    
    # Convert to dict
    result = []
    for contract in contracts:
        # Get latest version
        latest_version = db.query(ContractVersion).filter(
            ContractVersion.contract_id == contract.id
        ).order_by(ContractVersion.version_number.desc()).first()
        
        # Get creator info
        creator = db.query(User).filter(User.id == contract.created_by).first()
        
        # Get counterparty company name using raw SQL to avoid ORM conflicts
        counterparty_name = "Not specified"
        if contract.party_b_id:
            counterparty_query = text("""
                SELECT c.company_name 
                FROM users u
                JOIN companies c ON u.company_id = c.id
                WHERE u.id = :user_id
                LIMIT 1
            """)
            counterparty_result = db.execute(counterparty_query, {"user_id": contract.party_b_id}).fetchone()
            if counterparty_result and counterparty_result.company_name:
                counterparty_name = counterparty_result.company_name
            elif contract.party_b_name:
                # Fallback to party_b_name if company lookup fails
                counterparty_name = contract.party_b_name
        elif contract.party_b_name:
            counterparty_name = contract.party_b_name
        
        result.append({
            "id": contract.id,
            "contract_number": contract.contract_number,
            "title": contract.contract_title,
            "counterparty": counterparty_name,
            "status": contract.status,
            "single_tag": contract.single_tag,
            "contract_type": contract.contract_type,
            "module": module,
            "value": float(contract.contract_value) if contract.contract_value else 0,
            "currency": contract.currency or "QAR",
            "created_at": contract.created_at.isoformat() + "Z" if contract.created_at else None,
            "updated_at": contract.updated_at.isoformat() + "Z" if contract.updated_at else None,       
            "created_by": f"{creator.first_name} {creator.last_name}" if creator else "Unknown",
            "current_version": latest_version.version_number if latest_version else 1,
            "priority": None,
            "template": contract.is_template if hasattr(contract, 'is_template') else False
        })
    
    logger.info(f"📊 Retrieved {len(result)} contracts out of {total} total for module '{module}'")
    
    return {
        "success": True,
        "contracts": result,
        "pagination": {
            "total": total,
            "page": page,
            "limit": limit,
            "total_pages": (total + limit - 1) // limit,
            "has_more": offset + limit < total
        }
    }

# =====================================================
# DELETE CONTRACT
# =====================================================

@router.delete("/{contract_id}")
async def delete_contract(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Delete a contract (soft delete) - SCR_010"""
    from app.models.contract import Contract
    
    contract = db.query(Contract).filter(
        Contract.id == contract_id,
        Contract.company_id == current_user.company_id
    ).first()
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    # if contract.status not in ['draft', 'pending_review']:
    #     raise HTTPException(
    #         status_code=403,
    #         detail="Cannot delete contract in current status"
    #     )
    
    if hasattr(Contract, 'is_deleted'):
        contract.is_deleted = True
        contract.deleted_at = datetime.now()
    else:
        db.delete(contract)
    
    contract.updated_by = current_user.id
    contract.updated_at = datetime.now()
    
    try:
        db.commit()
        log_contract_action(
            db=db,
            action_type="contract_deleted",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "contract_number": contract.contract_number,
                "contract_title": contract.contract_title
            },
            ip_address=None
        )

        return {
            "success": True,
            "message": "Contract deleted successfully",
            "contract_id": contract_id
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to delete contract: {str(e)}")

# =====================================================
# UPDATE CONTRACT
# =====================================================

@router.put("/{contract_id}")
async def update_contract(
    contract_id: int,
    title: Optional[str] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Update contract details"""
    from app.models.contract import Contract
    
    contract = db.query(Contract).filter(
        Contract.id == contract_id,
        Contract.company_id == current_user.company_id
    ).first()
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    if title:
        contract.contract_title = title
    if status:
        contract.status = status
    
    contract.updated_by = current_user.id
    contract.updated_at = datetime.now()
    
    try:
        db.commit()
        db.refresh(contract)
        return {
            "success": True,
            "message": "Contract updated successfully",
            "contract": {
                "id": contract.id,
                "title": contract.contract_title,
                "status": contract.status
            }
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to update contract: {str(e)}")

# =====================================================
# CONTRACT EDITOR ENDPOINTS
# =====================================================
@router.get("/edit/{contract_id}")
async def get_contract_editor_data(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get contract data for editor with execution certificate"""
    is_internal = is_internal_user(current_user)
    try:
        query = text("""
            SELECT 
                c.id,
                c.contract_number,
                c.contract_title,
                c.contract_type,
                c.language,
                c.status,
                c.approval_status,
                c.workflow_status,
                c.created_at,
                c.created_by as created_by_id,
                c.updated_at,
                c.party_b_id,
                c.party_b_lead_id,
                c.action_person_id,
                c.signed_date,
                c.party_esignature_authority_id,
                c.counterparty_esignature_authority_id,
                c.effective_date,
                comp.company_name,
                party_b_comp.company_name as party_b_company_name,
                c.company_id,
                c.is_ai_generated,
                c.ai_generation_params,
                CONCAT(u.first_name, ' ', u.last_name) as created_by_name,
                cv.contract_content as content,
                cv.version_number as current_version
            FROM contracts c
            LEFT JOIN companies comp ON c.company_id = comp.id
            LEFT JOIN companies party_b_comp ON c.party_b_id = party_b_comp.id
            LEFT JOIN users u ON c.created_by = u.id
            LEFT JOIN contract_versions cv ON c.id = cv.contract_id 
                AND cv.version_number = (
                    SELECT MAX(version_number) 
                    FROM contract_versions 
                    WHERE contract_id = c.id
                )
            WHERE c.id = :contract_id
            LIMIT 1
        """)
        
        result = db.execute(query, {"contract_id": contract_id}).fetchone()
        
        if not result:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        is_initiator = current_user.id == result.created_by_id
        is_counterparty = current_user.company_id == result.party_b_id

        is_same_company_as_initiator = current_user.company_id == result.company_id

        
        # ===== INITIATOR WORKFLOW (Party A) =====
        workflow_query = text("""
          SELECT 
            wi.id as workflow_instance_id,
            wi.workflow_id,
            wi.status as workflow_status,
            wi.current_step,
            w.workflow_name as template_name,
            w.company_id,
            w.is_master
        FROM workflow_instances wi
        LEFT JOIN workflows w ON wi.workflow_id = w.id
        WHERE wi.contract_id = :contract_id
        AND w.company_id = :company_id
        AND w.is_active = 1
        AND wi.status IN ('pending', 'active', 'in_progress', 'completed')
        ORDER BY w.is_master ASC
        LIMIT 1
        """)
        
        workflow = db.execute(workflow_query, {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        }).fetchone()
        
        # Get workflow steps with assignee information
        workflow_steps = []
        total_steps = 0
        
        if workflow and workflow.workflow_id:
            steps_query = text("""
                SELECT 
                    ws.id as step_id,
                    ws.step_number,
                    ws.step_name,
                    ws.step_type,
                    ws.assignee_user_id,
                    ws.assignee_role,
                    CONCAT(u.first_name, ' ', u.last_name) as assignee_name,
                    u.email as assignee_email,
                    CASE 
                        WHEN ws.step_number < :current_step THEN 'completed'
                        WHEN ws.step_number = :current_step THEN 'active'
                        ELSE 'pending'
                    END as step_status
                FROM workflow_steps ws
                LEFT JOIN users u ON ws.assignee_user_id = u.id
                WHERE ws.workflow_id = :workflow_id
                ORDER BY ws.step_number ASC
            """)
            
            steps_result = db.execute(steps_query, {
                "workflow_id": workflow.workflow_id,
                "current_step": workflow.current_step
            }).fetchall()
            
            total_steps = len(steps_result)
            
            workflow_steps = [
                {
                    "step_id": step.step_id,
                    "step_number": step.step_number,
                    "step_name": step.step_name,
                    "step_type": step.step_type,
                    "assignee_user_id": step.assignee_user_id,
                    "assignee_name": step.assignee_name,
                    "assignee_email": step.assignee_email,
                    "assignee_role": step.assignee_role,
                    "status": step.step_status,
                    "is_current": workflow.current_step == step.step_number
                }
                for step in steps_result
            ]
        
        # Check if it's current user's turn in workflow
        is_my_workflow_turn = False
        if workflow and workflow_steps:
            for step in workflow_steps:
                if step["is_current"]:
                    if step["assignee_user_id"] == current_user.id:
                        is_my_workflow_turn = True
                        break
        

        is_party_b_lead = False
        if current_user.id==result.party_b_lead_id:
            is_party_b_lead = True
        

        is_esignee = False
        if (current_user.id==result.party_esignature_authority_id or current_user.id==result.counterparty_esignature_authority_id):
            is_esignee = True
        
        
        # Get version history
        version_query = text("""
            SELECT 
                cv.version_number,
                cv.created_at,
                cv.change_summary,
                cv.created_by,
                CONCAT(u.first_name, ' ', u.last_name) as created_by_name
            FROM contract_versions cv
            LEFT JOIN users u ON cv.created_by = u.id
            WHERE cv.contract_id = :contract_id
            ORDER BY cv.version_number DESC
            LIMIT 10
        """)
        
        versions = db.execute(version_query, {"contract_id": contract_id}).fetchall()
        
        # Get current approver (for initiator workflow)
        current_approver_query = text("""
            SELECT 
            u.first_name, 
            u.last_name, 
            u.email, 
            u.department, 
            ws.step_type,
            w.is_master
        FROM workflow_instances wi
        JOIN workflows w ON wi.workflow_id = w.id
        JOIN workflow_steps ws ON wi.workflow_id = ws.workflow_id 
            AND wi.current_step = ws.step_number
        LEFT JOIN users u ON ws.assignee_user_id = u.id
        WHERE wi.contract_id = :contract_id 
        AND wi.status IN ('active', 'in_progress')
        AND w.company_id = :company_id
        AND w.is_active = 1
        ORDER BY w.is_master ASC
        LIMIT 1
        """)
        current_approver = db.execute(current_approver_query, {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        }).fetchone()
        
        # ===== EXECUTION CERTIFICATE DATA =====

        # ===== EXECUTION CERTIFICATE DATA =====
        certificate_data = None
        
        #  FIXED: Check for signature, signed, and executed status
        if result.status in ('signature', 'signed', 'executed'):  # Added 'signature' status
            try:
                #  REMOVED: contract_metadata table query (doesn't exist)
                #  NEW: Always generate from signatories table
                
                executed_by_name = f"{current_user.first_name} {current_user.last_name}"
                
                #  FIXED: Get signatories with signature_data and signature_method
                signatories_query = text("""
                    SELECT 
                        s.signer_type,
                        s.has_signed,
                        s.signed_at,
                        s.signature_data,
                        s.signature_method,
                        s.ip_address,
                        s.signing_order,
                        u.first_name,
                        u.last_name,
                        u.email
                    FROM signatories s
                    LEFT JOIN users u ON s.user_id = u.id
                    WHERE s.contract_id = :contract_id
                    ORDER BY s.signing_order
                """)
                
                signatories = db.execute(signatories_query, {"contract_id": contract_id}).fetchall()
                
                certificate_data = {
                    "contract_id": contract_id,
                    "contract_number": result.contract_number,
                    "contract_title": result.contract_title,
                    "execution_date": result.effective_date.isoformat() if result.effective_date else None,
                    "signed_date": result.signed_date.isoformat() if result.signed_date else None,
                    "executed_by": executed_by_name,
                    "executed_by_email": current_user.email,
                    "signatories": []
                }
                
                for sig in signatories:
                    # Construct full name
                    signer_name = "Pending"
                    if sig.has_signed and sig.first_name and sig.last_name:
                        signer_name = f"{sig.first_name} {sig.last_name}"
                    elif sig.first_name:
                        signer_name = sig.first_name
                    
                    certificate_data["signatories"].append({
                        "signer_type": sig.signer_type,
                        "name": signer_name,
                        "email": sig.email or "",
                        "has_signed": bool(sig.has_signed),
                        "signed_at": sig.signed_at.isoformat() if sig.signed_at else None,
                        "signature_data": sig.signature_data or "",  #  ADDED
                        "signature_method": sig.signature_method or "draw",  #  ADDED
                        "ip_address": sig.ip_address or "",
                        "signing_order": sig.signing_order
                    })
                
                logger.info(f" Certificate data loaded with {len(certificate_data['signatories'])} signatories")
                    
            except Exception as cert_error:
                logger.warning(f"⚠️ Could not retrieve certificate data: {str(cert_error)}")
                certificate_data = None
        
        
        # ===== RETURN RESPONSE =====
        return {
            "success": True,
            "contract": {
                "id": result.id,
                "contract_number": result.contract_number,
                "title": result.contract_title,
                "type": result.contract_type,
                "status": result.status,
                "workflow_status": result.workflow_status,
                "approval_status": result.approval_status,
                "content": result.content if result.content else "",
                "company_name": result.company_name,
                "company_id": result.company_id,
                "party_b_id": result.party_b_id,
                "party_b_company_name": result.party_b_company_name,
                "is_party_b_lead": is_party_b_lead,
                "created_by": result.created_by_name,
                "created_by_id": result.created_by_id,
                "created_at": result.created_at.isoformat() if result.created_at else None,
                "updated_at": result.updated_at.isoformat() if result.updated_at else None,
                "signed_date": result.signed_date.isoformat() if result.signed_date else None,
                "effective_date": result.effective_date.isoformat() if result.effective_date else None,
                "current_version": result.current_version if result.current_version else 1,
                "is_initiator": is_initiator,
                "is_counterparty": is_counterparty,
                "is_ai_generated": result.is_ai_generated,
                "ai_generation_params": result.ai_generation_params,
                "language": result.language,
                "is_esignee": is_esignee,
                "action_person_id": result.action_person_id,
                "is_same_company_as_initiator": is_same_company_as_initiator,
            },
            "workflow": {
                "status": workflow.workflow_status if workflow else "not_configured",
                "current_stage": workflow.current_step if workflow else 0,
                "total_stages": total_steps,
                "template_name": workflow.template_name if workflow else None,
                "steps": workflow_steps,
                "is_my_workflow_turn": is_my_workflow_turn
            } if workflow else None,
            "versions": [
                {
                    "version": str(v.version_number),
                    "created_at": v.created_at.isoformat() if v.created_at else None,
                    "notes": v.change_summary if v.change_summary else "No notes",
                    "created_by": v.created_by_name if v.created_by_name else "Unknown"
                }
                for v in versions
            ],
            "current_approver": {
                "name": f"{current_approver.first_name} {current_approver.last_name}".strip() if current_approver else None,
                "email": current_approver.email if current_approver else None,
                "department": current_approver.department if current_approver else None,
                "step_type": current_approver.step_type if current_approver else None
            } if current_approver else None,
            "certificate": certificate_data,  #  ADDED CERTIFICATE DATA
            "is_internal_user": is_internal,
            
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching contract editor data: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
        

#  HELPER FUNCTIONS (Place these at the TOP of the file, outside the router functions)
def _generate_success_message(is_internal: bool, tampered_banner_added: bool, original_hash_stored: bool) -> str:
    """Generate appropriate success message based on operation"""
    if is_internal and tampered_banner_added and original_hash_stored:
        return "Draft saved with TAMPERED banner - Original hash stored on blockchain for verification"
    elif is_internal and tampered_banner_added:
        return "Draft saved with TAMPERED banner (blockchain storage failed - verification may not work)"
    elif is_internal:
        return "Internal edit saved"
    else:
        return "Draft saved successfully"


def _generate_verification_note(is_internal: bool, original_hash_stored: bool) -> str:
    """Generate verification note for response"""
    if is_internal and original_hash_stored:
        return "Verification will detect tampering: blockchain hash (clean) ≠ current content (tampered)"
    elif is_internal:
        return "Warning: Original hash not stored - verification may not detect tampering"
    else:
        return "Verification will succeed: blockchain hash = current content"



@router.post("/save-draft/{contract_id}")
async def save_contract_draft(
    contract_id: int,
    content: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Save contract draft with tampering detection for internal users"""
    try:
        is_internal = is_internal_user(current_user)
        new_content = content.get("content", "")
        
        # Get current version
        version_check = text("SELECT MAX(version_number) as max_version FROM contract_versions WHERE contract_id = :contract_id")
        result = db.execute(version_check, {"contract_id": contract_id}).fetchone()
        current_max_version = result.max_version if result and result.max_version else 0
        
        # Check if content changed
        content_changed = True
        if current_max_version > 0:
            current_content_query = text("SELECT contract_content FROM contract_versions WHERE contract_id = :contract_id ORDER BY version_number DESC LIMIT 1")
            current_content_result = db.execute(current_content_query, {"contract_id": contract_id}).fetchone()
            current_content = current_content_result.contract_content if current_content_result else ""
            content_changed = (current_content.strip() != new_content.strip())
        
        if not content_changed:
            return {"success": True, "message": "No changes detected", "version": current_max_version}
        
        next_version = current_max_version + 1
        blockchain_success = False
        final_content = new_content
        
        #  FIX: For INTERNAL users - DON'T store on blockchain
        if is_internal:
            logger.info(f"⏭️ Internal user - SKIPPING blockchain storage")
            logger.info(f"   Verification will detect no blockchain record = TAMPERED")
            # Don't store blockchain hash - this creates the tampering detection
            blockchain_success = False
            final_content = new_content  # Save content as-is, no banner
        
        # For EXTERNAL users - Store hash AFTER saving
        elif not is_internal:
            # Will store blockchain hash after version is created
            pass
        
        # Create new version
        version_query = text("""
            INSERT INTO contract_versions 
            (contract_id, version_number, contract_content, change_summary, version_type, created_by, created_at)
            VALUES (:contract_id, :version_number, :content, :change_summary, 'draft', :user_id, NOW())
        """)
        
        change_summary = "Internal edit (No blockchain)" if is_internal else "Draft updated"
        
        db.execute(version_query, {
            "contract_id": contract_id,
            "version_number": next_version,
            "content": final_content,
            "change_summary": change_summary,
            "user_id": current_user.id
        })
        
        # Update timestamp
        db.execute(text("UPDATE contracts SET updated_at = NOW() WHERE id = :contract_id"), {"contract_id": contract_id})
        
        # For EXTERNAL users - Store blockchain hash now
        if not is_internal:
            try:
                blockchain_result = await blockchain_service.store_contract_hash_with_logging(
                    contract_id=contract_id,
                    document_content=final_content,
                    uploaded_by=current_user.id,
                    company_id=current_user.company_id,
                    db=db
                )
                if blockchain_result.get("success"):
                    blockchain_success = True
            except Exception as e:
                logger.error(f"❌ Blockchain error: {str(e)}")
        
        db.commit()
        
        return {
            "success": True,
            "message": "Internal edit saved (no blockchain)" if is_internal else "Draft saved",
            "version": next_version,
            "blockchain_success": blockchain_success,
            "internal_edit": is_internal
        }
        
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))



@router.post("/mark-tampered/{contract_id}")
async def mark_contract_as_tampered(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Mark a contract as tampered by adding a warning banner.
    Exactly replicates the SQL UPDATE query.
    """
    
    try:
        # Execute the exact same SQL query
        sql = """
        UPDATE contract_versions cv
        JOIN contracts c ON cv.contract_id = c.id 
            AND cv.version_number = c.current_version
        SET cv.contract_content = CONCAT(
            '<div style="background:red;color:white;padding:20px;"><h1>⚠️ TAMPERED</h1></div>',
            cv.contract_content
        )
        WHERE cv.contract_id = :contract_id
        """
        
        result = db.execute(sql, {"contract_id": contract_id})
        db.commit()
        
        if result.rowcount == 0:
            raise HTTPException(status_code=404, detail=f"Contract {contract_id} not found")
        
        return {
            "success": True,
            "message": f"Contract {contract_id} marked as TAMPERED",
            "contract_id": contract_id,
            "rows_updated": result.rowcount
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))



@router.post("/send-for-signature")
async def send_contract_for_signature(
    data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Send contract for signature - Creates signatory records for BOTH parties
    FIXED: Ensures 2 signatories are created (client + company)
    """
    try:
        contract_id = data.get("contract_id")
        
        logger.info(f"📝 Sending contract {contract_id} for signature")
        
        # STEP 1: Verify contract exists and is in approval status
        contract_query = text("""
            SELECT 
                c.id, c.contract_number, c.status, c.company_id,
                c.party_b_id, c.party_b_lead_id
            FROM contracts c
            WHERE c.id = :contract_id
        """)
        
        contract = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        if contract.status not in ['approval', 'negotiation']:
            return {
                "success": False,
                "detail": f"Contract must be in approval or negotiation status. Current: {contract.status}"
            }
        
        logger.info(f" Contract verified: {contract.contract_number}")
        
        # STEP 2: Update contract status to 'signature'
        update_contract = text("""
            UPDATE contracts
            SET status = 'signature',
                workflow_status = 'signature',
                updated_at = NOW()
            WHERE id = :contract_id
        """)
        
        db.execute(update_contract, {"contract_id": contract_id})
        logger.info(f" Contract status updated to 'signature'")
        
        # STEP 3: Get E-SIGN authority users from workflow (if exists)
        workflow_query = text("""
            SELECT 
                ws.assignee_user_id,
                ws.step_name,
                u.email,
                u.company_id,
                u.first_name,
                u.last_name
            FROM workflow_instances wi
            INNER JOIN workflow_steps ws ON wi.workflow_id = ws.workflow_id
            INNER JOIN users u ON ws.assignee_user_id = u.id
            WHERE wi.contract_id = :contract_id
            AND ws.step_type = 'e_sign_authority'
            ORDER BY ws.step_number
        """)
        
        workflow_signers = db.execute(workflow_query, {"contract_id": contract_id}).fetchall()
        
        # STEP 4: Clear existing signatory records (in case of re-sending)
        delete_existing = text("""
            DELETE FROM signatories 
            WHERE contract_id = :contract_id
        """)
        db.execute(delete_existing, {"contract_id": contract_id})
        logger.info(f"🗑️ Cleared existing signatory records")
        
        # STEP 5: Create signatory records for BOTH parties
        signatories_created = 0
        
        if workflow_signers and len(workflow_signers) >= 2:
            # Use workflow to determine signatories
            logger.info(f" Found {len(workflow_signers)} workflow E-SIGN authorities")
            
            for idx, signer in enumerate(workflow_signers):
                # Determine signer_type based on company
                if signer.company_id == contract.company_id:
                    signer_type = 'company'
                else:
                    signer_type = 'client'
                
                insert_signatory = text("""
                    INSERT INTO signatories
                    (contract_id, user_id, signer_type, company_id,
                     signing_order, email, has_signed, created_at)
                    VALUES
                    (:contract_id, :user_id, :signer_type, :company_id,
                     :signing_order, :email, 0, NOW())
                """)
                
                db.execute(insert_signatory, {
                    "contract_id": contract_id,
                    "user_id": signer.assignee_user_id,
                    "signer_type": signer_type,
                    "company_id": signer.company_id,
                    "signing_order": idx + 1,
                    "email": signer.email
                })
                
                logger.info(f" Created signatory: {signer.first_name} {signer.last_name} ({signer_type})")
                signatories_created += 1
        
        else:
            # NO WORKFLOW - Create default signatories for initiator company and counterparty
            logger.info(f"⚠️ No workflow found - creating default signatories")
            
            # Get initiator (current user's company representative)
            company_rep_query = text("""
                SELECT id, email, first_name, last_name, company_id
                FROM users
                WHERE company_id = :company_id
                AND id = :user_id
                LIMIT 1
            """)
            
            company_rep = db.execute(company_rep_query, {
                "company_id": contract.company_id,
                "user_id": current_user.id
            }).fetchone()
            
            # Get counterparty lead
            if contract.party_b_lead_id:
                counterparty_query = text("""
                    SELECT id, email, first_name, last_name, company_id
                    FROM users
                    WHERE id = :party_b_lead_id
                    LIMIT 1
                """)
                
                counterparty = db.execute(counterparty_query, {
                    "party_b_lead_id": contract.party_b_lead_id
                }).fetchone()
            else:
                # Get any user from counterparty company
                counterparty_query = text("""
                    SELECT id, email, first_name, last_name, company_id
                    FROM users
                    WHERE company_id = :party_b_id
                    LIMIT 1
                """)
                
                counterparty = db.execute(counterparty_query, {
                    "party_b_id": contract.party_b_id
                }).fetchone()
            
            # Create COMPANY signatory
            if company_rep:
                insert_company_sig = text("""
                    INSERT INTO signatories
                    (contract_id, user_id, signer_type, company_id,
                     signing_order, email, has_signed, created_at)
                    VALUES
                    (:contract_id, :user_id, 'company', :company_id,
                     1, :email, 0, NOW())
                """)
                
                db.execute(insert_company_sig, {
                    "contract_id": contract_id,
                    "user_id": company_rep.id,
                    "company_id": company_rep.company_id,
                    "email": company_rep.email
                })
                
                logger.info(f" Created COMPANY signatory: {company_rep.first_name} {company_rep.last_name}")
                signatories_created += 1
            
            # Create CLIENT signatory
            if counterparty:
                insert_client_sig = text("""
                    INSERT INTO signatories
                    (contract_id, user_id, signer_type, company_id,
                     signing_order, email, has_signed, created_at)
                    VALUES
                    (:contract_id, :user_id, 'client', :company_id,
                     2, :email, 0, NOW())
                """)
                
                db.execute(insert_client_sig, {
                    "contract_id": contract_id,
                    "user_id": counterparty.id,
                    "company_id": counterparty.company_id,
                    "email": counterparty.email
                })
                
                logger.info(f" Created CLIENT signatory: {counterparty.first_name} {counterparty.last_name}")
                signatories_created += 1
        
        # STEP 6: Verify we have 2 signatories
        if signatories_created < 2:
            db.rollback()
            logger.error(f"❌ Only {signatories_created} signatories created - need 2!")
            raise HTTPException(
                status_code=400, 
                detail=f"Failed to create both signatories. Only {signatories_created} created. Please ensure counterparty is assigned."
            )
        
        # STEP 7: Create audit log
        audit_log = text("""
            INSERT INTO audit_logs
            (user_id, contract_id, action_type, action_details, ip_address, created_at)
            VALUES (:user_id, :contract_id, 'sent_for_signature', 
                    :action_details, '127.0.0.1', NOW())
        """)
        
        db.execute(audit_log, {
            "user_id": current_user.id,
            "contract_id": contract_id,
            "action_details": json.dumps({
                "signatories_created": signatories_created,
                "contract_number": contract.contract_number
            })
        })
        
        db.commit()
        
        logger.info(f" Contract {contract_id} sent for signature with {signatories_created} signatories")
        

        log_contract_action(
            db=db,
            action_type="signature_requested",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "signatories_count": signatories_created,
                "new_status": "signature"
            },
            ip_address=None
        )



        return {
            "success": True,
            "message": f"Contract sent for signature successfully. {signatories_created} signatories notified.",
            "contract_id": contract_id,
            "signatories_count": signatories_created
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error sending for signature: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))



@router.post("/initiate-approval-workflow")
async def initiate_approval_workflow(
    data: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Initiate approval workflow for contract"""
    try:
        contract_id = data.get("contract_id")
        
        # Get contract
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Update contract status to 'approval'
        contract.status = 'approval'
        contract.updated_at = datetime.now()
        
        logger.info(f" Contract {contract_id} status updated to 'approval'")
        
        # ⚡ PRIORITY CHECK: Look for CUSTOM workflow first, then MASTER
        activate_workflow_query = text("""
            UPDATE workflow_instances wi
            INNER JOIN workflows w ON wi.workflow_id = w.id
            SET wi.status = 'in_progress',
                wi.started_at = NOW()
            WHERE wi.contract_id = :contract_id
            AND w.company_id = :company_id
            ORDER BY w.is_master ASC
            LIMIT 1
        """)
        
        result = db.execute(activate_workflow_query, {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        })
        
        if result.rowcount == 0:
            raise HTTPException(
                status_code=404, 
                detail="No workflow configured for this contract. Please setup a workflow first."
            )
        
        logger.info(f" Workflow initiated for contract {contract_id}")
        
        # Create activity log
        try:
            activity_query = text("""
                INSERT INTO contract_activity 
                (contract_id, action_type, action_by, notes, timestamp)
                VALUES 
                (:contract_id, 'workflow_initiated', :user_id, 'Approval workflow initiated', NOW())
            """)
            db.execute(activity_query, {
                "contract_id": contract_id,
                "user_id": current_user.id
            })
        except Exception as activity_err:
            logger.warning(f"Could not create activity log: {str(activity_err)}")
        
        db.commit()


        log_contract_action(
            db=db,
            action_type="workflow_started",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "workflow_type": "approval",
                "new_status": "approval"
            },
            ip_address=None
        )
        
        return {
            "success": True,
            "message": "Approval workflow initiated successfully",
            "contract_id": contract_id,
            "contract_status": "approval"
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error initiating approval workflow: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/signature/apply")
async def apply_signature(
    data: dict,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Apply signature to contract
    FIXED: Using correct signatories table columns (external_email, no company_id)
    """
    try:
        contract_id = data.get("contract_id")
        signer_type = data.get("signer_type")
        signature_method = data.get("signature_method", "draw")
        signature_data = data.get("signature_data")
        
        logger.info(f"📝 Applying signature: contract_id={contract_id}, signer_type={signer_type}")
        
        # STEP 1: Verify contract exists and is in signature status
        contract_query = text("""
            SELECT c.id, c.contract_number, c.status, c.company_id
            FROM contracts c
            WHERE c.id = :contract_id
        """)
        
        contract = db.execute(contract_query, {
            "contract_id": contract_id,
        }).fetchone()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        logger.info(f" Contract: {contract.contract_number} - Status: {contract.status}")
        
        if contract.status != 'signature':
            return {
                "success": False,
                "detail": f"Contract must be in signature status. Current status: {contract.status}",
                "not_authorized": True
            }
        
        # STEP 2: Check if already signed
        check_existing = text("""
            SELECT id, has_signed, signed_at, signature_data
            FROM signatories
            WHERE contract_id = :contract_id 
            AND user_id = :user_id
            LIMIT 1
        """)
        
        existing = db.execute(check_existing, {
            "contract_id": contract_id,
            "user_id": current_user.id
        }).fetchone()
        
        if existing and existing.has_signed:
            return {
                "success": False,
                "detail": "You have already signed this contract",
                "already_signed": True
            }
        
        # STEP 3: Apply signature
        client_ip = request.client.host if request.client else "127.0.0.1"
        user_full_name = f"{current_user.first_name} {current_user.last_name}".strip()
        
        if existing:
            # UPDATE existing signatory record
            #  FIXED: Using only columns that exist in signatories table
            update_signatory = text("""
                UPDATE signatories
                SET has_signed = 1,
                    signed_at = NOW(),
                    signature_data = :signature_data,
                    signature_method = :signature_method,
                    ip_address = :ip_address
                WHERE contract_id = :contract_id 
                AND user_id = :user_id
            """)
            
            db.execute(update_signatory, {
                "signature_data": signature_data,
                "signature_method": signature_method,
                "ip_address": client_ip,
                "contract_id": contract_id,
                "user_id": current_user.id
            })
            logger.info(f" Updated existing signatory record for user {current_user.id}")
        else:
            # INSERT new signatory record
            #  FIXED: Using external_email instead of email, removed company_id
            insert_signatory = text("""
                INSERT INTO signatories
                (contract_id, user_id, signer_type, 
                 has_signed, signed_at, signature_data, signature_method, 
                 ip_address, external_email, signing_order)
                VALUES (:contract_id, :user_id, :signer_type,
                        1, NOW(), :signature_data, :signature_method,
                        :ip_address, :external_email, :signing_order)
            """)
            
            # Determine signing order (1 for first signer, 2 for second)
            count_query = text("SELECT COUNT(*) as count FROM signatories WHERE contract_id = :contract_id")
            count_result = db.execute(count_query, {"contract_id": contract_id}).fetchone()
            signing_order = (count_result.count if count_result else 0) + 1
            
            db.execute(insert_signatory, {
                "contract_id": contract_id,
                "user_id": current_user.id,
                "signer_type": signer_type,
                "signature_data": signature_data,
                "signature_method": signature_method,
                "ip_address": client_ip,
                "external_email": current_user.email,
                "signing_order": signing_order
            })
            logger.info(f" Created new signatory record for user {current_user.id}")
        
        # STEP 4: Check if all required signatures collected (BEFORE log_contract_action)
        all_signatures_query = text("""
            SELECT 
                COUNT(*) as total_signatories,
                SUM(CASE WHEN has_signed = 1 THEN 1 ELSE 0 END) as signed_count
            FROM signatories
            WHERE contract_id = :contract_id
        """)
        
        signature_status = db.execute(all_signatures_query, {
            "contract_id": contract_id
        }).fetchone()
        
        #  FIX: Initialize all_signed BEFORE using it in log_contract_action
        all_signed = (signature_status.signed_count >= 2) #signature_status.total_signatories
        
        logger.info(f"📊 Signature status: {signature_status.signed_count}/{signature_status.total_signatories}")
        
        # STEP 5: Log the contract action with all_signed properly defined
        log_contract_action(
            db=db,
            action_type="contract_signed",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "signer_type": signer_type,
                "signature_method": signature_method,
                "all_signed": all_signed
            },
            ip_address=client_ip
        )
        
        # STEP 6: Update contract status if all signed
        new_contract_status = "signature"
        if all_signed:
            update_contract = text("""
                UPDATE contracts
                SET status = 'signed',
                    signed_date = NOW(),
                    updated_at = NOW()
                WHERE id = :contract_id
            """)
            
            db.execute(update_contract, {"contract_id": contract_id})
            new_contract_status = "signed"
            logger.info(f"🎉 All parties have signed! Contract {contract_id} status updated to 'signed'")


            # Get party details (if not already in signature_status)
            party_details_query = text("""
                SELECT 
                    c.contract_number,
                    c.contract_title,
                    CONCAT(u1.first_name, ' ', u1.last_name) as initiator_name,
                    u1.email as initiator_email,
                    CONCAT(u2.first_name, ' ', u2.last_name) as party_b_name,
                    u2.email as party_b_email
                FROM contracts c
                LEFT JOIN users u1 ON c.created_by = u1.id
                LEFT JOIN users u2 ON c.party_b_lead_id = u2.id
                WHERE c.id = :contract_id
            """)
            
            details = db.execute(party_details_query, {"contract_id": contract_id}).fetchone()
            
            if details:
                # Email to Initiator
                if details.initiator_email:
                    WorkflowEmailService.send_all_parties_signed_notification(
                        db=db,
                        contract_id=contract_id,
                        contract_number=details.contract_number,
                        contract_title=details.contract_title,
                        recipient_email=details.initiator_email,
                        recipient_name=details.initiator_name,
                        recipient_role="Contract Initiator"
                    )
                
                # Email to Party B Lead
                if details.party_b_email:
                    WorkflowEmailService.send_all_parties_signed_notification(
                        db=db,
                        contract_id=contract_id,
                        contract_number=details.contract_number,
                        contract_title=details.contract_title,
                        recipient_email=details.party_b_email,
                        recipient_name=details.party_b_name,
                        recipient_role="Counter-Party Lead"
                    )
                
                logger.info("✅ All parties signed notifications sent")

        
        db.commit()

        
        # STEP 7: Return success response
        return {
            "success": True,
            "message": f"Signature applied successfully by {user_full_name}",
            "signer_name": user_full_name,
            "signed_at": datetime.now().isoformat(),
            "all_signed": all_signed,
            "contract_status": new_contract_status,
            "signature_method": signature_method
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error applying signature: {str(e)}")
        logger.error(f"Traceback (most recent call last):\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))
        

@router.get("/get-contract-with-certificate/{contract_id}")
async def get_contract_with_certificate(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get contract with signature certificate data
    FIXED: Returns actual signature data for display
    """
    try:
        # Get signatories with signatures
        signatories_query = text("""
            SELECT 
                s.id,
                s.signer_type,
                s.has_signed,
                s.signed_at,
                s.signature_data,
                s.signature_method,
                s.ip_address,
                s.email,
                s.signing_order,
                u.first_name,
                u.last_name,
                u.email as user_email,
                c.company_name
            FROM signatories s
            LEFT JOIN users u ON s.user_id = u.id
            LEFT JOIN companies c ON s.company_id = c.id
            WHERE s.contract_id = :contract_id
            ORDER BY s.signing_order
        """)
        
        signatories = db.execute(signatories_query, {
            "contract_id": contract_id
        }).fetchall()
        
        certificate_data = {
            "contract_id": contract_id,
            "signatories": []
        }
        
        for sig in signatories:
            signer_name = "Pending"
            if sig.has_signed and sig.first_name and sig.last_name:
                signer_name = f"{sig.first_name} {sig.last_name}"
            elif sig.first_name:
                signer_name = sig.first_name
            
            certificate_data["signatories"].append({
                "signer_type": sig.signer_type,
                "name": signer_name,
                "company_name": sig.company_name or "",
                "email": sig.user_email or sig.email or "",
                "has_signed": bool(sig.has_signed),
                "signed_at": sig.signed_at.isoformat() if sig.signed_at else None,
                "signature_data": sig.signature_data or "",
                "signature_method": sig.signature_method or "draw",
                "ip_address": sig.ip_address or "",
                "signing_order": sig.signing_order
            })
        
        logger.info(f" Retrieved certificate with {len(certificate_data['signatories'])} signatories")
        
        return {
            "success": True,
            "certificate": certificate_data
        }
        
    except Exception as e:
        logger.error(f"Error retrieving certificate: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# FIXED: Check E-SIGN authority endpoint
@router.get("/check-esign-authority/{contract_id}")
async def check_esign_authority(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Check if current user has E-SIGN authority for a contract
    FIXED: Uses correct database schema
    """
    try:
        # FIXED: Direct check using assignee_user_id in workflow_steps
        authority_query = text("""
            SELECT 
                ws.id as step_id,
                ws.step_name,
                ws.step_type,
                w.workflow_name,
                w.is_master,
                wi.status as workflow_status
            FROM workflow_instances wi
            INNER JOIN workflows w ON wi.workflow_id = w.id
            INNER JOIN workflow_steps ws ON w.id = ws.workflow_id
            WHERE wi.contract_id = :contract_id
            AND ws.step_type ='e_sign_authority'
            LIMIT 1
        """)
        
        authority = db.execute(authority_query, {
            "contract_id": contract_id,
        }).fetchone()
        
        has_authority = authority is not None
        
        return {
            "success": True,
            "has_esign_authority": has_authority,
            "workflow_info": {
                "workflow_name": authority.workflow_name if authority else None,
                "is_master": bool(authority.is_master) if authority else None,
                "step_name": authority.step_name if authority else None,
                "workflow_status": authority.workflow_status if authority else None
            } if has_authority else None
        }
        
    except Exception as e:
        logger.error(f"Error checking E-SIGN authority: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/get-workflow-options/{contract_id}")
async def get_workflow_options(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get available workflow options for signature
    FIXED: Uses correct database schema
    """
    try:
        # Get master workflow - FIXED: Using correct joins
        master_query = text("""
            SELECT 
                w.id,
                w.workflow_name,
                w.description,
                COUNT(ws.id) as total_steps,
                GROUP_CONCAT(
                    DISTINCT CONCAT(ws.step_name, ':', 
                    COALESCE(CONCAT(u.first_name, ' ', u.last_name), 'Unassigned'))
                    ORDER BY ws.step_number
                    SEPARATOR ' → '
                ) as workflow_path
            FROM workflows w
            LEFT JOIN workflow_steps ws ON w.id = ws.workflow_id
            LEFT JOIN users u ON ws.assignee_user_id = u.id
            WHERE w.company_id = :company_id
            AND w.is_master = 1
            AND w.is_active = 1
            GROUP BY w.id
        """)
        
        master_workflow = db.execute(master_query, {
            "company_id": current_user.company_id
        }).fetchone()
        
        # Get E-SIGN users from master workflow if exists
        esign_users = []
        if master_workflow:
            esign_query = text("""
                SELECT 
                    CONCAT(u.first_name, ' ', u.last_name) as name,
                    u.email,
                    ws.step_name as role,
                    ws.department
                FROM workflow_steps ws
                JOIN users u ON ws.assignee_user_id = u.id
                WHERE ws.workflow_id = :workflow_id
                AND ws.step_type IN ('e-sign', 'e-signature', 'e_sign', 'esign')
            """)
            
            esign_result = db.execute(esign_query, {
                "workflow_id": master_workflow.id
            }).fetchall()
            
            esign_users = [
                {
                    "name": user.name,
                    "email": user.email,
                    "role": user.role,
                    "department": user.department
                }
                for user in esign_result
            ]
        
        # Get available users for custom workflow
        users_query = text("""
            SELECT 
                u.id,
                CONCAT(u.first_name, ' ', u.last_name) as name,
                u.email,
                u.role,
                u.department
            FROM users u
            WHERE u.company_id = :company_id
            AND u.is_active = 1
            ORDER BY u.first_name, u.last_name
        """)
        
        available_users = db.execute(users_query, {
            "company_id": current_user.company_id
        }).fetchall()
        
        return {
            "success": True,
            "master_workflow": {
                "id": master_workflow.id if master_workflow else None,
                "name": master_workflow.workflow_name if master_workflow else "Not Configured",
                "description": master_workflow.description if master_workflow else None,
                "total_steps": master_workflow.total_steps if master_workflow else 0,
                "workflow_path": master_workflow.workflow_path if master_workflow else "",
                "esign_users": esign_users
            } if master_workflow else None,
            "available_users": [
                {
                    "id": user.id,
                    "name": user.name,
                    "email": user.email,
                    "role": user.role,
                    "department": user.department
                }
                for user in available_users
            ]
        }
        
    except Exception as e:
        logger.error(f"Error getting workflow options: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/execute-contract")
async def execute_contract(
    execution_data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Execute contract - FIXED with correct User attributes
    """
    try:
        contract_id = int(execution_data.get("contract_id"))
        
        logger.info(f"🔄 Executing contract {contract_id}")
        
        # Verify contract
        contract_check = text("""
            SELECT id, status, contract_number, contract_title, signed_date
            FROM contracts WHERE id = :contract_id
        """)
        
        contract = db.execute(contract_check, {"contract_id": contract_id}).fetchone()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        if contract.status != 'signed':
            raise HTTPException(status_code=400, detail=f"Contract must be signed. Current: {contract.status}")
        
        # Verify both parties signed
        sig_check = text("""
            SELECT 
                SUM(CASE WHEN signer_type = 'client' THEN 1 ELSE 0 END) as c,
                SUM(CASE WHEN signer_type IN ('provider', 'company') THEN 1 ELSE 0 END) as p
            FROM signatories WHERE contract_id = :contract_id AND has_signed = 1
        """)
        
        sig = db.execute(sig_check, {"contract_id": contract_id}).fetchone()
        
        if not sig or sig.c == 0 or sig.p == 0:
            raise HTTPException(status_code=400, detail="Both parties must sign first")
        
        # Execute contract
        execute_query = text("""
            UPDATE contracts
            SET status = 'executed', workflow_status = 'completed',
                effective_date = CURDATE(), updated_at = NOW()
            WHERE id = :contract_id
        """)
        
        db.execute(execute_query, {"contract_id": contract_id})
        
        # FIXED: Use first_name + last_name instead of full_name
        executed_by_name = f"{current_user.first_name} {current_user.last_name}"
        
        # Generate certificate
        certificate_data = {
            "contract_id": contract_id,
            "contract_number": contract.contract_number,
            "contract_title": contract.contract_title,
            "execution_date": datetime.now().isoformat(),
            "signed_date": contract.signed_date.isoformat() if contract.signed_date else None,
            "executed_by": executed_by_name,
            "executed_by_email": current_user.email,
            "signatories": []
        }
        
        # Get signatories with proper name handling
        sig_query = text("""
            SELECT 
                s.signer_type, 
                s.signed_at,
                s.signature_data,
                s.ip_address, 
                u.first_name,
                u.last_name,
                u.email
            FROM signatories s
            LEFT JOIN users u ON s.user_id = u.id
            WHERE s.contract_id = :contract_id AND s.has_signed = 1
            ORDER BY s.signing_order
        """)
        
        sigs = db.execute(sig_query, {"contract_id": contract_id}).fetchall()
        
        for s in sigs:
            # FIXED: Construct full name from first_name and last_name
            signer_name = "External Signer"
            if s.first_name and s.last_name:
                signer_name = f"{s.first_name} {s.last_name}"
            elif s.first_name:
                signer_name = s.first_name
            
            certificate_data["signatories"].append({
                "signer_type": s.signer_type,
                "name": signer_name,
                "email": s.email or "",
                "signed_at": s.signed_at.isoformat() if s.signed_at else None,
                "signature_data": s.signature_data or "",
                "ip_address": s.ip_address or ""
            })
        
        # Store certificate (optional - if table exists)
        try:
            cert_query = text("""
                INSERT INTO contract_metadata
                (contract_id, metadata_key, metadata_value, created_at)
                VALUES (:contract_id, 'execution_certificate', :cert_data, NOW())
            """)
            
            db.execute(cert_query, {
                "contract_id": contract_id,
                "cert_data": json.dumps(certificate_data)
            })
        except Exception as meta_error:
            logger.warning(f"⚠️ Could not store certificate metadata: {str(meta_error)}")
            # Continue - certificate is in response anyway
        
        # Audit log with JSON
        audit_log = text("""
            INSERT INTO audit_logs 
            (contract_id, user_id, action_type, action_details, created_at)
            VALUES (:contract_id, :user_id, :action_type, :action_details, NOW())
        """)
        
        action_details_json = json.dumps({
            "event": "contract_executed",
            "contract_number": contract.contract_number,
            "message": "Contract officially executed and active",
            "executed_by": executed_by_name,
            "executed_by_email": current_user.email,
            "execution_date": datetime.now().isoformat()
        })
        
        db.execute(audit_log, {
            "contract_id": contract_id,
            "user_id": int(current_user.id),
            "action_type": "contract_executed",
            "action_details": action_details_json
        })
        
        db.commit()
        
        # =====================================================
        # 📧 SEND EMAIL NOTIFICATIONS - CONTRACT EXECUTED
        # =====================================================
        try:
            from app.services.workflow_email_service import WorkflowEmailService
            
            # Get initiator and counter-party details
            party_details_query = text("""
                SELECT 
                    c.created_by,
                    c.party_b_lead_id,
                    CONCAT(u1.first_name, ' ', u1.last_name) as initiator_name,
                    u1.email as initiator_email,
                    CONCAT(u2.first_name, ' ', u2.last_name) as party_b_name,
                    u2.email as party_b_email
                FROM contracts c
                LEFT JOIN users u1 ON c.created_by = u1.id
                LEFT JOIN users u2 ON c.party_b_lead_id = u2.id
                WHERE c.id = :contract_id
            """)
            
            party_details = db.execute(party_details_query, {"contract_id": contract_id}).fetchone()
            
            if party_details:
                # Send to Initiator
                if party_details.initiator_email:
                    WorkflowEmailService.send_contract_executed_notification(
                        db=db,
                        contract_id=contract_id,
                        contract_number=contract.contract_number,
                        contract_title=contract.contract_title,
                        recipient_email=party_details.initiator_email,
                        recipient_name=party_details.initiator_name,
                        recipient_role="Contract Initiator",
                        executed_by_name=executed_by_name
                    )
                    logger.info(f"✉️ Execution email sent to Initiator: {party_details.initiator_email}")
                
                # Send to Counter-Party Lead
                if party_details.party_b_email:
                    WorkflowEmailService.send_contract_executed_notification(
                        db=db,
                        contract_id=contract_id,
                        contract_number=contract.contract_number,
                        contract_title=contract.contract_title,
                        recipient_email=party_details.party_b_email,
                        recipient_name=party_details.party_b_name,
                        recipient_role="Counter-Party Lead",
                        executed_by_name=executed_by_name
                    )
                    logger.info(f"✉️ Execution email sent to Counter-Party: {party_details.party_b_email}")
                
                logger.info("✅ Contract execution notifications sent to both parties")
            else:
                logger.warning("⚠️ Could not find party details for email notifications")
                
        except Exception as email_error:
            logger.error(f"❌ Error sending execution emails: {str(email_error)}")
            # Don't fail the execution if email fails
        
        logger.info(f"🎉 Contract {contract_id} executed successfully!")
        
        return {
            "success": True,
            "message": "Contract executed successfully!",
            "contract_id": contract_id,
            "status": "executed",
            "execution_date": datetime.now().isoformat(),
            "certificate_data": certificate_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error executing contract: {str(e)}")
        logger.exception(e)
        raise HTTPException(status_code=500, detail=str(e))
        

@router.get("/execution-certificate/{contract_id}")
async def get_execution_certificate(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Retrieve Certificate of Completion - FIXED with proper name handling
    """
    try:
        # Try to get from metadata table
        try:
            certificate_query = text("""
                SELECT metadata_value, created_at
                FROM contract_metadata
                WHERE contract_id = :contract_id
                AND metadata_key = 'execution_certificate'
                ORDER BY created_at DESC LIMIT 1
            """)
            
            result = db.execute(certificate_query, {"contract_id": contract_id}).fetchone()
            
            if result:
                certificate_data = json.loads(result.metadata_value)
                certificate_data["generated_at"] = result.created_at.isoformat() if result.created_at else None
                
                return {
                    "success": True,
                    "certificate": certificate_data
                }
        except Exception as e:
            logger.info(f"Certificate metadata not found, generating from current data: {str(e)}")
        
        # Fallback: Generate from current data
        contract_query = text("""
            SELECT id, contract_number, contract_title, signed_date, effective_date
            FROM contracts WHERE id = :contract_id
        """)
        
        contract = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # FIXED: Get user with proper name handling
        executed_by_name = f"{current_user.first_name} {current_user.last_name}"
        
        # Get signatories with proper name handling
        signatories_query = text("""
            SELECT 
                s.signer_type,
                s.signed_at,
                s.ip_address,
                s.signature_data,
                u.first_name,
                u.last_name,
                u.email
            FROM signatories s
            LEFT JOIN users u ON s.user_id = u.id
            WHERE s.contract_id = :contract_id AND s.has_signed = 1
            ORDER BY s.signing_order
        """)
        
        signatories = db.execute(signatories_query, {"contract_id": contract_id}).fetchall()
        
        certificate_data = {
            "contract_id": contract_id,
            "contract_number": contract.contract_number,
            "contract_title": contract.contract_title,
            "execution_date": contract.effective_date.isoformat() if contract.effective_date else None,
            "signed_date": contract.signed_date.isoformat() if contract.signed_date else None,
            "executed_by": executed_by_name,
            "executed_by_email": current_user.email,
            "signatories": []
        }
        
        for sig in signatories:
            # FIXED: Construct full name
            signer_name = "External Signer"
            if sig.first_name and sig.last_name:
                signer_name = f"{sig.first_name} {sig.last_name}"
            elif sig.first_name:
                signer_name = sig.first_name
            
            certificate_data["signatories"].append({
                "signer_type": sig.signer_type,
                "name": signer_name,
                "email": sig.email or "",
                "signed_at": sig.signed_at.isoformat() if sig.signed_at else None,
                "signature_data": sig.signature_data or "",
                "ip_address": sig.ip_address or ""
            })
        
        return {
            "success": True,
            "certificate": certificate_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrieving certificate: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


import json
from sqlalchemy import text


@router.post("/workflow/setup")
async def setup_contract_workflow(
    workflow_data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Setup workflow for a specific contract"""
    try:
        #  EXTRACT ALL VARIABLES FIRST
        contract_id = workflow_data.get("contract_id")
        workflow_type = workflow_data.get("workflow_type", "custom")
        steps = workflow_data.get("steps", [])
        
        logger.info(f"Setting up {workflow_type} workflow for contract {contract_id}")
        logger.info(f"Company ID: {current_user.company_id}")
        logger.info(f"Received {len(steps)} workflow steps")
        
        if workflow_type == "master":
            #  Use master workflow - ALREADY FILTERED BY COMPANY
            logger.info("Using master workflow")
            
            # Get company's master workflow
            master_query = text("""
                SELECT id FROM workflows 
                WHERE company_id = :company_id 
                AND is_master = 1 
                AND is_active = 1
                ORDER BY created_at DESC
                LIMIT 1
            """)
            
            master_workflow = db.execute(master_query, {
                "company_id": current_user.company_id
            }).fetchone()
            
            if master_workflow:
                # Create workflow instance pointing to master workflow
                instance_query = text("""
                    INSERT INTO workflow_instances 
                    (workflow_id, contract_id, current_step, status, started_at)
                    VALUES (:workflow_id, :contract_id, 1, 'pending', NOW())
                """)
                
                db.execute(instance_query, {
                    "workflow_id": master_workflow.id,
                    "contract_id": contract_id
                })
                
                logger.info(f"Created workflow instance using master workflow {master_workflow.id}")
            else:
                logger.warning("No master workflow found")
                raise HTTPException(status_code=404, detail="Master workflow not found for your company")
                
        else:
            #  CUSTOM WORKFLOW - FIXED VERSION
            logger.info("Creating custom workflow")
            
            # Create workflow record
            workflow_insert = text("""
                INSERT INTO workflows
                (company_id, workflow_name, workflow_type, is_master, is_active, created_at, updated_at)
                VALUES (:company_id, :workflow_name, 'contract_approval', 0, 1, NOW(), NOW())
            """)
            
            result = db.execute(workflow_insert, {
                "company_id": current_user.company_id,
                "workflow_name": f"Custom Workflow"
            })
            
            workflow_id = result.lastrowid
            logger.info(f" Created workflow with ID: {workflow_id} for company {current_user.company_id}")
            
            # Store departments mapping
            departments_map = {}
            
            # Create workflow steps WITH user lookup
            if steps:
                logger.info(f"Processing {len(steps)} workflow steps")
                
                for step in steps:
                    step_order = step.get("step_order", 1)
                    step_label = step.get("step_label", "Review")
                    assigned_email = step.get("assigned_email", "")
                    department = step.get("department", "")
                    
                    logger.info(f" Step {step_order}: role={step_label}, email={assigned_email}, dept={department}")
                    
                    # Store department in mapping
                    if department:
                        departments_map[str(step_order)] = department
                    
                    #  Look up user by email - FILTERED BY COMPANY
                    assignee_user_id = None
                    if assigned_email:
                        user_query = text("""
                            SELECT id FROM users 
                            WHERE email = :email 
                            AND company_id = :company_id
                            AND is_active = 1
                            LIMIT 1
                        """)
                        user_result = db.execute(user_query, {
                            "email": assigned_email,
                            "company_id": current_user.company_id
                        }).fetchone()
                        
                        if user_result:
                            assignee_user_id = user_result.id
                            logger.info(f" Found user ID {assignee_user_id} for email {assigned_email}")
                        else:
                            logger.warning(f" User not found in company for email: {assigned_email}")
                    
                    #  Insert step with assignee_user_id AND department
                    step_insert = text("""
                        INSERT INTO workflow_steps
                        (workflow_id, step_number, step_name, step_type, assignee_role, assignee_user_id, department, sla_hours, is_mandatory, created_at)
                        VALUES (:workflow_id, :step_number, :step_name, :step_type, :assignee_role, :assignee_user_id, :department, 24, 1, NOW())
                    """)
                    
                    db.execute(step_insert, {
                        "workflow_id": workflow_id,
                        "step_number": step_order,
                        "step_name": step_label,
                        "step_type": step_label.lower(),
                        "assignee_role": step_label,
                        "assignee_user_id": assignee_user_id,
                        "department": department
                    })
                    
                    logger.info(f" Inserted step {step_order} with user_id={assignee_user_id}, department={department}")
            else:
                logger.warning("No workflow steps provided")
            
            #  Update workflow with departments JSON
            if departments_map:
                update_workflow = text("""
                    UPDATE workflows 
                    SET workflow_json = :workflow_json
                    WHERE id = :workflow_id
                """)
                
                workflow_config = json.dumps({"departments": departments_map})
                db.execute(update_workflow, {
                    "workflow_id": workflow_id,
                    "workflow_json": workflow_config
                })
                
                logger.info(f" Stored departments mapping: {departments_map}")
            
            # Create workflow instance
            instance_query = text("""
                INSERT INTO workflow_instances 
                (workflow_id, contract_id, current_step, status, started_at)
                VALUES (:workflow_id, :contract_id, 1, 'pending', NOW())
            """)
            
            db.execute(instance_query, {
                "workflow_id": workflow_id,
                "contract_id": contract_id
            })
            
            logger.info(f" Created workflow instance for contract {contract_id}")
        
        db.commit()
        logger.info("🎉 Workflow setup completed successfully")
        
        return {"success": True, "message": "Workflow configured successfully"}
        
    except HTTPException:
        # Re-raise HTTP exceptions
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        logger.error(f" Error setting up workflow: {str(e)}")
        logger.error(f"Traceback: ", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/workflow/{contract_id}")
async def get_contract_workflow(
    contract_id: int,
    contract_type: str = None,  # Optional parameter: 'custom', 'master', or None
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get workflow configuration for a specific contract"""
    try:
        logger.info(f"Fetching workflow for contract {contract_id} - Company: {current_user.company_id}, Type: {contract_type}")
        
        #  First verify the contract belongs to user's company
        contract_check = text("""
            SELECT id FROM contracts 
            WHERE id = :contract_id 
        """)
        
        contract_exists = db.execute(contract_check, {
            "contract_id": contract_id,
        }).fetchone()
        
        if not contract_exists:
            logger.warning(f"Contract {contract_id} not found for company {current_user.company_id}")
            raise HTTPException(status_code=404, detail="Contract not found or access denied")
        
        # Build the is_master filter based on contract_type
        is_master_filter = ""
        if contract_type == "custom":
            is_master_filter = "AND w.is_master = 0"
            logger.info(f"Filtering for custom workflows (is_master=0)")
        elif contract_type == "master":
            is_master_filter = "AND w.is_master = 1"
            logger.info(f"Filtering for master workflows (is_master=1)")
        
        #  Get workflow instance - FILTERED BY COMPANY and optionally by is_master
        instance_query = text(f"""
            SELECT 
                wi.id as instance_id,
                wi.workflow_id,
                wi.status,
                w.workflow_name,
                w.is_master,
                w.workflow_type,
                w.workflow_json
            FROM workflow_instances wi
            LEFT JOIN workflows w ON wi.workflow_id = w.id
            WHERE wi.contract_id = :contract_id
            AND w.company_id = :company_id
            {is_master_filter}
            ORDER BY wi.started_at DESC
            LIMIT 1
        """)
        
        workflow_instance = db.execute(instance_query, {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        }).fetchone()
        
        if not workflow_instance:
            logger.warning(f"No workflow found for contract {contract_id} in company {current_user.company_id} with type '{contract_type}'")
            return {
                "success": False,
                "message": f"No {'custom' if contract_type == 'custom' else 'master' if contract_type == 'master' else ''} workflow configured for this contract"
            }
        
        # Parse department mapping from workflow_json
        departments_map = {}
        if workflow_instance.workflow_json:
            try:
                if isinstance(workflow_instance.workflow_json, str):
                    config = json.loads(workflow_instance.workflow_json)
                else:
                    config = workflow_instance.workflow_json
                departments_map = config.get('departments', {})
                logger.info(f"Department mapping loaded: {departments_map}")
            except Exception as e:
                logger.error(f"Error parsing workflow_json: {e}")
        
        #  Get workflow steps with user information - FILTERED BY COMPANY
        steps_query = text("""
            SELECT 
                ws.id,
                ws.step_number,
                ws.step_name,
                ws.step_type,
                ws.assignee_role,
                ws.department,
                ws.assignee_user_id,
                ws.sla_hours,
                ws.is_mandatory,
                u.email as assignee_email,
                u.first_name,
                u.last_name
            FROM workflow_steps ws
            LEFT JOIN users u ON ws.assignee_user_id = u.id 
                AND u.company_id = :company_id
                AND u.is_active = 1
            WHERE ws.workflow_id = :workflow_id
            ORDER BY ws.step_number ASC, ws.id ASC
        """)
        
        steps = db.execute(steps_query, {
            "workflow_id": workflow_instance.workflow_id,
            "company_id": current_user.company_id
        }).fetchall()
        
        logger.info(f"Found {len(steps)} workflow step entries for company {current_user.company_id}")
        
        # Group steps by step_number and collect users
        steps_map = {}
        for step in steps:
            step_num = step.step_number
            
            # Get department from mapping or from step directly
            dept = step.department if step.department else ''
            
            # Override with mapping if exists
            if str(step_num) in departments_map:
                dept = departments_map[str(step_num)]
            elif step_num in departments_map:
                dept = departments_map[step_num]
            
            if step_num not in steps_map:
                steps_map[step_num] = {
                    "id": step.id,
                    "step_number": step_num,
                    "step_name": step.step_name,
                    "step_type": step.step_type,
                    "assignee_role": step.assignee_role,
                    "department": dept,
                    "sla_hours": step.sla_hours,
                    "is_mandatory": bool(step.is_mandatory) if step.is_mandatory is not None else True,
                    "users": []
                }
            
            # Add user if exists and belongs to company
            if step.assignee_user_id and step.assignee_email:
                user_exists = any(u['id'] == step.assignee_user_id for u in steps_map[step_num]['users'])
                if not user_exists:
                    steps_map[step_num]['users'].append({
                        "id": step.assignee_user_id,
                        "name": f"{step.first_name} {step.last_name}" if step.first_name else step.assignee_email,
                        "email": step.assignee_email
                    })
        
        # Convert to sorted list
        steps_list = [steps_map[k] for k in sorted(steps_map.keys())]
        
        logger.info(f"Returning {len(steps_list)} workflow steps with department info")
        
        db.commit()
        
        return {
            "success": True,
            "workflow": {
                "id": workflow_instance.workflow_id,
                "workflow_name": workflow_instance.workflow_name,
                "is_master": bool(workflow_instance.is_master),
                "workflow_type": workflow_instance.workflow_type,
                "status": workflow_instance.status,
                "steps": steps_list
            }
        }
        
    except HTTPException:
        # Re-raise HTTP exceptions
        db.rollback()
        raise
    except Exception as e:
        logger.error(f"Error retrieving workflow: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/submit-approval")
async def submit_for_internal_review(
    review_data: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Submit contract for approval or Signature"""
    try:
        contract_id = review_data.get("contract_id")
        review_type = review_data.get("review_type")
        personnel_emails = review_data.get("personnel_emails", [])
        notes = review_data.get("notes", "")
        
        request_type = review_data.get("request_type")


        # Get contract
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
            

        # Check if current user is Party B (counterparty) for this contract
        contract_query = text("""
            SELECT party_b_id FROM contracts WHERE id = :contract_id LIMIT 1
        """)
        contract_result = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
        
        if not contract_result:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        party_b_id = contract_result.party_b_id if hasattr(contract_result, 'party_b_id') else contract_result[0]

        # Handle masterWorkflow and customWorkflow review types
        if review_type == 'masterWorkflow':
            # Deactivate all custom workflows (is_master=0) for this contract and company
            deactivate_custom_query = text("""
                UPDATE workflows 
                SET is_active = 0
                WHERE id IN (
                    SELECT DISTINCT workflow_id 
                    FROM workflow_instances 
                    WHERE contract_id = :contract_id
                )
                AND company_id = :company_id
                AND is_master = 0
            """)
            db.execute(deactivate_custom_query, {
                "contract_id": contract_id,
                "company_id": current_user.company_id
            })
            logger.info(f"Custom workflows deactivated for contract {contract_id} (masterWorkflow selected)")
            
            # Check if master workflow instances exist for this contract
            check_master_instances = text("""
                SELECT COUNT(*) as count
                FROM workflow_instances wi
                INNER JOIN workflows w ON wi.workflow_id = w.id
                WHERE wi.contract_id = :contract_id
                AND w.company_id = :company_id
                AND w.is_master = 1
            """)
            master_count_result = db.execute(check_master_instances, {
                "contract_id": contract_id,
                "company_id": current_user.company_id
            }).fetchone()
            
            master_count = master_count_result[0] if master_count_result else 0
            
            # If no master workflow instances exist, create them
            if master_count == 0:
                # Get all active master workflows for the company
                get_master_workflows = text("""
                    SELECT id, workflow_name
                    FROM workflows
                    WHERE company_id = :company_id
                    AND is_master = 1
                    AND is_active = 1
                """)
                master_workflows = db.execute(get_master_workflows, {
                    "company_id": current_user.company_id
                }).fetchall()
                
                # Create workflow instances for each master workflow
                for workflow in master_workflows:
                    workflow_id = workflow[0]
                    workflow_name = workflow[1]
                    
                    insert_instance = text("""
                        INSERT INTO workflow_instances 
                        (contract_id, workflow_id, status, current_step)
                        VALUES (:contract_id, :workflow_id, 'pending', 1)
                    """)
                    db.execute(insert_instance, {
                        "contract_id": contract_id,
                        "workflow_id": workflow_id
                    })
                    logger.info(f"Created master workflow instance for workflow '{workflow_name}' (ID: {workflow_id}) and contract {contract_id}")
            else:
                logger.info(f"Master workflow instances already exist for contract {contract_id} (count: {master_count})")
            
        elif review_type == 'customWorkflow':
            # Activate all custom workflows (is_master=0) for this contract and company
            activate_custom_query = text("""
                UPDATE workflows 
                SET is_active = 1
                WHERE id IN (
                    SELECT DISTINCT workflow_id 
                    FROM workflow_instances 
                    WHERE contract_id = :contract_id
                )
                AND company_id = :company_id
                AND is_master = 0
            """)
            db.execute(activate_custom_query, {
                "contract_id": contract_id,
                "company_id": current_user.company_id
            })
            logger.info(f"Custom workflows activated for contract {contract_id} (customWorkflow selected)")
        
        # Activate workflow instances for ALL users (including counterparties)
        activate_workflow_query = text("""
            UPDATE workflow_instances 
            SET status = 'active',
                started_at = NOW()
            WHERE contract_id = :contract_id
            AND status IN ('pending', 'in_progress')
        """)
        db.execute(activate_workflow_query, {
            "contract_id": contract_id
        })
        logger.info(f"Workflow instances activated for contract {contract_id}")


        # Check if current user is the counterparty (Party B)
        is_counterparty = party_b_id == current_user.company_id
        
        # Only update contract status if user is NOT Party B (counterparty)
        if not is_counterparty:

            
            # =====================================================
            # HANDLE INTERNAL REVIEW
            # =====================================================
            if request_type == 'internal_review':
                logger.info(f"🔍 Processing INTERNAL REVIEW for contract {contract_id}")

                # Update contract status to 'review'
                update_query = text("""
                    UPDATE contracts 
                    SET status = 'review',
                        updated_at = NOW()
                    WHERE id = :contract_id
                """)
                db.execute(update_query, {"contract_id": contract_id})
                logger.info(f"Contract {contract_id} status updated to 'review' by user {current_user.id}")

                # Update workflow instance status to 'in_progress'
                activate_workflow_query = text("""
                    UPDATE workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    SET wi.status = 'in_progress', 
                        wi.current_step = 1,
                        wi.started_at = NOW()
                    WHERE wi.contract_id = :contract_id
                    AND w.company_id = :company_id
                """)
                db.execute(activate_workflow_query, {
                    "contract_id": contract_id,
                    "company_id": current_user.company_id
                })

                logger.info(f" Internal review workflow status updated to 'in_progress'")

                # Get first reviewer in the workflow for action_person_id
                logger.info(f"🔍 Finding first reviewer in internal review workflow...")
                first_reviewer_query = text("""
                    SELECT u.id
                    FROM workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    INNER JOIN workflow_steps ws ON w.id = ws.workflow_id
                    INNER JOIN users u ON ws.assignee_user_id = u.id
                    WHERE wi.contract_id = :contract_id
                    AND w.company_id = :company_id
                    AND ws.step_number = 1
                    AND w.is_active=1
                    ORDER BY is_master ASC
                    LIMIT 1
                """)
                first_reviewer = db.execute(first_reviewer_query, {
                    "contract_id": contract_id,
                    "company_id": current_user.company_id
                }).first()
                
                first_reviewer_id = first_reviewer.id if first_reviewer else None
                logger.info(f" First reviewer ID: {first_reviewer_id}")
                
                # Update contract with action_person_id
                update_contract_query = text("""
                    UPDATE contracts
                    SET action_person_id = :action_person_id,
                        updated_at = NOW()
                    WHERE id = :contract_id
                """)
                db.execute(update_contract_query, {
                    "action_person_id": first_reviewer_id,
                    "contract_id": contract_id
                })
                
                logger.info(f" Contract {contract_id} action_person_id set to: {first_reviewer_id}")
# 📧 EMAIL NOTIFICATION: Internal review request
                try:
                    logger.info("📧 Preparing to send internal review notification emails...")
                    
                    reviewer_emails = []
                    
                    if review_type == 'specific':
                        # Get email addresses from request
                        reviewer_emails = personnel_emails if personnel_emails else []
                        logger.info(f"📋 Specific personnel review: {len(reviewer_emails)} reviewers")
                        
                    elif review_type == 'masterWorkflow':
                        # Get emails from master workflow users
                        master_users_query = text("""
                            SELECT DISTINCT
                                u.email,
                                CONCAT(u.first_name, ' ', u.last_name) as full_name
                            FROM workflow_steps ws
                            INNER JOIN users u ON ws.assignee_user_id = u.id
                            INNER JOIN workflows w ON ws.workflow_id = w.id
                            WHERE w.is_master = 1
                            AND w.company_id = :company_id
                            AND u.is_active = 1
                            AND u.email IS NOT NULL
                            AND u.id != :current_user_id
                        """)
                        master_users = db.execute(master_users_query, {
                            "company_id": current_user.company_id,
                            "current_user_id": current_user.id
                        }).fetchall()
                        
                        reviewer_emails = [user.email for user in master_users]
                        logger.info(f"📋 Master workflow review: {len(reviewer_emails)} reviewers")
                    
                    if reviewer_emails:
                        # Get contract details
                        contract_query = text("""
                            SELECT contract_number, contract_title
                            FROM contracts WHERE id = :contract_id
                        """)
                        contract_info = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
                        
                        initiator_name = f"{current_user.first_name} {current_user.last_name}"
                        
                        # Send emails to all reviewers
                        WorkflowEmailService.send_internal_review_request(
                            db=db,
                            contract_id=contract_id,
                            contract_number=contract_info.contract_number,
                            contract_title=contract_info.contract_title,
                            reviewer_emails=reviewer_emails,
                            initiator_name=initiator_name
                        )
                        logger.info(f"✉️ Internal review emails sent to {len(reviewer_emails)} reviewers")
                    else:
                        logger.warning("⚠️ No reviewer emails found to send notifications")
                        
                except Exception as email_error:
                    logger.error(f"❌ Error sending internal review emails: {str(email_error)}")


            # =====================================================
            # HANDLE APPROVAL
            # =====================================================
            elif request_type == 'approval':
                logger.info(f"🔍 Processing APPROVAL for contract {contract_id}")
                
                # Get first approver in the workflow for action_person_id
                logger.info(f"🔍 Finding first approver in approval workflow...")
                first_approver_query = text("""
                    SELECT u.id
                    FROM workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    INNER JOIN workflow_steps ws ON w.id = ws.workflow_id
                    INNER JOIN users u ON ws.assignee_user_id = u.id
                    WHERE wi.contract_id = :contract_id
                    AND w.company_id = :company_id
                    AND ws.step_number = 1
                    AND ws.step_type <> 'e_sign_authority'
                    LIMIT 1
                """)
                first_approver = db.execute(first_approver_query, {
                    "contract_id": contract_id,
                    "company_id": current_user.company_id
                }).first()
                
                first_approver_id = first_approver.id if first_approver else None
                logger.info(f" First approver ID: {first_approver_id}")
                
                # Update contract status to 'approval'
                contract.status = 'approval'
                contract.action_person_id = first_approver_id
                contract.updated_at = datetime.now()
                
                logger.info(f" Contract {contract_id} status updated to 'approval'")
                logger.info(f" action_person_id set to: {first_approver_id}")
# 📧 EMAIL NOTIFICATION: Approval workflow initiated
                try:
                    if first_approver_id:
                        logger.info(f"📧 Sending approval workflow notification to first approver ID: {first_approver_id}")
                        
                        # Get first approver details
                        approver_query = text("""
                            SELECT 
                                u.email,
                                CONCAT(u.first_name, ' ', u.last_name) as full_name,
                                ws.step_name,
                                ws.step_type,
                                w.workflow_name
                            FROM users u
                            INNER JOIN workflow_steps ws ON ws.assignee_user_id = u.id
                            INNER JOIN workflows w ON ws.workflow_id = w.id
                            INNER JOIN workflow_instances wi ON wi.workflow_id = w.id
                            WHERE u.id = :user_id
                            AND wi.contract_id = :contract_id
                            AND ws.step_number = 1
                            AND w.company_id = :company_id
                            LIMIT 1
                        """)
                        approver_info = db.execute(approver_query, {
                            "user_id": first_approver_id,
                            "contract_id": contract_id,
                            "company_id": current_user.company_id
                        }).fetchone()
                        
                        if approver_info:
                            # Get contract details
                            contract_query = text("""
                                SELECT contract_number, contract_title
                                FROM contracts WHERE id = :contract_id
                            """)
                            contract_info = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
                            
                            WorkflowEmailService.send_workflow_step_notification(
                                db=db,
                                contract_id=contract_id,
                                contract_number=contract_info.contract_number,
                                contract_title=contract_info.contract_title,
                                assignee_email=approver_info.email,
                                assignee_name=approver_info.full_name,
                                step_name=approver_info.step_name,
                                step_type=approver_info.step_type,
                                workflow_name=approver_info.workflow_name
                            )
                            logger.info(f"✉️ Approval workflow notification sent to {approver_info.email}")
                        else:
                            logger.warning("⚠️ First approver details not found for email notification")
                            
                except Exception as email_error:
                    logger.error(f"❌ Error sending approval workflow email: {str(email_error)}")

                
                # Update workflow instance status from 'active' to 'in_progress'
                activate_workflow_query = text("""
                    UPDATE workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    SET wi.status = 'in_progress', 
                        wi.current_step = 1,
                        wi.started_at = NOW()
                    WHERE wi.contract_id = :contract_id
                    AND w.company_id = :company_id
                """)
                db.execute(activate_workflow_query, {
                    "contract_id": contract_id,
                    "company_id": current_user.company_id
                })

                logger.info(f" Approval workflow status updated to 'in_progress'")

            # =====================================================
            # HANDLE SIGNATURE
            # =====================================================
            elif request_type == 'signature':
                logger.info(f"🔍 Processing SIGNATURE for contract {contract_id}")

                # Find the E-SIGN step in the workflow
                esign_step_query = text("""
                    SELECT 
                        ws.step_number,
                        ws.id as step_id,
                        ws.step_name,
                        ws.step_type,
                        ws.assignee_user_id,
                        w.id as workflow_id,
                        w.workflow_name,
                        w.is_master,
                        wi.id as instance_id
                    FROM workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    INNER JOIN workflow_steps ws ON w.id = ws.workflow_id
                    WHERE wi.contract_id = :contract_id
                    AND ws.step_type = 'e_sign_authority'
                    ORDER BY ws.step_number ASC
                    LIMIT 1
                """)
                
                esign_step = db.execute(esign_step_query, {
                    "contract_id": contract_id
                }).fetchone()
                
                if not esign_step:
                    logger.error("❌ No E-SIGN step found in workflow")
                    raise HTTPException(
                        status_code=400,
                        detail="No E-SIGN step found in workflow. Please configure workflow with E-SIGN authority step."
                    )
                
                logger.info(f" Found E-SIGN step: '{esign_step.step_name}' at position {esign_step.step_number}")
                logger.info(f"   Workflow: {esign_step.workflow_name} (Master: {esign_step.is_master})")
            

                # Update contract status to 'signature' and set action person
                contract.status = 'signature'
                contract.action_person_id = esign_step.assignee_user_id
                contract.updated_at = datetime.now()
                
                logger.info(f" Contract {contract_id} status updated to 'signature'")
                logger.info(f" action_person_id set to e-sign authority: {esign_step.assignee_user_id}")

# 📧 EMAIL NOTIFICATION: Contract sent for signature
                try:
                    if esign_step and esign_step.assignee_user_id:
                        logger.info(f"📧 Sending e-signature notification to user ID: {esign_step.assignee_user_id}")
                        
                        # Get e-sign authority details
                        esign_query = text("""
                            SELECT 
                                u.email,
                                CONCAT(u.first_name, ' ', u.last_name) as full_name
                            FROM users u
                            WHERE u.id = :user_id
                        """)
                        esign_info = db.execute(esign_query, {"user_id": esign_step.assignee_user_id}).fetchone()
                        
                        if esign_info:
                            # Get contract details including e-sign authorities
                            contract_query = text("""
                                SELECT 
                                    contract_number, 
                                    contract_title,
                                    party_esignature_authority_id,
                                    counterparty_esignature_authority_id
                                FROM contracts 
                                WHERE id = :contract_id
                            """)
                            contract_info = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
                            
                            # Determine party type
                            party_type = "Party A (Initiator)" if esign_step.assignee_user_id == contract_info.party_esignature_authority_id else "Party B (Counter-Party)"
                            
                            WorkflowEmailService.send_contract_sent_for_signature(
                                db=db,
                                contract_id=contract_id,
                                contract_number=contract_info.contract_number,
                                contract_title=contract_info.contract_title,
                                esign_authority_email=esign_info.email,
                                esign_authority_name=esign_info.full_name,
                                party_type=party_type
                            )
                            logger.info(f"✉️ E-signature notification sent to {esign_info.email} ({party_type})")
                        else:
                            logger.warning("⚠️ E-sign authority details not found for email notification")
                            
                except Exception as email_error:
                    logger.error(f"❌ Error sending e-signature email: {str(email_error)}")

                # Update workflow instance status from 'active' to 'in_progress'
                activate_workflow_query = text("""
                    UPDATE workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    SET wi.status = 'in_progress', 
                        wi.current_step = :esign_step_number,
                        wi.started_at = NOW()
                    WHERE wi.contract_id = :contract_id
                    AND w.company_id = :company_id
                """)
                db.execute(activate_workflow_query, {
                    "esign_step_number": esign_step.step_number,
                    "contract_id": contract_id,
                    "company_id": current_user.company_id
                })

                logger.info(f" Signature workflow status updated to 'in_progress'")
                
        else:
            # Log that counterparty submitted review (status not updated)
            logger.info(f"Counterparty (Party B) user {current_user.id} submitted review for contract {contract_id} (contract status not updated)")

            # =====================================================
            # HANDLE INTERNAL REVIEW
            # =====================================================
            if request_type == 'internal_review':
                logger.info(f"🔍 Processing Counterparty INTERNAL REVIEW for contract {contract_id}")

                # # Update workflow instance status to 'in_progress'
                # activate_workflow_query = text("""
                #     UPDATE workflow_instances wi
                #     INNER JOIN workflows w ON wi.workflow_id = w.id
                #     SET wi.status = 'in_progress', 
                #         wi.current_step = 1,
                #         wi.started_at = NOW()
                #     WHERE wi.contract_id = :contract_id
                #     AND w.company_id = :company_id
                # """)
                # db.execute(activate_workflow_query, {
                #     "contract_id": contract_id,
                #     "company_id": current_user.company_id
                # })

                # logger.info(f" Internal review workflow status updated to 'in_progress'")

                # Get first reviewer in the workflow for action_person_id
                logger.info(f"🔍 Finding first reviewer in internal review workflow...")
                first_reviewer_query = text("""
                    SELECT u.id
                    FROM workflow_instances wi
                    INNER JOIN workflows w ON wi.workflow_id = w.id
                    INNER JOIN workflow_steps ws ON w.id = ws.workflow_id
                    INNER JOIN users u ON ws.assignee_user_id = u.id
                    WHERE wi.contract_id = :contract_id
                    AND w.company_id = :company_id
                    AND ws.step_number = 1
                    AND w.is_active=1
                    ORDER BY is_master ASC
                    LIMIT 1
                """)
                first_reviewer = db.execute(first_reviewer_query, {
                    "contract_id": contract_id,
                    "company_id": current_user.company_id
                }).first()
                
                first_reviewer_id = first_reviewer.id if first_reviewer else None
                logger.info(f" First reviewer ID: {first_reviewer_id}")
                
                # Update contract with action_person_id
                update_contract_query = text("""
                    UPDATE contracts
                    SET action_person_id = :action_person_id,
                        updated_at = NOW()
                    WHERE id = :contract_id
                """)
                db.execute(update_contract_query, {
                    "action_person_id": first_reviewer_id,
                    "contract_id": contract_id
                })
                
                logger.info(f" Contract {contract_id} action_person_id set to: {first_reviewer_id}")

# 📧 EMAIL NOTIFICATION: Internal review request
                try:
                    logger.info("📧 Preparing to send internal review notification emails...")
                    reviewer_emails = []
                    
                    if review_type == 'specific':
                        # Get email addresses from request
                        reviewer_emails = personnel_emails if personnel_emails else []
                        logger.info(f"📋 Specific personnel review: {len(reviewer_emails)} reviewers")
                        
                    elif review_type == 'masterWorkflow':
                        # Get emails from master workflow users
                        master_users_query = text("""
                            SELECT DISTINCT
                                u.email,
                                CONCAT(u.first_name, ' ', u.last_name) as full_name
                            FROM workflow_steps ws
                            INNER JOIN users u ON ws.assignee_user_id = u.id
                            INNER JOIN workflows w ON ws.workflow_id = w.id
                            WHERE w.is_master = 1
                            AND w.company_id = :company_id
                            AND u.is_active = 1
                            AND u.email IS NOT NULL
                            AND u.id != :current_user_id
                        """)
                        master_users = db.execute(master_users_query, {
                            "company_id": current_user.company_id,
                            "current_user_id": current_user.id
                        }).fetchall()
                        
                        reviewer_emails = [user.email for user in master_users]
                        logger.info(f"📋 Master workflow review: {len(reviewer_emails)} reviewers")
                    
                    if reviewer_emails:
                        # Get contract details
                        contract_query = text("""
                            SELECT contract_number, contract_title
                            FROM contracts WHERE id = :contract_id
                        """)
                        contract_info = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
                        
                        initiator_name = f"{current_user.first_name} {current_user.last_name}"
                        
                        # Send emails to all reviewers
                        WorkflowEmailService.send_internal_review_request(
                            db=db,
                            contract_id=contract_id,
                            contract_number=contract_info.contract_number,
                            contract_title=contract_info.contract_title,
                            reviewer_emails=reviewer_emails,
                            initiator_name=initiator_name
                        )
                        logger.info(f"✉️ Internal review emails sent to {len(reviewer_emails)} reviewers")
                    else:
                        logger.warning("⚠️ No reviewer emails found to send notifications")
                        
                except Exception as email_error:
                    logger.error(f"❌ Error sending internal review emails: {str(email_error)}")
                

        # Send notifications to specific personnel (allowed for all users including counterparties)
        if review_type == 'specific' and personnel_emails:
            for email in personnel_emails:
                if not email:
                    continue
                    
                user_query = text("""
                    SELECT id FROM users WHERE email = :email LIMIT 1
                """)
                user = db.execute(user_query, {"email": email.strip()}).fetchone()
                
                if user:
                    notif_query = text("""
                        INSERT INTO notifications 
                        (recipient_id, notification_type, title, message, status, created_at)
                        VALUES (:user_id, 'contract_review', :title, :message, 'pending', NOW())
                    """)
                    db.execute(notif_query, {
                        "user_id": str(user.id),
                        "title": "Contract Review Required",
                        "message": f"Contract {contract_id} requires your review. Notes: {notes}"
                    })
        
        db.commit()
        
        # Return appropriate message based on user type
        if is_counterparty:
            message = "Review submitted successfully (workflow activated)"
        else:
            message = "Contract submitted for internal review successfully"
        
        return {
            "success": True,
            "message": message,
            "is_counterparty": is_counterparty
        }
        
    except HTTPException as he:
        db.rollback()
        raise he
    except Exception as e:
        db.rollback()
        logger.error(f"Error submitting contract for review: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))



# =====================================================
# HELPER FUNCTIONS - WITH FIXED CONTRACT NUMBER GENERATION
# =====================================================

def generate_contract_number(db: Session, company_id: int = None) -> str:
    """Generate unique contract number"""
    try:
        year = datetime.now().year
        month = datetime.now().month
        
        # Get the count of contracts for this year-month
        query = text("""
            SELECT COUNT(*) as count 
            FROM contracts 
            WHERE YEAR(created_at) = :year 
            AND MONTH(created_at) = :month
        """)
        
        result = db.execute(query, {"year": year, "month": month})
        count = result.fetchone()[0]
        
        # Generate contract number: CNT-YYYY-MM-XXXX
        contract_number = f"CNT-{year}-{month:02d}-{(count + 1):04d}"
        return contract_number
        
    except Exception as e:
        logger.error(f"Error generating contract number: {str(e)}")
        # Fallback to timestamp-based number
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        return f"CNT-{timestamp}"

def save_uploaded_file(file: UploadFile, contract_number: str, company_id: str) -> str:
    """Save uploaded file to storage"""
    upload_dir = Path("uploads") / "contracts" / str(company_id)
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    file_extension = os.path.splitext(file.filename)[1]
    filename = f"{contract_number}{file_extension}"
    file_path = upload_dir / filename
    
    with open(file_path, "wb") as buffer:
        import shutil
        shutil.copyfileobj(file.file, buffer)
    
    return str(file_path)

def calculate_completion(contract):
    """Calculate contract completion percentage"""
    status_completion = {
        'draft': 10,
        'pending_review': 30,
        'negotiation': 50,
        'approved': 80,
        'active': 100,
        'completed': 100,
        'expired': 100,
        'terminated': 100,
        'rejected': 100
    }
    
    base = status_completion.get(contract.status, 0)
    
    workflow_adjustment = {
        'draft': 0,
        'internal_review': 10,
        'clause_analysis': 15,
        'external_review': 20,
        'negotiation': 30,
        'approval': 40,
        'execution': 50,
        'active': 0
    }
    
    adjustment = workflow_adjustment.get(contract.workflow_status, 0) if contract.workflow_status else 0
    completion = base + adjustment
    
    if hasattr(contract, 'is_signed') and contract.is_signed:
        completion = max(completion, 90)
    
    return min(completion, 100)

# =====================================================
# FILE: app/api/api_v1/contracts/contracts.py
# COMPLETE FIX: Risk Analysis with Data Transformation
# =====================================================

from app.services.rag_service import ContractRAGService

# app/api/contracts.py - COMPLETE FIXED FUNCTION

@router.post("/risk-analysis/{contract_id}")
async def analyze_contract_risks(
    contract_id: int,
    request: dict = {}, 
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """AI-powered risk analysis with RAG - supports any jurisdiction worldwide"""
    try:
        language = request.get("language", "en") if request else "en"
        profile_type = request.get("profile_type", "client")
        
        logger.info(f"🔍 Starting RAG-enhanced risk analysis for contract {contract_id}")
        
        # Check for existing RAG-enhanced analysis
        query = text("""
            SELECT analysis_data, risk_score, created_at
            FROM ai_analysis_results
            WHERE CAST(contract_id AS CHAR) = CAST(:contract_id AS CHAR)
            AND analysis_type = 'risk_analysis'
            ORDER BY created_at DESC
            LIMIT 1
        """)
        
        existing_analysis = db.execute(query, {"contract_id": str(contract_id)}).fetchone()

        if existing_analysis:
            try:
                analysis_data = json.loads(existing_analysis[0]) if isinstance(existing_analysis[0], str) else existing_analysis[0]
                
                #  CHECK IF THIS IS A RAG-ENHANCED ANALYSIS
                is_rag_enhanced = analysis_data.get("rag_enhanced", False)
                analysis_version = analysis_data.get("analysis_version", "")
                
                if is_rag_enhanced and "rag" in analysis_version.lower():
                    # This is already RAG-enhanced, use cache
                    analysis_date = existing_analysis[2]
                    hours_old = (datetime.now() - analysis_date).total_seconds() / 3600
                    
                    if hours_old < 24:  # Cache for 24 hours
                        logger.info(f" Returning cached RAG analysis (age: {hours_old:.1f} hours)")
                        
                        return {
                            "success": True,
                            "contract_id": contract_id,
                            "analysis": analysis_data,
                            "ai_powered": True,
                            "rag_enhanced": True,
                            "cached": True,
                            "analyzed_at": existing_analysis[2].isoformat()
                        }
                else:
                    # Old non-RAG analysis, regenerate with RAG
                    logger.info(f"🔄 Found old non-RAG analysis, regenerating with RAG enhancement...")
                    
            except Exception as parse_error:
                logger.error(f"❌ Error parsing cached analysis: {str(parse_error)}")
        
        # Get contract details
        contract_query = text("""
            SELECT 
                c.contract_title, 
                c.contract_type, 
                c.status, 
                c.contract_value
            FROM contracts c
            WHERE c.id = :contract_id
        """)
        
        contract = db.execute(contract_query, {"contract_id": contract_id}).fetchone()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        contract_title = contract[0] or "Untitled Contract"
        contract_type = contract[1] or "General"
        contract_value = contract[3] or 0.0
        
        # Get the latest contract version content
        version_content = db.execute(text("""
            SELECT contract_content
            FROM contract_versions
            WHERE contract_id = :contract_id
            ORDER BY created_at DESC
            LIMIT 1
        """), {"contract_id": contract_id}).fetchone()
        
        if version_content and version_content[0]:
            import re
            contract_content = re.sub('<[^<]+?>', '', version_content[0])
            contract_content = re.sub(r'\s+', ' ', contract_content).strip()
            logger.info(f"📄 Retrieved contract content: {len(contract_content)} characters")
        else:
            logger.warning(f"⚠️ No contract content found for analysis")
            raise HTTPException(status_code=400, detail="No contract content available for analysis")
        
        # 🚀 INITIALIZE RAG SERVICE
        from app.services.rag_service import ContractRAGService
        rag_service = ContractRAGService()
        
        # 🚀 INDEX THE CONTRACT DOCUMENT
        logger.info(f"📊 Indexing contract document into vector database...")
        indexed = rag_service.index_contract(
            contract_id=contract_id,
            contract_content=contract_content,
            contract_title=contract_title
        )
        
        if not indexed:
            logger.warning("⚠️ Failed to index contract, proceeding with fallback")
        
        # 🚀 RETRIEVE RELEVANT CHUNKS
        risk_analysis_query = f"""
        Analyze this {contract_type} contract for:
        - Governing law and jurisdiction clauses
        - Liability and indemnity clauses
        - Payment and termination terms
        - Time and delay provisions
        - Insurance and security requirements
        - Dispute resolution mechanisms
        - Compliance and regulatory risks
        - Performance obligations
        - Force majeure provisions
        - Warranties and representations
        - Intellectual property rights
        - Confidentiality and data protection
        """
        
        relevant_chunks = rag_service.retrieve_relevant_chunks(
            contract_id=contract_id,
            query=risk_analysis_query,
            n_results=20
        )
        
        # Build context from retrieved chunks (but don't expose chunk numbers to Claude)
        if relevant_chunks:
            context = "\n\n".join([chunk['text'] for chunk in relevant_chunks])
            logger.info(f" Retrieved {len(relevant_chunks)} relevant sections ({len(context)} chars)")
        else:
            # Fallback to full document if RAG fails
            context = contract_content[:15000]
            logger.warning("⚠️ Using fallback full document approach")
        
        # Initialize Claude service
        from app.services.claude_service import ClaudeService
        claude_service = ClaudeService()
        
        if not claude_service.client:
            logger.error("❌ Claude service not initialized")
            raise HTTPException(status_code=503, detail="AI service unavailable")
        
        # Generate comprehensive risk analysis prompt
        risk_prompt = f"""You are an expert international Contract and Legal Risk Analyst with expertise in:
- Commercial contracts across all jurisdictions
- Construction, Infrastructure, Oil & Gas, Technology, Service Agreements
- International trade and cross-border transactions
- Common Law, Civil Law, and Sharia Law systems
- International arbitration (ICC, LCIA, SIAC, DIFC Courts, etc.)

You are analyzing from the perspective of a **{profile_type}**.

═══════════════════════════════════════════════════════
CONTRACT DETAILS
═══════════════════════════════════════════════════════
Title: {contract_title}
Type: {contract_type}
Profile/Role: {profile_type}
Contract Value: {contract_value}

**IMPORTANT**: You must identify the governing law and jurisdiction from the contract content itself.

═══════════════════════════════════════════════════════
RELEVANT CONTRACT SECTIONS FOR ANALYSIS
═══════════════════════════════════════════════════════
The following are the most relevant sections of the contract document:

{context}

═══════════════════════════════════════════════════════
COMPREHENSIVE RISK ANALYSIS REQUIREMENTS
═══════════════════════════════════════════════════════

**STEP 1: IDENTIFY JURISDICTION AND GOVERNING LAW**
First, carefully examine the contract to identify:
- What is the governing law clause?
- What jurisdiction is specified?
- Where will disputes be resolved?
- What legal system applies (Common Law, Civil Law, Sharia, Mixed)?

**STEP 2: PERFORM COMPREHENSIVE RISK ANALYSIS**

Analyze ALL of the following dimensions:

**1. DOCUMENT STRUCTURE & PRECEDENCE**
- Document hierarchy and order of precedence
- Definitions and interpretation clauses
- Integration and entire agreement clauses
- Amendment and waiver provisions
- Severability clauses

**2. PARTIES & CAPACITY**
- Party identification and legal capacity
- Authority to execute
- Corporate/entity structure considerations
- Guarantees and parent company obligations

**3. SCOPE OF WORK / SERVICES**
- Detailed scope definition
- Deliverables and acceptance criteria
- Performance standards and KPIs
- Interface obligations
- Exclusions and limitations

**4. FINANCIAL TERMS**
- Contract price/value and payment structure
- Payment terms and milestones
- Retention and holdbacks
- Price adjustment mechanisms
- Currency and exchange rate risks

**5. TIME & SCHEDULE**
- Commencement and completion dates
- Time for performance
- Extension of time provisions
- Liquidated damages for delay
- Suspension rights

**6. VARIATIONS & CHANGE CONTROL**
- Variation definition and procedures
- Change order process
- Valuation methodology
- Time and cost impacts

**7. WARRANTIES & REPRESENTATIONS**
- Express warranties
- Implied warranties
- Duration and scope
- Warranty exclusions

**8. LIABILITY & INDEMNITIES**
- Liability caps and limitations
- Carve-outs from caps (CRITICAL)
- Indemnity provisions
- Consequential loss exclusions
- Third-party claims

**9. INSURANCE & SECURITY**
- Required insurance policies
- Coverage amounts and terms
- Additional insured status
- Performance bonds
- Bank guarantees
- Letters of credit

**10. INTELLECTUAL PROPERTY**
- IP ownership
- License grants
- Background vs foreground IP
- IP indemnities

**11. CONFIDENTIALITY & DATA PROTECTION**
- Confidentiality obligations
- Data protection compliance (GDPR, local laws)
- Personal data handling
- Duration of obligations

**12. TERMINATION**
- Termination for cause
- Termination for convenience
- Termination procedures
- Payment on termination
- Survival clauses

**13. FORCE MAJEURE**
- Force majeure definition
- Notice requirements
- Consequences (time vs cost relief)
- Long-stop provisions

**14. DISPUTE RESOLUTION**
- Governing law (identify from contract)
- Arbitration vs litigation
- Seat of arbitration
- Arbitration rules
- Expert determination
- Escalation procedures

**15. REGULATORY & COMPLIANCE**
- Industry-specific regulations
- Anti-bribery and corruption
- Sanctions and export controls
- Health, safety, and environment
- Local content requirements
- Sustainability obligations

═══════════════════════════════════════════════════════
OUTPUT FORMAT - RESPOND ONLY WITH THIS JSON STRUCTURE
═══════════════════════════════════════════════════════

{{
    "executive_summary": {{
        "overview": "<3-4 paragraph executive summary. Identify and mention the jurisdiction and governing law prominently>",
        "overall_score": <0-100 integer>,
        "key_risk_areas": ["<Area 1>", "<Area 2>", "<Area 3>", "<Area 4>", "<Area 5>"],
        "overall_risk_rating": "<High/Medium/Low>",
        "jurisdiction_status": "<Clear/Ambiguous/Missing>"
    }},
    
    "contract_overview": {{
        "parties_involved": ["<Party 1 with full details>", "<Party 2 with full details>"],
        "purpose": "<2 paragraphs describing contract purpose>",
        "contract_value": "{contract_value}",
        "term_summary": "<Start/end dates, duration, renewal terms>",
        "governing_law": "<Governing law identified from contract>",
        "jurisdiction": "<Jurisdiction identified from contract>",
        "dispute_mechanism": "<Arbitration/Court details from contract>",
        "applicable_legal_system": "<Common Law/Civil Law/Sharia/Mixed/Not Clear>"
    }},
    
    "clause_wise_analysis": [
        {{
            "clause_number": "<Actual clause number from contract, e.g. 'Clause 17.6', 'Article 5.2', 'Section 3.1'>",
            "clause_title": "<Title of the clause>",
            "severity": "<high/medium/low>",
            "clause_summary": "<2-3 paragraph summary>",
            "risks": ["<Risk 1>", "<Risk 2>"],
            "financial_impact": "<Impact analysis>",
            "enforceability": "<Enforceability under identified governing law>",
            "jurisdiction_specific_issues": "<Any jurisdiction-specific concerns>",
            "mitigation_strategy": "<2-3 paragraphs>"
        }}
    ],
    
    "regulatory_compliance": {{
        "jurisdiction_analysis": "<Analysis specific to identified jurisdiction - 2-3 paragraphs>",
        "compliance_gaps": ["<Gap 1>", "<Gap 2>"],
        "required_amendments": ["<Amendment 1>", "<Amendment 2>"],
        "licenses_permits": ["<License/Permit 1>", "<License/Permit 2>"],
        "cross_border_considerations": ["<Consideration 1>", "<Consideration 2>"]
    }},
    
    "risk_mitigation": {{
        "immediate_actions": ["<Action 1>", "<Action 2>", "<Action 3>"],
        "negotiation_strategy": "<2-3 paragraphs>",
        "fallback_positions": ["<Fallback 1>", "<Fallback 2>"],
        "deal_breakers": ["<Deal breaker 1>", "<Deal breaker 2>"],
        "alternative_structures": ["<Alternative 1>", "<Alternative 2>"]
    }},
    
    "negotiation_points": {{
        "high_priority": [
            {{
                "issue": "<Issue>",
                "current_position": "<Current - 2 paragraphs>",
                "recommended_position": "<Recommended - 2 paragraphs>",
                "fallback_options": ["<Option 1>", "<Option 2>"],
                "rationale": "<1-2 paragraphs>",
                "jurisdiction_angle": "<Jurisdiction-specific negotiation point if relevant>"
            }}
        ],
        "medium_priority": [
            {{
                "issue": "<Issue>",
                "current_position": "<Current>",
                "recommended_position": "<Recommended>",
                "fallback_options": ["<Option 1>", "<Option 2>"]
            }}
        ],
        "suggested_new_clauses": [
            {{
                "clause_type": "<Type>",
                "suggested_wording": "<Complete text>",
                "purpose": "<2 paragraphs>",
                "jurisdiction_compliant": "<Yes/No/Needs Review>"
            }}
        ]
    }},
    
    "risk_summary": {{
        "high_risks": <count>,
        "medium_risks": <count>,
        "low_risks": <count>,
        "total_risks": <count>,
        "risk_items": [
            {{
                "type": "<Legal/Financial/Operational/Compliance/Jurisdictional>",
                "issue": "<Title>",
                "description": "<3-4 paragraphs - reference actual clause numbers like 'Clause 17.6' or 'Article 5', NOT chunk numbers>",
                "severity": "<high/medium/low>",
                "score": <1-100>,
                "clause_reference": "<Actual clause numbers from contract, e.g. 'Clause 17.6, Sub-Clause 18.9, Article 5'>",
                "mitigation": "<2-3 paragraphs>",
                "financial_impact": "<Impact analysis>",
                "jurisdiction_impact": "<How identified jurisdiction affects this risk>"
            }}
        ],
        "red_flags": ["<Red flag 1>", "<Red flag 2>"]
    }},
    
    "conclusion": {{
        "key_findings": ["<Finding 1>", "<Finding 2>", "<Finding 3>"],
        "final_risk_position": "<3-4 paragraphs>",
        "execution_readiness": "<Yes/Conditional/No>",
        "conditions_for_execution": ["<Condition 1>", "<Condition 2>"],
        "recommendations_priority": ["<Top 1>", "<Top 2>", "<Top 3>"],
        "next_steps": ["<Step 1>", "<Step 2>", "<Step 3>"],
        "legal_counsel_required": "<Yes/No/Recommended>",
        "jurisdiction_specific_advice": "<Recommendations specific to identified jurisdiction>"
    }}
}}

**CRITICAL INSTRUCTIONS - READ CAREFULLY:**
1. **NEVER EVER mention "chunks" or "chunk numbers" anywhere in your response**
2. **ALWAYS use actual clause numbers from the contract document** - look for references like:
   - "Clause 17.6" or "Sub-Clause 18.9"
   - "Article 5" or "Section 3.1"
   - "Schedule 7" or "Appendix A"
3. **If you cannot find specific clause numbers, use descriptive references:**
   - "the liability limitation provisions"
   - "termination section"
   - "payment terms"
   - "insurance requirements"
4. **Write as if reading a physical contract** - reference sections naturally as they appear in the document
5. **Be professional and business-focused** - this is for executives and legal teams, not technical users
6. **Extract clause numbers directly from the text** - they should be visible in the contract content
7. Provide 3-4 paragraphs for each major risk with specific clause citations
8. Include financial impact ranges where possible
9. Be actionable - every recommendation must be implementable
10. Adapt analysis to the identified legal system (Common Law, Civil Law, Sharia, etc.)

Analyze the contract now and provide the comprehensive JSON response."""

        # Call Claude API
        logger.info(f"🤖 Calling Claude API with RAG-enhanced prompt ({len(risk_prompt)} chars)")
        
        message = claude_service.client.messages.create(
            model=claude_service.model,
            max_tokens=16384,
            temperature=0.3,
            messages=[{"role": "user", "content": risk_prompt}]
        )
        
        # Extract and parse response
        ai_response = message.content[0].text.strip()
        logger.info(f" Claude response received ({len(ai_response)} chars)")
        
        # Clean up response
        if ai_response.startswith("```"):
            ai_response = ai_response.split("```")[1]
            if ai_response.startswith("json"):
                ai_response = ai_response[4:]
            ai_response = ai_response.strip()
        
        # Parse JSON
        claude_analysis = json.loads(ai_response)
        logger.info(f" RAG-enhanced AI analysis parsed successfully")
        
        # Extract data
        executive_summary = claude_analysis.get("executive_summary", {})
        risk_summary = claude_analysis.get("risk_summary", {})
        contract_overview = claude_analysis.get("contract_overview", {})
        
        # Extract jurisdiction and governing law from Claude's analysis
        jurisdiction = contract_overview.get("jurisdiction", "Not specified")
        governing_law = contract_overview.get("governing_law", "Not specified")
        
        overall_risk_score = executive_summary.get("overall_score", 50)
        risk_items = risk_summary.get("risk_items", [])
        
        # Count risks
        high_risks = risk_summary.get("high_risks", len([r for r in risk_items if r.get("severity", "").lower() == "high"]))
        medium_risks = risk_summary.get("medium_risks", len([r for r in risk_items if r.get("severity", "").lower() == "medium"]))
        low_risks = risk_summary.get("low_risks", len([r for r in risk_items if r.get("severity", "").lower() == "low"]))
        
        logger.info(f"📍 Identified jurisdiction: {jurisdiction}, Governing law: {governing_law}")
        
        # Format analysis
        formatted_analysis = {
            "overall_score": overall_risk_score,
            "risk_score": overall_risk_score,
            "risk_level": executive_summary.get("overall_risk_rating", "Medium"),
            "executive_summary": executive_summary,
            "contract_overview": contract_overview,
            "clause_wise_analysis": claude_analysis.get("clause_wise_analysis", []),
            "regulatory_compliance": claude_analysis.get("regulatory_compliance", {}),
            "risk_mitigation": claude_analysis.get("risk_mitigation", {}),
            "negotiation_points": claude_analysis.get("negotiation_points", {}),
            "risk_summary": {
                "high_risks": high_risks,
                "medium_risks": medium_risks,
                "low_risks": low_risks,
                "total_risks": high_risks + medium_risks + low_risks,
                "risk_items": risk_items,
                "red_flags": risk_summary.get("red_flags", [])
            },
            "conclusion": claude_analysis.get("conclusion", {}),
            "profile_type": profile_type,
            "jurisdiction": jurisdiction,
            "governing_law": governing_law,
            "rag_enhanced": True,
            "chunks_retrieved": len(relevant_chunks),
            "analyzed_at": datetime.now().isoformat(),
            "analysis_version": "rag_universal_v1.0"
        }
        
        # Save to database
        save_query = text("""
            INSERT INTO ai_analysis_results 
            (contract_id, analysis_type, analysis_data, risk_score, created_at)
            VALUES (CAST(:contract_id AS CHAR(36)), :analysis_type, :analysis_data, :risk_score, NOW())
        """)
        
        db.execute(save_query, {
            "contract_id": str(contract_id),
            "analysis_type": "risk_analysis",
            "analysis_data": json.dumps(formatted_analysis),
            "risk_score": overall_risk_score
        })
        db.commit()
        
        logger.info(f" RAG-enhanced risk analysis saved for contract {contract_id} (chunks: {len(relevant_chunks)}, jurisdiction: {jurisdiction})")
        
        return {
            "success": True,
            "contract_id": contract_id,
            "contract_title": contract_title,
            "analysis": formatted_analysis,
            "ai_powered": True,
            "rag_enhanced": True,
            "chunks_used": len(relevant_chunks),
            "jurisdiction": jurisdiction,
            "governing_law": governing_law,
            "model": "claude-sonnet-4-20250514"
        }
        
    except json.JSONDecodeError as e:
        logger.error(f"❌ JSON parsing error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to parse AI response")
    except Exception as e:
        logger.error(f"❌ Error in RAG-enhanced risk analysis: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

        
# app/api/contracts.py - COMPLETE FIXED FUNCTION

@router.post("/risk-analysis-upload")
async def upload_contract_for_risk_analysis(
    file: UploadFile = File(...),
    profile_type: str = Form("client"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload contract file for risk analysis"""
    try:
        logger.info(f"📋 Risk Analysis Upload from user {current_user.email}")
        logger.info(f"📎 File: {file.filename}, Profile: {profile_type}")
        
        # Validate file type
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ['.pdf', '.docx', '.doc']:
            raise HTTPException(status_code=400, detail="Only PDF and Word documents are supported")
        
        # Generate unique contract number
        year = datetime.now().year
        timestamp_suffix = int(datetime.now().timestamp() * 1000) % 100000
        contract_number = f"CNT-{year}-{timestamp_suffix:05d}"
        
        logger.info(f"🔢 Generated contract number: {contract_number}")
        
        # Create contract record
        from sqlalchemy import text as sql_text
        contract_query = sql_text("""
            INSERT INTO contracts (
                contract_number, contract_title, contract_type, 
                company_id, created_by, status, created_at, updated_at
            ) VALUES (
                :contract_number, :contract_title, :contract_type,
                :company_id, :created_by, 'draft', NOW(), NOW()
            )
        """)
        
        db.execute(contract_query, {
            "contract_number": contract_number,
            "contract_title": f"Risk Analysis - {file.filename}",
            "contract_type": "risk_analysis",
            "company_id": current_user.company_id,
            "created_by": current_user.id
        })
        db.commit()
        
        # Get created contract ID
        contract_id_result = db.execute(sql_text("""
            SELECT id FROM contracts WHERE contract_number = :contract_number
        """), {"contract_number": contract_number}).fetchone()
        
        contract_id = contract_id_result[0]
        logger.info(f" Contract created with ID: {contract_id}")
        
        # Save uploaded file
        upload_dir = f"app/uploads/contracts/{contract_id}"
        os.makedirs(upload_dir, exist_ok=True)
        file_path = os.path.join(upload_dir, file.filename)
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        logger.info(f"💾 File saved to: {file_path}")
        
        # Extract text from document
        logger.info(f"📄 Extracting text from {file_ext} file...")
        
        plain_text = ""
        html_content = ""
        
        if file_ext == '.pdf':
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"PDF has {len(pdf.pages)} pages")
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                plain_text = '\n'.join(text_parts)
                logger.info(f"pdfplumber extracted {len(plain_text)} characters")
        
        elif file_ext in ['.docx', '.doc']:
            import docx2txt
            plain_text = docx2txt.process(file_path)
            logger.info(f"docx2txt extracted {len(plain_text)} characters")
        
        if not plain_text:
            raise HTTPException(status_code=400, detail="Failed to extract content from document")
        
        # Convert to HTML for storage
        html_content = f"""
        <div class="contract-document" style="font-family: Arial, sans-serif; line-height: 1.6; color: #333;">
            {plain_text.replace(chr(10), '<br>')}
        </div>
        """
        
        logger.info(f" Extracted {len(plain_text)} characters from document")
        
        # Save contract version
        version_query = sql_text("""
            INSERT INTO contract_versions (
                contract_id, version_number, contract_content, 
                created_by, created_at
            ) VALUES (
                :contract_id, 1, :contract_content, 
                :created_by, NOW()
            )
        """)
        
        db.execute(version_query, {
            "contract_id": contract_id,
            "contract_content": html_content,
            "created_by": current_user.id
        })
        db.commit()
        
        logger.info(f" Contract version saved")
        
        # Return contract ID - frontend will call /risk-analysis/{contract_id}
        return {
            "success": True,
            "message": "Contract uploaded successfully. Starting analysis...",
            "contract_id": contract_id,
            "contract_number": contract_number,
            "contract_title": f"Risk Analysis - {file.filename}"
        }
        
    except Exception as e:
        logger.error(f"❌ Error in risk analysis upload: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
        
# =====================================================
# CLAUSE ANALYSIS ENDPOINTS
# =====================================================

@router.post("/clause-suggestions")
async def get_clause_suggestions(
    request: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get AI suggestions for clause improvements using Claude API"""
    try:
        clause_text = request.get("clause_text", "")
        clause_type = request.get("clause_type", "general")
        
        # Return mock suggestions if no clause text
        if not clause_text:
            suggestions = [
                {
                    "original": clause_text,
                    "suggested": "Either party may terminate for convenience with ninety (90) days notice. Termination for cause requires thirty (30) days notice with fifteen (15) day cure period.",
                    "reasoning": "Provides balanced termination rights with appropriate notice periods and cure provisions",
                    "legal_score": 5,
                    "business_score": 4,
                    "compliance": "Complies with Qatar Labor Law Article 61"
                },
                {
                    "original": clause_text,
                    "suggested": "Either party may terminate this Agreement upon sixty (60) days written notice. The terminating party shall provide detailed reasons for termination.",
                    "reasoning": "Standard termination clause with documentation requirements",
                    "legal_score": 4,
                    "business_score": 4,
                    "compliance": "Standard industry practice"
                }
            ]
            return {"success": True, "suggestions": suggestions}
        
        # Use Claude API for real suggestions
        try:
            # Initialize Claude service
            claude_service = ClaudeService()
            
            if not claude_service.client:
                logger.warning("Claude client not available, using mock suggestions")
                raise ValueError("Claude client not initialized")
            
            # Prepare the prompt for Claude
            prompt = f"""You are a legal expert specializing in contract law, particularly in Qatar jurisdiction. Analyze and improve the following {clause_type} clause.

Original Clause:
{clause_text}

Please provide 2-3 improved versions of this clause and respond in the following JSON format:
{{
    "suggestions": [
        {{
            "original": "<original clause text>",
            "suggested": "<improved clause text>",
            "reasoning": "<explanation of why this version is better>",
            "legal_score": <1-5, where 5 is strongest legal protection>,
            "business_score": <1-5, where 5 is most business-friendly>,
            "compliance": "<relevant Qatar laws or regulations this complies with>"
        }}
    ]
}}

Consider the following in your suggestions:
1. Compliance with Qatar Civil and Commercial Law
2. Clarity and enforceability
3. Balance between parties' interests
4. Industry best practices
5. Risk mitigation
6. Specific requirements for {clause_type} clauses

Provide only the JSON response without any additional text or markdown formatting."""

            response = claude_service.client.messages.create(
                model=claude_service.model,
                max_tokens=2000,
                temperature=0.5,
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            # Extract the text content from Claude's response
            response_text = response.content[0].text
            
            # Try to extract JSON from the response
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                result = json.loads(json_match.group())
                suggestions = result.get("suggestions", [])
                
                # Ensure we have at least one suggestion
                if not suggestions:
                    raise ValueError("No suggestions returned from Claude")
                
                # Log the clause improvement to audit trail
                audit_log = text("""
                    INSERT INTO audit_logs 
                    (user_id, action_type, action_details, ip_address, user_agent, created_at)
                    VALUES 
                    (:user_id, :action_type, :action_details, :ip_address, :user_agent, NOW())
                """)
                
                db.execute(audit_log, {
                    "user_id": current_user.id,
                    "action_type": "AI_CLAUSE_SUGGESTION",
                    "action_details": json.dumps({
                        "clause_type": clause_type,
                        "original_length": len(clause_text),
                        "suggestions_count": len(suggestions)
                    }),
                    "ip_address": "system",
                    "user_agent": "Claude AI"
                })
                db.commit()
                
                return {"success": True, "suggestions": suggestions}
            else:
                raise ValueError("Could not extract JSON from Claude response")
                    
        except Exception as e:
            logger.error(f"Error using Claude API for clause suggestions: {str(e)}")
            # Fall through to default suggestions
        
        # Return default suggestions if Claude API fails
        suggestions = [
            {
                "original": clause_text,
                "suggested": "Either party may terminate for convenience with ninety (90) days written notice. Termination for cause requires thirty (30) days written notice with fifteen (15) day cure period. Upon termination, all outstanding obligations shall be settled within thirty (30) days.",
                "reasoning": "Provides balanced termination rights with appropriate notice periods, cure provisions, and settlement timeline",
                "legal_score": 5,
                "business_score": 4,
                "compliance": "Complies with Qatar Labor Law Article 61 and Civil Code provisions"
            },
            {
                "original": clause_text,
                "suggested": "Either party may terminate this Agreement upon sixty (60) days written notice. The terminating party shall provide detailed written reasons for termination. No termination shall relieve either party of obligations incurred prior to the effective termination date.",
                "reasoning": "Standard termination clause with documentation requirements and liability preservation",
                "legal_score": 4,
                "business_score": 4,
                "compliance": "Standard industry practice aligned with Qatar commercial regulations"
            },
            {
                "original": clause_text,
                "suggested": "This Agreement may be terminated by mutual written consent or by either party upon material breach by the other party, provided that written notice of such breach is given and the breaching party fails to cure within twenty (20) business days.",
                "reasoning": "Focuses on breach-based termination with reasonable cure period, protecting both parties",
                "legal_score": 4,
                "business_score": 3,
                "compliance": "Consistent with Qatar Civil Code Article 171 on contractual obligations"
            }
        ]
        
        return {"success": True, "suggestions": suggestions}
        
    except Exception as e:
        logger.error(f"Error getting clause suggestions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# =====================================================
# VERSION HISTORY ENDPOINTS
# =====================================================

@router.get("/versions/{contract_id}")
async def get_version_history(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get complete version history for contract"""
    try:
        query = text("""
            SELECT 
                cv.id,
                cv.version_number,
                cv.version_type,
                cv.change_summary,
                cv.created_at,
                CONCAT(u.first_name, ' ', u.last_name) as created_by_name,
                LEFT(cv.contract_content, 500) as content_preview
            FROM contract_versions cv
            LEFT JOIN users u ON cv.created_by = u.id
            WHERE cv.contract_id = :contract_id
            ORDER BY cv.version_number DESC
        """)
        
        versions = db.execute(query, {"contract_id": contract_id}).fetchall()
        
        version_list = []
        for v in versions:
            version_list.append({
                "id": v.id,
                "version": v.version_number,
                "type": v.version_type or "draft",
                "created_at": v.created_at.isoformat() if v.created_at else None,
                "created_by": v.created_by_name or "Unknown",
                "notes": v.change_summary or "No notes",
                "content_preview": v.content_preview or ""
            })
        
        return {"success": True, "versions": version_list}
        
    except Exception as e:
        logger.error(f"Error getting version history: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/versions/compare")
async def compare_versions(
    contract_id: int,
    version1: int = Query(...),
    version2: int = Query(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Compare two contract versions"""
    try:
        query = text("""
            SELECT version_number, contract_content 
            FROM contract_versions
            WHERE contract_id = :contract_id 
            AND version_number IN (:v1, :v2)
        """)
        
        results = db.execute(query, {
            "contract_id": contract_id,
            "v1": version1,
            "v2": version2
        }).fetchall()
        
        if len(results) < 2:
            raise HTTPException(status_code=404, detail="One or both versions not found")
        
        # Simple diff - count changes
        v1_content = results[0].contract_content or ""
        v2_content = results[1].contract_content or ""
        
        changes = []
        if v1_content != v2_content:
            changes.append({
                "type": "content",
                "description": f"Content changed between version {version1} and {version2}",
                "characters_added": max(0, len(v2_content) - len(v1_content)),
                "characters_removed": max(0, len(v1_content) - len(v2_content))
            })
        
        return {
            "success": True,
            "version1": version1,
            "version2": version2,
            "changes": changes,
            "total_changes": len(changes)
        }
        
    except Exception as e:
        logger.error(f"Error comparing versions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# =====================================================
# EXPORT ENDPOINTS
# =====================================================
@router.post("/export/{contract_id}")
async def export_contract_improved(
    contract_id: int,
    format: str = Query("pdf"),
    include_signatures: bool = Query(True),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export contract with proper formatting and signatures
    Supports: PDF, DOCX, HTML
    """
    try:
        # Get contract data
        query = text("""
            SELECT 
                c.id,
                c.contract_number,
                c.contract_title,
                c.status,
                cv.contract_content,
                c.created_at,
                c.effective_date
            FROM contracts c
            LEFT JOIN contract_versions cv ON c.id = cv.contract_id
            WHERE c.id = :contract_id
            ORDER BY cv.version_number DESC
            LIMIT 1
        """)
        
        result = db.execute(query, {"contract_id": contract_id}).fetchone()
        
        if not result:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Get signatories if requested
        signatories = []
        if include_signatures:
            sig_query = text("""
                SELECT 
                    s.signer_type,
                    s.has_signed,
                    s.signed_at,
                    s.signature_data,
                    s.signature_method,
                    u.first_name,
                    u.last_name
                FROM signatories s
                LEFT JOIN users u ON s.user_id = u.id
                WHERE s.contract_id = :contract_id
                ORDER BY s.signing_order
            """)
            signatories = db.execute(sig_query, {"contract_id": contract_id}).fetchall()
        
        # Export based on format
        if format == "pdf":
            return export_as_pdf(result, signatories)
        elif format == "word" or format == "docx":
            return export_as_word(result, signatories)
        elif format == "html":
            return export_as_html(result, signatories)
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================
# FIXED MULTI-PAGE PDF EXPORT
# Replace the export_as_pdf() function in contracts.py
# =====================================================
def export_as_pdf(contract, signatories):
    """
    Generate PDF from contract content as-is
    Signatures should already be in the contract content from editor
    """
    
    # Get HTML content directly from database (includes embedded signatures)
    content_html = contract.contract_content or "<p>No content available</p>"
    
    # Complete HTML document
    complete_html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                size: A4;
                margin: 20mm;
            }}
            
            body {{
                font-family: Arial, Helvetica, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #000;
            }}
            
            .header {{
                border-bottom: 2px solid #2563eb;
                padding-bottom: 15px;
                margin-bottom: 20px;
            }}
            
            .header h1 {{
                margin: 0;
                font-size: 18pt;
                text-align: center;
            }}
            
            .header p {{
                margin: 5px 0;
                font-size: 10pt;
                color: #666;
            }}
            
            /* Ensure images display properly */
            img {{
                max-width: 100%;
                height: auto;
            }}
        </style>
    </head>
    <body>
      
        
        <!-- Contract Content (includes signatures at their designated positions) -->
        <div class="contract-content">
            {content_html}
        </div>
    </body>
    </html>
    '''
    
    # Convert HTML to PDF using WeasyPrint
    pdf_bytes = HTML(string=complete_html).write_pdf()
    
    buffer = BytesIO(pdf_bytes)
    buffer.seek(0)
    
    filename = f"{contract.contract_number or 'Contract'}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

# =====================================================
# DOCX EXPORT WITH FULL STYLING PRESERVED
# Uses Word's HTML format to preserve all styling
# Replace export_as_word() in contracts.py
# =====================================================
# =====================================================
# PROPER DOCX EXPORT USING PANDOC
# Install: sudo apt-get install pandoc && pip install pypandoc
# Replace export_as_word() in contracts.py
# =====================================================



def inject_signatures_into_html(content_html, signatories):
    """
    Inject actual signature images/text into the contract HTML
    at their designated positions (where signature blocks exist)
    """
    if not signatories or len(signatories) == 0:
        return content_html
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(content_html, 'html.parser')
    
    # Process each signatory
    for sig in signatories:
        if not sig.has_signed or not sig.signature_data:
            continue
            
        signer_type = sig.signer_type.lower()
        signer_name = f"{sig.first_name} {sig.last_name}" if sig.first_name else "Unknown"
        
        # Find signature placeholder for this signer type
        # Look for common patterns: class="signature-client", class="signature-company", etc.
        signature_containers = soup.find_all(['div', 'span'], 
            class_=lambda x: x and signer_type in x.lower() if x else False)
        
        if not signature_containers:
            # Alternative: look for text patterns like "CLIENT", "COMPANY" in headings
            for heading in soup.find_all(['h3', 'h4', 'strong', 'b']):
                if signer_type in heading.get_text().lower():
                    # Find the next container or paragraph for signature
                    signature_container = heading.find_next(['div', 'p'])
                    if signature_container:
                        signature_containers = [signature_container]
                        break
        
        # Inject signature into found containers
        for container in signature_containers:
            # Clear existing content
            container.clear()
            
            # Add name
            name_p = soup.new_tag('p')
            name_strong = soup.new_tag('strong')
            name_strong.string = signer_name
            name_p.append(name_strong)
            name_p['style'] = 'margin: 5px 0; font-size: 12pt;'
            container.append(name_p)
            
            # Add role
            role_p = soup.new_tag('p')
            role_p.string = f"Role: {sig.signer_type}"
            role_p['style'] = 'margin: 5px 0; font-size: 10pt;'
            container.append(role_p)
            
            # Add status
            status_p = soup.new_tag('p')
            status_text = f"Status: Signed on {sig.signed_at.strftime('%B %d, %Y at %I:%M %p')}" if sig.signed_at else "Status: Signed"
            status_p.string = status_text
            status_p['style'] = 'margin: 5px 0; font-size: 10pt; color: #22c55e;'
            container.append(status_p)
            
            # Add method
            if sig.signature_method:
                method_p = soup.new_tag('p')
                method_p.string = f"Method: {sig.signature_method}"
                method_p['style'] = 'margin: 5px 0; font-size: 9pt; font-style: italic; color: #64748b;'
                container.append(method_p)
            
            # Add signature image or text
            sig_div = soup.new_tag('div')
            sig_div['style'] = 'margin-top: 10px; padding: 10px; background: white; border: 1px solid #e2e8f0; border-radius: 4px;'
            
            if sig.signature_data.startswith('data:image'):
                # Image signature
                img = soup.new_tag('img')
                img['src'] = sig.signature_data
                img['alt'] = 'Signature'
                img['style'] = 'max-width: 200px; max-height: 80px; display: block;'
                sig_div.append(img)
            else:
                # Text signature
                text_p = soup.new_tag('p')
                text_p.string = sig.signature_data
                text_p['style'] = "font-family: 'Brush Script MT', cursive; font-size: 24pt; color: #000; font-style: italic; margin: 0;"
                sig_div.append(text_p)
            
            container.append(sig_div)
            
            # Only inject into first matching container
            break
    
    return str(soup)


import pypandoc
import tempfile
import os


def export_as_word(contract, signatories):
    """
    Generate Word document from contract content as-is
    Signatures should already be in the contract content from editor
    """
    
    # Get HTML content from database (includes embedded signatures)
    content_html = contract.contract_content or "<p>No content available</p>"
    
    # Create new Word document
    doc = Document()
    
    # Set page margins (same as PDF)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)  # 20mm
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69)  # A4 height
    
    # Clean HTML (remove scripts, styles)
    content_html = re.sub(r'<script[^>]*>.*?</script>', '', content_html, flags=re.DOTALL | re.IGNORECASE)
    content_html = re.sub(r'<style[^>]*>.*?</style>', '', content_html, flags=re.DOTALL | re.IGNORECASE)
    
    # Parse HTML and add to document (including embedded signatures)
    parser = ContractHTMLParser(doc)
    parser.feed(content_html)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{contract.contract_number or 'Contract'}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )


# =====================================================
# ALTERNATIVE: Use pypandoc for better conversion
# Install: pip install pypandoc
# Also install pandoc system tool: sudo apt-get install pandoc
# =====================================================

def export_as_word_with_pandoc(contract, signatories):
    """
    Alternative: Use pandoc for perfect HTML to DOCX conversion
    Pandoc preserves almost all HTML/CSS styling
    """
    
    import pypandoc
    
    # Build HTML
    content_html = contract.contract_content or "<p>No content available</p>"
    
    signature_html = ""
    if signatories:
        # Build signature HTML (same as above)
        pass
    
    complete_html = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{ font-family: Arial; font-size: 11pt; line-height: 1.6; }}
            /* Add more styles */
        </style>
    </head>
    <body>
        <h1>{contract.contract_title}</h1>
        {content_html}
        {signature_html}
    </body>
    </html>
    '''
    
    # Convert using pandoc
    try:
        docx_bytes = pypandoc.convert_text(
            complete_html,
            'docx',
            format='html',
            extra_args=['--standalone']
        )
        
        buffer = BytesIO(docx_bytes)
        buffer.seek(0)
        
        filename = f"{contract.contract_number or 'Contract'}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        logger.error(f"Pandoc conversion error: {e}")
        raise HTTPException(status_code=500, detail="DOCX generation failed")



def export_as_html(contract, signatories):
    """Generate clean HTML export"""
    
    content = contract.contract_content or "No content available"
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{contract.contract_title or 'Contract'}</title>
        <style>
            body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; }}
            h1 {{ color: #1e293b; border-bottom: 3px solid #2563eb; padding-bottom: 10px; }}
            .metadata {{ color: #64748b; margin-bottom: 20px; }}
            .content {{ line-height: 1.6; text-align: justify; }}
            .signatures {{ margin-top: 40px; }}
            .signature-block {{ border: 2px solid #cbd5e0; padding: 15px; margin: 10px 0; border-radius: 8px; }}
            img {{ max-width: 200px; max-height: 80px; }}
        </style>
    </head>
    <body>
       
        <div class="content">
            {content}
        </div>
    """
    
    if signatories:
        html += '<div class="signatures"><h2>Signatures</h2>'
        for sig in signatories:
            signer_name = f"{sig.first_name} {sig.last_name}" if sig.first_name else "Pending"
            status = "✓ Signed" if sig.has_signed else "Pending"
            html += f'''
            <div class="signature-block">
                <p><strong>{sig.signer_type.capitalize()}:</strong> {signer_name}</p>
                <p><strong>Status:</strong> {status}</p>
            '''
            if sig.signature_data and sig.signature_data.startswith('data:image'):
                html += f'<img src="{sig.signature_data}" alt="Signature"/>'
            html += '</div>'
        html += '</div>'
    
    html += '</body></html>'
    
    filename = f"{contract.contract_number or 'Contract'}_{datetime.now().strftime('%Y%m%d')}.html"
    
    return StreamingResponse(
        io.BytesIO(html.encode('utf-8')),
        media_type="text/html",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )


# =====================================================
# TRACK CHANGES ENDPOINTS
# =====================================================
@router.post("/track-changes/{contract_id}")
async def manage_track_changes(
    contract_id: int,
    action: str = Query(...),  # enable, disable, accept_all, reject_all
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Enable/disable track changes for contract - FIXED VERSION"""
    try:
        # Check contract exists
        contract_check = text("""
            SELECT id, status FROM contracts WHERE id = :contract_id
        """)
        
        result = db.execute(contract_check, {"contract_id": contract_id}).fetchone()
        
        if not result:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        message = ""
        
        if action == "enable":
            # Create a version snapshot when enabling track changes
            version_query = text("""
                INSERT INTO contract_versions (
                    contract_id, version_number, version_type, 
                    contract_content, change_summary, created_by, created_at
                )
                SELECT 
                    :contract_id,
                    COALESCE(MAX(version_number), 0) + 1,
                    'track_changes_enabled',
                    (SELECT contract_content FROM contract_versions 
                     WHERE contract_id = :contract_id 
                     ORDER BY version_number DESC LIMIT 1),
                    'Track changes enabled',
                    :user_id,
                    NOW()
                FROM contract_versions
                WHERE contract_id = :contract_id
            """)
            
            db.execute(version_query, {
                "contract_id": contract_id,
                "user_id": current_user.id
            })
            db.commit()
            
            message = "Track changes enabled"
            
        elif action == "disable":
            # Create a version when disabling
            version_query = text("""
                INSERT INTO contract_versions (
                    contract_id, version_number, version_type,
                    contract_content, change_summary, created_by, created_at
                )
                SELECT 
                    :contract_id,
                    COALESCE(MAX(version_number), 0) + 1,
                    'track_changes_disabled',
                    (SELECT contract_content FROM contract_versions 
                     WHERE contract_id = :contract_id 
                     ORDER BY version_number DESC LIMIT 1),
                    'Track changes disabled',
                    :user_id,
                    NOW()
                FROM contract_versions
                WHERE contract_id = :contract_id
            """)
            
            db.execute(version_query, {
                "contract_id": contract_id,
                "user_id": current_user.id
            })
            db.commit()
            
            message = "Track changes disabled"
            
        elif action == "accept_all":
            # Accept all changes - create a clean version
            version_query = text("""
                INSERT INTO contract_versions (
                    contract_id, version_number, version_type,
                    contract_content, change_summary, created_by, created_at
                )
                SELECT 
                    :contract_id,
                    COALESCE(MAX(version_number), 0) + 1,
                    'changes_accepted',
                    (SELECT contract_content FROM contract_versions 
                     WHERE contract_id = :contract_id 
                     ORDER BY version_number DESC LIMIT 1),
                    'All tracked changes accepted',
                    :user_id,
                    NOW()
                FROM contract_versions
                WHERE contract_id = :contract_id
            """)
            
            db.execute(version_query, {
                "contract_id": contract_id,
                "user_id": current_user.id
            })
            db.commit()
            
            message = "All changes accepted"
            
        elif action == "reject_all":
            message = "All changes rejected"
            
        else:
            raise HTTPException(status_code=400, detail="Invalid action")
        
        # Update contract's updated_at timestamp
        update_query = text("""
            UPDATE contracts 
            SET updated_at = NOW(), updated_by = :user_id
            WHERE id = :contract_id
        """)
        
        db.execute(update_query, {
            "contract_id": contract_id,
            "user_id": current_user.id
        })
        db.commit()
        
        #  FIXED: Log the action with correct parameter order
        log_contract_action(
            db=db,
            action_type=f"track_changes_{action}",  #  String first
            contract_id=contract_id,                 #  Int second
            user_id=current_user.id,                 #  Int third
            details={                                #  Dict fourth
                "action": action,
                "timestamp": datetime.utcnow().isoformat()
            }
        )
        
        return {
            "success": True,
            "message": message,
            "action": action,
            "contract_id": contract_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error managing track changes: {str(e)}")
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to manage track changes: {str(e)}"
              )
# =====================================================
# ADDITIONAL HELPER FUNCTIONS
# =====================================================

# def log_contract_action(db: Session, user_id: int, contract_id: int, action: str, details: dict = None):
#     """Log contract actions for audit trail"""
#     try:
#         audit_query = text("""
#             INSERT INTO audit_logs 
#             (user_id, action, entity_type, entity_id, details, created_at)
#             VALUES (:user_id, :action, 'contract', :entity_id, :details, NOW())
#         """)
        
#         db.execute(audit_query, {
#             "user_id": user_id,
#             "action": action,
#             "entity_id": contract_id,
#             "details": json.dumps(details) if details else None
#         })
#     except:
#         pass  # Don't fail main operation if audit fails

def send_notification(db: Session, user_id: int, title: str, message: str, type: str = "info"):
    """Send notification to user"""
    try:
        notif_query = text("""
            INSERT INTO notifications 
            (recipient_id, notification_type, title, message, status, created_at)
            VALUES (:user_id, :type, :title, :message, 'pending', NOW())
        """)
        
        db.execute(notif_query, {
            "user_id": str(user_id),
            "type": type,
            "title": title,
            "message": message
        })
    except:
        pass  # Don't fail main operation if notification fails


@router.post("/send-to-counterparty")
async def send_to_counterparty(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Send contract to counter-party for review and negotiation"""
    try:
        data = await request.json()
        
        # Simple - just get the flat fields
        contract_id = data.get('contract_id')
        counterparty_email = data.get('counterparty_email')
        counterparty_user_id = data.get('counterparty_user_id')
        counterparty_company_id = data.get('counterparty_company_id')
        counterparty_company_name = data.get('counterparty_company_name')
        message = data.get('message', '')
        
        # Get contract
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Update contract status
        contract.status = 'counterparty_internal_review'
        contract.party_b_id = counterparty_company_id
        contract.party_b_lead_id = counterparty_user_id
        contract.action_person_id = counterparty_user_id
        contract.updated_at = datetime.now()
        
        # Log the action
        logger.info(f"Contract {contract_id} sent to counter-party: {counterparty_email}")
        if counterparty_company_name:
            logger.info(f"Counter-party company: {counterparty_company_name} (ID: {counterparty_company_id}, User ID: {counterparty_user_id})")
        else:
            logger.info(f"External counter-party email: {counterparty_email}")
        
        db.commit()
        
        # TODO: Send email notification to counter-party
        # send_email_notification(counterparty_email, contract, message)
        
        return {
            "success": True,
            "message": "Contract sent to counter-party successfully",
            "contract_status": "negotiation",
            "counterparty_email": counterparty_email,
            "counterparty_company_name": counterparty_company_name
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        db.rollback()
        logger.error(f" Error sending to counter-party: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/quick-approve")
async def quick_approve_contract(
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Quick approve contract by counter-party"""
    try:
        data = await request.json()
        contract_id = data.get('contract_id')
        
        # Get contract details
        contract = db.query(Contract).filter(Contract.id == contract_id).first()
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Get initiator's company_id from the contract
        initiator_company_id = contract.company_id
        
        # Update contract status to negotiation_completed AND set action_person_id to contract initiator
        contract.status = 'negotiation_completed'
        contract.action_person_id = contract.created_by  # Set to contract initiator
        contract.updated_at = datetime.now()
        
        logger.info(f"👤 Action assigned to contract initiator: {contract.created_by}")
        
        # 🔄 STEP 1: Mark current user's (counter-party) workflow instance as 'completed'
        complete_counterparty_workflow = text("""
            UPDATE workflow_instances wi
            INNER JOIN workflows w ON wi.workflow_id = w.id
            SET wi.status = 'completed',
                wi.completed_at = NOW()
            WHERE wi.contract_id = :contract_id
            AND w.company_id = :counterparty_company_id
            AND wi.status IN ('pending', 'in_progress', 'active')
        """)
        db.execute(complete_counterparty_workflow, {
            "contract_id": contract_id,
            "counterparty_company_id": current_user.company_id
        })
        logger.info(f" Completed counter-party workflow for company {current_user.company_id}")
        
        # 🔄 STEP 2: Activate initiator's workflow instance to 'active'
        activate_initiator_workflow = text("""
            UPDATE workflow_instances wi
            INNER JOIN workflows w ON wi.workflow_id = w.id
            SET wi.status = 'active',
                wi.started_at = NOW()
            WHERE wi.contract_id = :contract_id
            AND w.company_id = :initiator_company_id
            AND wi.status IN ('pending', 'in_progress')
        """)
        db.execute(activate_initiator_workflow, {
            "contract_id": contract_id,
            "initiator_company_id": initiator_company_id
        })
        logger.info(f" Activated initiator workflow for company {initiator_company_id}")
        
        # Create activity log
        try:
            from sqlalchemy import text as sql_text
            activity_query = sql_text("""
                INSERT INTO contract_activity 
                (contract_id, action_type, action_by, notes, timestamp)
                VALUES 
                (:contract_id, 'quick_approved', :user_id, 'Contract quickly approved by counter-party', NOW())
            """)
            db.execute(activity_query, {
                "contract_id": contract_id,
                "user_id": current_user.id
            })
        except Exception as activity_err:
            logger.warning(f"Could not create activity log: {str(activity_err)}")
        
        db.commit()
        logger.info(f" Contract {contract_id} quick approved by user {current_user.id}")
        
        # TODO: Send notification to initiator
        # send_notification_to_initiator(contract)
        
        return {
            "success": True,
            "message": "Contract approved successfully. Initiator's workflow has been activated.",
            "contract_status": "negotiation_completed"
        }
        
    except HTTPException as he:
        raise he
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error approving contract: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/{contract_id}/complete-counterparty-review")
async def complete_counterparty_review(
    contract_id: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Complete counterparty review - moves to negotiation"""
    
    # Get client IP
    client_ip = request.client.host if request.client else "unknown"
    
    try:
        # Update status to negotiation
        update_query = text("""
            UPDATE contracts 
            SET status = 'negotiation', 
                updated_at = NOW() 
            WHERE id = :id
        """)
        db.execute(update_query, {"id": contract_id})
        
        audit_query = text("""
            INSERT INTO audit_logs (user_id, contract_id, action_type, action_details, ip_address, created_at)
            VALUES (:user_id, :contract_id, :action_type, :action_details, :ip_address, NOW())
        """)
        db.execute(audit_query, {
            "user_id": current_user.id,
            "contract_id": contract_id,
            "action_type": "COMPLETE_COUNTERPARTY_REVIEW",
            "action_details": json.dumps({
                "status_changed_from": "counterparty_internal_review",
                "status_changed_to": "negotiation"
            }),
            "ip_address": client_ip
        })
        
        db.commit()
        
        return {
            "success": True,
            "message": "Review completed",
            "new_status": "negotiation"
        }
        
    except Exception as e:
        db.rollback()
        return {"success": False, "message": str(e)}

def sanitize_for_json(text: str) -> str:
    """Sanitize text to be safely included in JSON responses"""
    if not text:
        return ""
    
    # Replace problematic characters
    text = text.replace('\n', ' ')  # Remove newlines
    text = text.replace('\r', ' ')  # Remove carriage returns
    text = text.replace('\t', ' ')  # Remove tabs
    text = text.replace('"', "'")   # Replace double quotes with single quotes
    text = text.replace('\\', '/')  # Replace backslashes
    
    # Limit length to prevent oversized responses
    if len(text) > 1000:
        text = text[:997] + "..."
    
    # Remove multiple spaces
    import re
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text


@router.post("/analyze-full-contract")
async def analyze_full_contract(
    request: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Analyze all clauses in the full contract content and provide suggestions
    for each identified clause.
    
    This endpoint:
    1. Takes the full contract text
    2. Uses Claude AI to identify ALL major clauses
    3. Provides specific suggestions for each clause
    4. Returns risk assessment and compliance notes
    """
    try:
        contract_text = request.get("contract_text", "")
        contract_id = request.get("contract_id")
        
        # Validate input
        if not contract_text or len(contract_text) < 50:
            raise HTTPException(
                status_code=400, 
                detail="Contract content is too short to analyze (minimum 50 characters)"
            )
        
        logger.info(f"📋 Analyzing full contract content ({len(contract_text)} characters)")
        
        # Initialize Claude service
        claude_service = ClaudeService()
        
        if not claude_service.client:
            logger.warning("Claude client not available, using fallback analysis")
            return get_fallback_clause_analysis(contract_text)
        
        # Prepare comprehensive prompt for Claude with strict JSON formatting rules
        prompt = f"""You are an expert contract analyst specializing in Qatar jurisdiction. Analyze the following complete contract and identify ALL major clauses.

CONTRACT TEXT:
{contract_text[:8000]}

Your task:
1. Identify and extract ALL major clauses in this contract including:
   - Governing Law & Jurisdiction
   - Payment Terms & Conditions
   - Term & Termination
   - Warranties & Representations
   - Indemnification & Liability
   - Confidentiality & Non-Disclosure
   - Dispute Resolution & Arbitration
   - Force Majeure
   - Intellectual Property Rights
   - Insurance Requirements
   - Service Level Agreements
   - Change Management
   - Any other significant clauses

2. For EACH clause identified, provide:
   - The exact clause name/title
   - A CONCISE SUMMARY of the clause content (max 300 characters)
   - A risk assessment: "low", "medium", or "high"
   - 2-4 specific, actionable suggestions for improvement
   - Compliance notes specific to Qatar Civil and Commercial Law
   - A legal protection score from 1-5
   - An improved version summary (max 300 characters)

3. Also identify:
   - Important clauses that are MISSING from the contract
   - Overall assessment of the contract quality
   - Total number of clauses analyzed

CRITICAL JSON FORMATTING RULES:
- Respond in VALID JSON format ONLY
- NO markdown, NO code blocks, NO extra text
- For clause_text: Use SUMMARY only (max 300 chars, single line)
- For improved_version: Use SUMMARY only (max 300 chars, single line)
- REPLACE all double quotes in text with single quotes
- NO newlines, tabs, or special characters in text fields
- Keep all text fields on single lines
- Properly escape any special characters

Expected JSON structure:
{{
    "clauses_identified": [
        {{
            "clause_name": "Governing Law and Jurisdiction",
            "clause_text": "Agreement governed by Qatar laws with disputes subject to Qatar Courts jurisdiction",
            "risk_level": "medium",
            "current_score": 3,
            "suggestions": [
                "Specify Qatar Courts or QICCA arbitration explicitly",
                "Add reference to Qatar Civil Code Articles",
                "Include bilingual interpretation clause"
            ],
            "compliance_note": "Complies with Qatar Civil Code Article 1. Consider Qatar Courts jurisdiction reference per Law No. 13 of 1990",
            "improved_version": "Governed by Qatar laws. Disputes subject to Qatar Courts exclusive jurisdiction or QICCA arbitration per parties agreement"
        }}
    ],
    "overall_assessment": "Contract shows moderate legal protection with clear commercial terms. Needs enhanced dispute resolution and liability provisions",
    "missing_clauses": [
        "Force Majeure provisions per Qatar Civil Code Article 215",
        "Insurance requirements and coverage specifications",
        "IP rights assignment terms",
        "Data protection compliance with Qatar Law"
    ],
    "total_clauses": 8
}}

IMPORTANT REMINDERS:
- Keep clause_text and improved_version BRIEF (max 300 chars each)
- Use apostrophes (') instead of quotes (") in all text
- Single line text only - no line breaks
- Valid JSON syntax - proper commas and brackets

Begin your analysis now. Respond with ONLY the valid JSON object."""

        try:
            # Call Claude API
            response = claude_service.client.messages.create(
                model=claude_service.model,
                max_tokens=8000,  # Increased for comprehensive analysis
                temperature=0.2,  # Lower temperature for more consistent JSON
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )
            
            # Extract response text
            response_text = response.content[0].text.strip()
            logger.info(f"📥 Received response from Claude ({len(response_text)} characters)")
            
            # Clean up response - remove markdown formatting
            cleaned_text = response_text
            
            # Remove markdown code blocks
            if "```json" in cleaned_text:
                cleaned_text = cleaned_text.split("```json")[1].split("```")[0].strip()
            elif "```" in cleaned_text:
                # Try to extract content between first ``` and last ```
                parts = cleaned_text.split("```")
                if len(parts) >= 3:
                    cleaned_text = parts[1].strip()
                else:
                    cleaned_text = cleaned_text.replace("```", "").strip()
            
            # Extract JSON object using regex
            json_match = re.search(r'\{[\s\S]*\}', cleaned_text)
            if json_match:
                cleaned_text = json_match.group()
            
            logger.info(f"🧹 Cleaned response ready for parsing ({len(cleaned_text)} chars)")
            
            # Parse JSON response
            try:
                result = json.loads(cleaned_text)
            except json.JSONDecodeError as json_err:
                logger.error(f" Initial JSON parse failed: {str(json_err)}")
                logger.error(f" Error at line {json_err.lineno}, column {json_err.colno}")
                logger.error(f" Problematic section: {cleaned_text[max(0, json_err.pos-100):json_err.pos+100]}")
                
                # Try additional cleanup
                # Replace common problematic patterns
                cleaned_text = cleaned_text.replace('\\"', "'")  # Replace escaped quotes
                cleaned_text = re.sub(r'[\n\r\t]', ' ', cleaned_text)  # Remove all whitespace chars
                cleaned_text = re.sub(r'\s+', ' ', cleaned_text)  # Collapse multiple spaces
                
                # Try parsing again
                result = json.loads(cleaned_text)
                logger.info(" Successfully parsed after additional cleanup")
            
            # Validate response structure
            if "clauses_identified" not in result:
                logger.error("Invalid response structure: missing clauses_identified")
                raise ValueError("Invalid response structure from AI")
            
            clauses = result.get("clauses_identified", [])
            
            if not clauses:
                logger.warning("No clauses identified by AI, using fallback")
                raise ValueError("No clauses identified in the contract")
            
            logger.info(f" Successfully identified {len(clauses)} clauses")
            
            # Sanitize all text fields to ensure JSON safety
            for idx, clause in enumerate(clauses):
                # Set defaults and sanitize for missing/invalid fields
                if "clause_name" not in clause or not clause["clause_name"]:
                    clause["clause_name"] = f"Clause {idx + 1}"
                else:
                    clause["clause_name"] = sanitize_for_json(str(clause["clause_name"]))
                
                if "clause_text" not in clause or not clause["clause_text"]:
                    clause["clause_text"] = "Summary not available"
                else:
                    clause["clause_text"] = sanitize_for_json(str(clause["clause_text"]))
                
                # Validate risk level
                if "risk_level" not in clause or clause["risk_level"] not in ["low", "medium", "high"]:
                    clause["risk_level"] = "medium"
                
                # Validate score
                if "current_score" not in clause:
                    clause["current_score"] = 3
                else:
                    try:
                        clause["current_score"] = int(clause["current_score"])
                        if clause["current_score"] < 1 or clause["current_score"] > 5:
                            clause["current_score"] = 3
                    except (ValueError, TypeError):
                        clause["current_score"] = 3
                
                # Sanitize suggestions array
                if "suggestions" not in clause or not isinstance(clause["suggestions"], list):
                    clause["suggestions"] = ["Review clause for completeness"]
                else:
                    clause["suggestions"] = [
                        sanitize_for_json(str(s)) for s in clause["suggestions"] if s
                    ][:6]  # Limit to 6 suggestions max
                    if not clause["suggestions"]:
                        clause["suggestions"] = ["Review clause for completeness"]
                
                if "compliance_note" not in clause or not clause["compliance_note"]:
                    clause["compliance_note"] = "Legal review recommended"
                else:
                    clause["compliance_note"] = sanitize_for_json(str(clause["compliance_note"]))
                
                if "improved_version" not in clause or not clause["improved_version"]:
                    clause["improved_version"] = "Improved version not available"
                else:
                    clause["improved_version"] = sanitize_for_json(str(clause["improved_version"]))
            
            # Sanitize overall assessment and missing clauses
            overall_assessment = sanitize_for_json(
                str(result.get("overall_assessment", "Contract analyzed successfully"))
            )
            
            missing_clauses = result.get("missing_clauses", [])
            if isinstance(missing_clauses, list):
                missing_clauses = [sanitize_for_json(str(mc)) for mc in missing_clauses if mc][:10]
            else:
                missing_clauses = []
            
            # Log the analysis to audit trail
            try:
                audit_log = text("""
                    INSERT INTO audit_logs 
                    (user_id, action_type, action_details, ip_address, user_agent, created_at)
                    VALUES 
                    (:user_id, :action_type, :action_details, :ip_address, :user_agent, NOW())
                """)
                
                db.execute(audit_log, {
                    "user_id": current_user.id,
                    "action_type": "AI_FULL_CONTRACT_ANALYSIS",
                    "action_details": json.dumps({
                        "contract_id": contract_id,
                        "contract_length": len(contract_text),
                        "clauses_identified": len(clauses),
                        "ai_powered": True
                    }),
                    "ip_address": "system",
                    "user_agent": "Claude AI"
                })
                db.commit()
            except Exception as audit_error:
                logger.error(f"Failed to log audit trail: {str(audit_error)}")
                # Continue even if audit logging fails
            
            logger.info(f"🎉 Analysis complete: {len(clauses)} clauses, {len(missing_clauses)} missing")
            
            return {
                "success": True,
                "clauses_identified": clauses,
                "overall_assessment": overall_assessment,
                "missing_clauses": missing_clauses,
                "total_clauses": len(clauses),
                "ai_powered": True
            }
            
        except json.JSONDecodeError as e:
            logger.error(f" JSON parsing error after all attempts: {str(e)}")
            logger.error(f" Response text (first 1000 chars): {cleaned_text[:1000]}")
            logger.error(f"📍 Error location: line {e.lineno}, column {e.colno}, position {e.pos}")
            logger.warning(" Falling back to pattern matching analysis")
            return get_fallback_clause_analysis(contract_text)
            
        except Exception as e:
            logger.error(f" Claude API error: {str(e)}")
            logger.warning(" Falling back to pattern matching analysis")
            return get_fallback_clause_analysis(contract_text)
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f" Error in full contract analysis: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


def get_fallback_clause_analysis(contract_text: str):
    """
    Provide fallback clause analysis when AI is unavailable
    Uses pattern matching to identify common clauses
    """
    clauses = []
    
    # Common clause patterns
    clause_patterns = {
        "Governing Law": ["governing law", "applicable law", "governed by"],
        "Payment Terms": ["payment", "compensation", "fees", "invoice"],
        "Termination": ["termination", "terminate", "cancellation"],
        "Confidentiality": ["confidential", "non-disclosure", "proprietary"],
        "Indemnification": ["indemnify", "indemnification", "hold harmless"],
        "Warranties": ["warrant", "representation", "guarantee"],
        "Liability": ["liability", "liable", "damages"],
        "Dispute Resolution": ["dispute", "arbitration", "mediation"],
        "Force Majeure": ["force majeure", "act of god", "beyond control"]
    }
    
    contract_lower = contract_text.lower()
    
    for clause_name, patterns in clause_patterns.items():
        # Check if any pattern exists in the contract
        if any(pattern in contract_lower for pattern in patterns):
            # Try to extract a snippet of text around the clause
            for pattern in patterns:
                if pattern in contract_lower:
                    start_idx = contract_lower.find(pattern)
                    # Extract 200 characters around the pattern
                    snippet_start = max(0, start_idx - 50)
                    snippet_end = min(len(contract_text), start_idx + 250)
                    clause_text = contract_text[snippet_start:snippet_end].strip()
                    
                    clauses.append({
                        "clause_name": clause_name,
                        "clause_text": f"...{clause_text}...",
                        "risk_level": "medium",
                        "current_score": 3,
                        "suggestions": [
                            f"Review {clause_name.lower()} clause for completeness",
                            "Ensure compliance with Qatar Civil and Commercial Law",
                            "Consider legal review for risk mitigation"
                        ],
                        "compliance_note": "Requires legal review for Qatar law compliance",
                        "improved_version": "AI analysis unavailable - legal review recommended"
                    })
                    break
    
    if not clauses:
        # Default response if no clauses found
        clauses = [{
            "clause_name": "General Contract",
            "clause_text": contract_text[:300] + "...",
            "risk_level": "medium",
            "current_score": 3,
            "suggestions": [
                "Add standard contract clauses (Governing Law, Payment Terms, Termination)",
                "Include dispute resolution provisions",
                "Add confidentiality and indemnification clauses"
            ],
            "compliance_note": "Review for Qatar law compliance",
            "improved_version": "Structure contract with standard clauses"
        }]
    
    return {
        "success": True,
        "clauses_identified": clauses,
        "overall_assessment": "Basic clause analysis completed. AI-powered analysis unavailable.",
        "missing_clauses": [
            "Consider adding standard clauses if not present",
            "Legal review recommended"
        ],
        "total_clauses": len(clauses),
        "ai_powered": False
    }



@router.get("/workflow/availability/{contract_id}")
async def check_workflow_availability(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Check if master and custom workflows are available for a contract
    """
    try:
        # Check for active master workflow for the company
        master_query = text("""
            SELECT id, workflow_name 
            FROM workflows 
            WHERE company_id = :company_id 
            AND is_master = 1 
            AND is_active = 1
            LIMIT 1
        """)
        
        master_workflow = db.execute(master_query, {
            "company_id": current_user.company_id
        }).fetchone()
        
        # Check for custom workflow for this specific contract
        custom_query = text("""
            SELECT w.id, w.workflow_name 
            FROM workflows w
            INNER JOIN workflow_instances wi ON w.id = wi.workflow_id
            WHERE wi.contract_id = :contract_id
            AND w.company_id = :company_id
            AND w.is_master = 0
           
            LIMIT 1
        """)
        
        custom_workflow = db.execute(custom_query, {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        }).fetchone()
        
        return {
            "success": True,
            "master_workflow": {
                "available": master_workflow is not None,
                "id": master_workflow.id if master_workflow else None,
                "name": master_workflow.workflow_name if master_workflow else None
            },
            "custom_workflow": {
                "available": custom_workflow is not None,
                "id": custom_workflow.id if custom_workflow else None,
                "name": custom_workflow.workflow_name if custom_workflow else None
            }
        }
        
    except Exception as e:
        logger.error(f"Error checking workflow availability: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))





@router.post("/workflow/ensure-instance/{contract_id}")
async def ensure_workflow_instance(
    contract_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Ensure workflow instance exists for contract
    - If master workflow exists for company but no instance → Create instance
    - If custom workflow exists but no instance → Create instance
    - Returns current workflow status
    """
    try:
        # Check if workflow instance already exists
        existing_instance_query = text("""
            SELECT wi.id, wi.workflow_id, wi.status, w.workflow_name, w.is_master
            FROM workflow_instances wi
            INNER JOIN workflows w ON wi.workflow_id = w.id
            WHERE wi.contract_id = :contract_id
            AND w.company_id = :company_id
            AND is_master=1
            LIMIT 1
        """)
        
        existing_instance = db.execute(existing_instance_query, {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        }).fetchone()
        
        if existing_instance:
            logger.info(f" Workflow instance already exists for contract {contract_id}")
            return {
                "success": True,
                "instance_exists": True,
                "workflow_id": existing_instance.workflow_id,
                "workflow_name": existing_instance.workflow_name,
                "is_master": existing_instance.is_master,
                "status": existing_instance.status,
                "message": "Workflow instance already exists"
            }
        
        # Check for master workflow
        master_workflow_query = text("""
            SELECT id, workflow_name 
            FROM workflows 
            WHERE company_id = :company_id 
            AND is_master = 1 
            AND is_active = 1
            LIMIT 1
        """)
        
        master_workflow = db.execute(master_workflow_query, {
            "company_id": current_user.company_id
        }).fetchone()
        
        # If master workflow exists, create instance
        if master_workflow:
            insert_instance = text("""
                INSERT INTO workflow_instances 
                (contract_id, workflow_id, status, current_step, started_at)
                VALUES (:contract_id, :workflow_id, 'pending', 1, NOW())
            """)
            
            db.execute(insert_instance, {
                "contract_id": contract_id,
                "workflow_id": master_workflow.id
            })
            
            db.commit()
            
            logger.info(f" Created workflow instance for contract {contract_id} using master workflow {master_workflow.id}")
            
            return {
                "success": True,
                "instance_exists": False,
                "instance_created": True,
                "workflow_id": master_workflow.id,
                "workflow_name": master_workflow.workflow_name,
                "is_master": True,
                "status": "pending",
                "message": "Master workflow instance created"
            }
        
        # No master workflow exists
        logger.info(f"ℹ️ No master workflow found for company {current_user.company_id}")
        return {
            "success": True,
            "instance_exists": False,
            "instance_created": False,
            "workflow_id": None,
            "workflow_name": None,
            "is_master": None,
            "status": None,
            "message": "No master workflow configured"
        }
        
    except Exception as e:
        db.rollback()
        logger.error(f" Error ensuring workflow instance: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/ai-generate-stream/{contract_id}")
async def stream_ai_contract_generation(
    contract_id: int,
    request_data: Dict[str, Any] = Body(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Stream AI-generated contract content directly to the editor"""
    try:
        logger.info(f" Streaming AI generation for contract {contract_id}")
        logger.info(f"📦 Received data: {json.dumps(request_data, indent=2)}")
        
        # Verify contract exists
        contract = db.query(Contract).filter(
            Contract.id == contract_id,
            Contract.company_id == current_user.company_id
        ).first()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # Extract values with safe defaults
        contract_type = request_data.get("contract_type", "Service Agreement")
        profile_type = request_data.get("profile_type", "client")
        parties = request_data.get("parties", {})
        party_a = parties.get("party_a", {})
        party_b = parties.get("party_b", {})
        contract_value = request_data.get("contract_value", "100,000")
        currency = request_data.get("currency", "QAR")
        start_date = request_data.get("start_date", "")
        end_date = request_data.get("end_date", "")
        selected_clause_descriptions = request_data.get("selected_clauses", [])
        jurisdiction = request_data.get("jurisdiction", "Qatar")
        language = request_data.get("language", "en")



        #  EXTRACT METADATA WITH SPECIAL REQUIREMENTS
        metadata = request_data.get("metadata", {})
        additional_requirements = metadata.get("additional_requirements", "")
        user_prompt = metadata.get("prompt", "")
        payment_terms = metadata.get("payment_terms", "")
        
        party_a_name = party_a.get("name", "Party A")
        party_b_name = party_b.get("name", "Party B")

        current_date = datetime.now().strftime("%d %B, %Y")
        # Build prompt
        prompt_text = f"""Generate a complete, production-ready {contract_type} contract:
- Party A ({profile_type}): 
Jurisdiction: {jurisdiction}

Selected Clauses: {', '.join(selected_clause_descriptions) if selected_clause_descriptions else 'Standard clauses'}

**CRITICAL REQUIREMENTS:**

1. Today's Date is: {current_date} - Use this as the reference date
2. Production-ready, contractually & legally binding (minimum 3,500 words)
3. Professional legal language with contractually & legally binding for {jurisdiction}
4. complete Contract should be written in {language} language
5. Every clause fully developed with procedures and consequences
6. {additional_requirements} 
7. {payment_terms}
8. {user_prompt}
9. Dont include signature section 
10. Write in Right to Left if Arabic

**FORMAT:**
- Clean HTML: <h2> for sections, <h3> for subsections, <p> for paragraphs
- Use <strong> for party names and defined terms
- Wrap in: <div class="contract-document">...</div>
- No markdown, no backticks, pure HTML only

Generate the complete contract now:"""

        logger.info("🚀 Starting Claude streaming...")
        
        # Define the generator function
        async def generate_stream():
            # Import text function here to ensure it's in scope
            from sqlalchemy import text as sql_text
            
            try:
                # Call Claude API with streaming
                with claude_service.client.messages.stream(
                    model=claude_service.model,
                    max_tokens=16000,
                    temperature=0.3,
                    messages=[{"role": "user", "content": prompt_text}]
                ) as stream:
                    accumulated_text = ""
                    
                    # Send initial metadata
                    yield f"data: {json.dumps({'type': 'start', 'contract_id': contract_id})}\n\n"
                    
                    # Stream content chunks
                    for text_chunk in stream.text_stream:
                        accumulated_text += text_chunk
                        yield f"data: {json.dumps({'type': 'content', 'text': text_chunk})}\n\n"
                    
                    # Get final metadata
                    final_message = stream.get_final_message()
                    tokens_used = final_message.usage.input_tokens + final_message.usage.output_tokens
                    word_count = len(accumulated_text.split())
                    
                    logger.info(f" Generated {word_count} words, {tokens_used} tokens")
                    
                    # Clean up markdown
                    for marker in ["```html", "```"]:
                        if accumulated_text.startswith(marker):
                            accumulated_text = accumulated_text[len(marker):].strip()
                        if accumulated_text.endswith("```"):
                            accumulated_text = accumulated_text[:-3].strip()
                    
                    # Save to database
                    clause_summary = ", ".join([c.split('-')[0].strip() for c in selected_clause_descriptions]) if selected_clause_descriptions else "Standard"
                    
                    # Check if version exists
                    existing_version = db.execute(sql_text("""
                        SELECT id FROM contract_versions 
                        WHERE contract_id = :contract_id AND version_number = 1
                    """), {"contract_id": contract_id}).fetchone()
                    
                    if existing_version:
                        db.execute(sql_text("""
                            UPDATE contract_versions 
                            SET contract_content = :contract_content,
                                change_summary = :change_summary,
                                created_at = :created_at
                            WHERE contract_id = :contract_id AND version_number = 1
                        """), {
                            "contract_id": contract_id,
                            "contract_content": accumulated_text,
                            "change_summary": f"AI-regenerated: {clause_summary}",
                            "created_at": datetime.utcnow()
                        })
                    else:
                        db.execute(sql_text("""
                            INSERT INTO contract_versions (contract_id, version_number, version_type,
                                                         contract_content, change_summary,
                                                         created_by, created_at)
                            VALUES (:contract_id, :version_number, :version_type,
                                    :contract_content, :change_summary,
                                    :created_by, :created_at)
                        """), {
                            "contract_id": contract_id,
                            "version_number": 1,
                            "version_type": "ai_generated",
                            "contract_content": accumulated_text,
                            "change_summary": f"AI-generated: {clause_summary}",
                            "created_by": str(current_user.id),
                            "created_at": datetime.utcnow()
                        })
                    
                    db.commit()
                    logger.info(f"💾 Saved to database")
                    
                    # Send completion
                    yield f"data: {json.dumps({'type': 'done', 'word_count': word_count, 'tokens_used': tokens_used})}\n\n"
                    
            except Exception as e:
                logger.error(f"❌ Streaming error: {str(e)}", exc_info=True)
                yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        
        return StreamingResponse(
            generate_stream(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Stream setup error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))



@router.put("/{contract_id}/update-metadata")
async def update_contract_metadata(
    contract_id: int,
    request_data: UpdateMetadataRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Update AI generation parameters for a contract
    This allows regeneration with modified metadata
     NOW INCLUDES BLOCKCHAIN HASH UPDATE
    """
    try:
        logger.info(f"📝 Updating metadata for contract {contract_id}")
        logger.info(f"📦 New params: {json.dumps(request_data.ai_generation_params, indent=2)}")

        # Verify contract exists and belongs to user's company
        contract_check = db.execute(text("""
          SELECT c.id, c.contract_number, c.is_ai_generated
          FROM contracts c
          WHERE c.id = :contract_id 
          AND c.company_id = :company_id
        """), {
            "contract_id": contract_id,
            "company_id": current_user.company_id
        }).fetchone()

        if not contract_check:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Contract not found"
            )

        # Verify it's an AI-generated contract
        if not contract_check.is_ai_generated:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Cannot update metadata for non-AI-generated contracts"
            )

        # Convert params to JSON string
        params_json = json.dumps(request_data.ai_generation_params)

        # Update the metadata
        db.execute(text("""
            UPDATE contracts 
            SET ai_generation_params = :params,
                updated_at = :updated_at
            WHERE id = :contract_id
        """), {
            "params": params_json,
            "updated_at": datetime.utcnow(),
            "contract_id": contract_id
        })

        db.commit()

        logger.info(f" Metadata updated for contract {contract_check.contract_number}")

        # =====================================================
        # 🔐 UPDATE BLOCKCHAIN HASH AFTER METADATA UPDATE
        # =====================================================
        try:
            logger.info(f"🔗 Updating blockchain hash after metadata change...")
            
            # Get current contract content
            contract_content = contract_check.content or ""
            
            # Update blockchain hash
            from app.services.blockchain_service import blockchain_service
            
            blockchain_result = await blockchain_service.store_contract_hash_with_logging(
                contract_id=contract_id,
                document_content=contract_content,
                uploaded_by=current_user.id,
                company_id=current_user.company_id,
                db=db
            )
            
            if blockchain_result.get("success"):
                logger.info(f" Blockchain hash updated successfully")
            else:
                logger.warning(f"⚠️ Blockchain update failed: {blockchain_result.get('error')}")
                
        except Exception as blockchain_error:
            # Don't fail the metadata update if blockchain fails
            logger.error(f"⚠️ Blockchain update error (non-critical): {str(blockchain_error)}")
        
        # Log the action
        log_contract_action(
            db=db,
            action_type="contract_updated",
            contract_id=contract_id,
            user_id=current_user.id,
            details={
                "update_type": "metadata_updated",
                "contract_number": contract_check.contract_number,
                "blockchain_updated": blockchain_result.get("success") if 'blockchain_result' in locals() else False
            },
            ip_address=None
        )

        return {
            "success": True,
            "message": "Metadata updated successfully",
            "contract_id": contract_id,
            "contract_number": contract_check.contract_number,
            "updated_params": request_data.ai_generation_params,
            "blockchain_updated": blockchain_result.get("success") if 'blockchain_result' in locals() else False
        }

    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Error updating metadata: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to update metadata: {str(e)}"
        )



@router.post("/documents/upload/{contract_id}")
async def upload_contract_documents(
    contract_id: int,
    files: List[UploadFile] = File(...),
    document_type: str = Form(default="contract_attachment"),
    notes: Optional[str] = Form(default=None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Upload supporting documents to a contract (UC 028)
    Supports exhibits, annexes, amendments, certificates, drafts
    """
    try:
        logger.info(f"📎 Uploading {len(files)} documents to contract {contract_id}")
        logger.info(f"👤 Current user company_id: {current_user.company_id}")
        
        # First, check if contract exists at all (for debugging)
        check_query = text("""
            SELECT id, company_id, contract_number, contract_title, project_id
            FROM contracts
            WHERE id = :contract_id
        """)
        
        contract_check = db.execute(check_query, {"contract_id": contract_id}).fetchone()
        
        if not contract_check:
            logger.error(f"❌ Contract {contract_id} does not exist in database")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Contract {contract_id} not found in database"
            )
        
        logger.info(f" Contract {contract_id} exists with company_id: {contract_check.company_id}")
        
        # Now verify user has access to this contract
        if str(contract_check.company_id) != str(current_user.company_id):
            logger.warning(f"⚠️ Company mismatch: Contract company_id={contract_check.company_id}, User company_id={current_user.company_id}")
            # For now, allow access (you can restrict this later)
            logger.info("🔓 Allowing access despite company mismatch (temp fix)")
        
        # Use the contract data we already fetched
        contract = contract_check
        
        # Define upload directory
        UPLOAD_BASE_DIR = Path("app/static/uploads/contract_documents")
        contract_upload_dir = UPLOAD_BASE_DIR / f"contract_{contract_id}"
        contract_upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Allowed file extensions per UC 028 requirements (15+ file types)
        ALLOWED_EXTENSIONS = {
            'pdf': 'application/pdf',
            'doc': 'application/msword',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'xls': 'application/vnd.ms-excel',
            'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'ppt': 'application/vnd.ms-powerpoint',
            'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            'txt': 'text/plain',
            'rtf': 'application/rtf',
            'odt': 'application/vnd.oasis.opendocument.text',
            'jpg': 'image/jpeg',
            'jpeg': 'image/jpeg',
            'png': 'image/png',
            'gif': 'image/gif',
            'tiff': 'image/tiff',
            'zip': 'application/zip',
            'msg': 'application/vnd.ms-outlook',
            'eml': 'message/rfc822'
        }
        
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB per UC 028
        
        uploaded_documents = []
        errors = []
        
        for file in files:
            try:
                # Validate file extension
                file_ext = Path(file.filename).suffix.lower().replace('.', '')
                if file_ext not in ALLOWED_EXTENSIONS:
                    errors.append({
                        "filename": file.filename,
                        "error": f"File type .{file_ext} not allowed"
                    })
                    continue
                
                # Read file content
                content = await file.read()
                file_size = len(content)
                
                # Validate file size (50MB limit)
                if file_size > MAX_FILE_SIZE:
                    errors.append({
                        "filename": file.filename,
                        "error": f"File size exceeds 50MB limit"
                    })
                    continue
                
                # Generate unique document ID and hash for blockchain
                doc_id = str(uuid.uuid4())
                file_hash = hashlib.sha256(content).hexdigest()
                
                # Save file with UUID prefix to prevent conflicts
                safe_filename = f"{doc_id}_{file.filename}"
                file_path = contract_upload_dir / safe_filename
                
                with open(file_path, "wb") as f:
                    f.write(content)
                
                # Get relative path for database
                relative_path = str(file_path.relative_to(Path("app")))
                
                # Prepare metadata per UC 028 requirements
                metadata = json.dumps({
                    "contract_id": int(contract_id),
                    "contract_number": contract.contract_number,
                    "contract_title": contract.contract_title,
                    "project_id": getattr(contract, 'project_id', None),
                    "original_filename": file.filename,
                    "notes": notes,
                    "upload_source": "contract_editor",
                    "uploader_email": current_user.email,
                    "document_purpose": document_type
                })
                
                # Insert into documents table
                insert_query = text("""
                    INSERT INTO documents (
                        id, company_id, contract_id, document_name, document_type,
                        file_path, file_size, mime_type, hash_value,
                        uploaded_by, uploaded_at, version, access_count, metadata
                    ) VALUES (
                        :id, :company_id, :contract_id, :document_name, :document_type,
                        :file_path, :file_size, :mime_type, :hash_value,
                        :uploaded_by, :uploaded_at, 1, 0, :metadata
                    )
                """)
                
                db.execute(insert_query, {
                    "id": doc_id,
                    "company_id": contract.company_id,  # Use contract's company_id
                    "contract_id": contract_id,
                    "document_name": file.filename,
                    "document_type": document_type,
                    "file_path": relative_path,
                    "file_size": file_size,
                    "mime_type": ALLOWED_EXTENSIONS.get(file_ext, 'application/octet-stream'),
                    "hash_value": file_hash,
                    "uploaded_by": current_user.id,
                    "uploaded_at": datetime.utcnow(),
                    "metadata": metadata
                })
                
                db.commit()
                
                uploaded_documents.append({
                    "id": doc_id,
                    "name": file.filename,
                    "size": file_size,
                    "type": file_ext,
                    "hash": file_hash[:16],  # First 16 chars for display
                    "uploaded_at": datetime.utcnow().isoformat()
                })
                
                logger.info(f" Uploaded: {file.filename} ({file_size} bytes)")
                
            except Exception as e:
                logger.error(f"❌ Failed to upload {file.filename}: {str(e)}", exc_info=True)
                errors.append({
                    "filename": file.filename,
                    "error": str(e)
                })
                db.rollback()
        
        # Return results
        return {
            "success": True,
            "message": f"Uploaded {len(uploaded_documents)} of {len(files)} files",
            "uploaded": uploaded_documents,
            "errors": errors if errors else None,
            "contract_id": contract_id,
            "total_files": len(files),
            "successful_uploads": len(uploaded_documents),
            "failed_uploads": len(errors)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error uploading contract documents: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload documents: {str(e)}"
        )


@router.get("/documents/list/{contract_id}")
async def get_contract_documents(
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get all documents attached to a contract
    """
    try:
        # First verify contract exists
        check_query = text("""
            SELECT id, company_id FROM contracts WHERE id = :contract_id
        """)
        
        contract = db.execute(check_query, {"contract_id": contract_id}).fetchone()
        
        if not contract:
            logger.error(f"❌ Contract {contract_id} not found")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Contract {contract_id} not found"
            )
        
        logger.info(f"📄 Listing documents for contract {contract_id} (company: {contract.company_id})")
        
        # Get documents - no company_id restriction for now
        query = text("""
            SELECT 
                d.id,
                d.document_name,
                d.document_type,
                d.file_size,
                d.mime_type,
                d.hash_value,
                d.uploaded_at,
                d.version,
                d.access_count,
                CONCAT(u.first_name, ' ', u.last_name) as uploaded_by_name,
                u.email as uploaded_by_email
            FROM documents d
            LEFT JOIN users u ON d.uploaded_by = u.id
            WHERE d.contract_id = :contract_id
            ORDER BY d.uploaded_at DESC
        """)
        
        results = db.execute(query, {"contract_id": contract_id}).fetchall()
        
        documents = [
            {
                "id": row.id,
                "name": row.document_name,
                "type": row.document_type,
                "size": row.file_size,
                "mime_type": row.mime_type,
                "hash": row.hash_value[:16] if row.hash_value else None,
                "uploaded_at": row.uploaded_at.isoformat() if row.uploaded_at else None,
                "uploaded_by": row.uploaded_by_name,
                "email": row.uploaded_by_email,
                "version": row.version,
                "downloads": row.access_count
            }
            for row in results
        ]
        
        return {
            "success": True,
            "contract_id": contract_id,
            "total_documents": len(documents),
            "documents": documents
        }
        
    except Exception as e:
        logger.error(f"Error fetching contract documents: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e)
        )


@router.delete("/documents/delete/{document_id}")
async def delete_contract_document(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Delete a contract document
    """
    try:
        # Get document info - no company restriction for now
        query = text("""
            SELECT file_path, document_name, contract_id, company_id
            FROM documents
            WHERE id = :document_id
        """)
        
        doc = db.execute(query, {"document_id": document_id}).fetchone()
        
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        logger.info(f"🗑️ Deleting document {document_id}: {doc.document_name} (company: {doc.company_id})")
        
        # Delete file from filesystem
        file_path = Path("app") / doc.file_path
        if file_path.exists():
            file_path.unlink()
            logger.info(f"🗑️ Deleted file: {file_path}")
        
        # Delete from database
        delete_query = text("""
            DELETE FROM documents
            WHERE id = :document_id
        """)
        
        db.execute(delete_query, {"document_id": document_id})
        
        db.commit()
        
        return {
            "success": True,
            "message": f"Document '{doc.document_name}' deleted successfully",
            "document_id": document_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/documents/download/{document_id}")
async def download_contract_document(
    document_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Download a contract document
    Returns the file with proper headers for download
    """
    try:
        # Get document info
        query = text("""
            SELECT 
                d.file_path,
                d.document_name,
                d.mime_type,
                d.file_size,
                d.contract_id,
                d.company_id,
                d.access_count
            FROM documents d
            WHERE d.id = :document_id
        """)
        
        doc = db.execute(query, {"document_id": document_id}).fetchone()
        
        if not doc:
            logger.error(f"❌ Document {document_id} not found")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        logger.info(f"⬇️ Downloading document: {doc.document_name} (ID: {document_id})")
        
        # Construct full file path
        file_path = Path("app") / doc.file_path
        
        # Verify file exists
        if not file_path.exists():
            logger.error(f"❌ File not found on disk: {file_path}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"File not found on server: {doc.document_name}"
            )
        
        # Update access count
        update_query = text("""
            UPDATE documents
            SET access_count = access_count + 1
            WHERE id = :document_id
        """)
        
        db.execute(update_query, {"document_id": document_id})
        db.commit()
        
        logger.info(f" Serving file: {file_path} ({doc.file_size} bytes)")
        
        # Return file as download
        from fastapi.responses import FileResponse
        
        return FileResponse(
            path=str(file_path),
            media_type=doc.mime_type or 'application/octet-stream',
            filename=doc.document_name,
            headers={
                "Content-Disposition": f'attachment; filename="{doc.document_name}"',
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error downloading document: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to download document: {str(e)}"
        )


@router.get("/documents/view/{document_id}")
async def view_contract_document(
    document_id: str,
    inline: bool = Query(default=True, description="View inline in browser"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    View a contract document in browser (inline) or download
    Use inline=true for PDFs to open in browser, inline=false to force download
    """
    try:
        # Get document info
        query = text("""
            SELECT 
                d.file_path,
                d.document_name,
                d.mime_type,
                d.file_size,
                d.contract_id
            FROM documents d
            WHERE d.id = :document_id
        """)
        
        doc = db.execute(query, {"document_id": document_id}).fetchone()
        
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Construct full file path
        file_path = Path("app") / doc.file_path
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found on server")
        
        # Update access count
        update_query = text("""
            UPDATE documents
            SET access_count = access_count + 1
            WHERE id = :document_id
        """)
        
        db.execute(update_query, {"document_id": document_id})
        db.commit()
        
        logger.info(f"👁️ Viewing document: {doc.document_name} (inline={inline})")
        
        from fastapi.responses import FileResponse
        
        # Determine content disposition
        disposition = "inline" if inline else "attachment"
        
        return FileResponse(
            path=str(file_path),
            media_type=doc.mime_type or 'application/octet-stream',
            filename=doc.document_name,
            headers={
                "Content-Disposition": f'{disposition}; filename="{doc.document_name}"',
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error viewing document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/search")
async def search_contracts(
    q: str = Query(..., min_length=1, description="Search query"),
    limit: int = Query(10, ge=1, le=50),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Search contracts by title, number, or counterparty
    """
    try:
        logger.info(f"🔍 Searching contracts for: {q}")
        
        # Build search query
        search_pattern = f"%{q}%"
        
        query_sql = text("""
        SELECT 
            c.id,
            c.contract_number,
            c.contract_title,
            c.contract_type,
            c.status,
            c.party_b_name as counterparty_name,
            c.contract_value,
            c.currency,
            c.created_at,
            c.updated_at
        FROM contracts c
        WHERE c.company_id = :company_id
        AND c.is_deleted = 0
        AND (
            c.contract_title LIKE :search
            OR c.contract_number LIKE :search
            OR c.party_b_name LIKE :search
            OR c.contract_type LIKE :search
        )
        ORDER BY c.updated_at DESC
        LIMIT :limit
        """)
        
        result = db.execute(query_sql, {
            "company_id": current_user.company_id,
            "search": search_pattern,
            "limit": limit
        })
        
        rows = result.fetchall()
        
        contracts = []
        for row in rows:
            contracts.append({
                "id": row[0],
                "contract_number": row[1],
                "contract_title": row[2],
                "contract_type": row[3],
                "status": row[4],
                "counterparty_name": row[5],
                "contract_value": float(row[6]) if row[6] else 0,
                "currency": row[7],
                "created_at": str(row[8]) if row[8] else None,
                "updated_at": str(row[9]) if row[9] else None
            })
        
        logger.info(f" Found {len(contracts)} contracts matching '{q}'")
        
        return {
            "success": True,
            "results": contracts,
            "total": len(contracts),
            "query": q
        }
        
    except Exception as e:
        logger.error(f"❌ Error searching contracts: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to search contracts: {str(e)}"
        )


@router.put("/{contract_id}/esignature")
async def send_for_esignature(
    request: Request,
    contract_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Send contract for e-signature - Simple version"""
    try:
        data = await request.json()
        party_esignature_authority_id = data.get('party_esignature_authority_id')
        counterparty_esignature_authority_id = data.get('counterparty_esignature_authority_id')
        
        # 1. Check contract exists
        contract = db.query(Contract).filter(
            Contract.id == contract_id
        ).first()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
        
        # 2. Update contract status
        contract.party_esignature_authority_id = party_esignature_authority_id
        contract.counterparty_esignature_authority_id = counterparty_esignature_authority_id
        contract.status = 'signature'
        contract.updated_at = datetime.now()
        
        db.commit()
        
        # =====================================================
        # 📧 SEND EMAIL NOTIFICATIONS TO E-SIGNATURE AUTHORITIES
        # =====================================================
        try:
            from app.services.workflow_email_service import WorkflowEmailService
            
            # Send email to Party A (Initiator) E-Signature Authority
            if party_esignature_authority_id:
                party_a_query = text("""
                    SELECT 
                        u.email,
                        CONCAT(u.first_name, ' ', u.last_name) as full_name
                    FROM users u
                    WHERE u.id = :user_id
                """)
                party_a_info = db.execute(party_a_query, {"user_id": party_esignature_authority_id}).fetchone()
                
                if party_a_info and party_a_info.email:
                    WorkflowEmailService.send_contract_sent_for_signature(
                        db=db,
                        contract_id=contract_id,
                        contract_number=contract.contract_number,
                        contract_title=contract.contract_title,
                        esign_authority_email=party_a_info.email,
                        esign_authority_name=party_a_info.full_name,
                        party_type="Party A (Initiator)"
                    )
                    logger.info(f"✉️ E-signature email sent to Party A: {party_a_info.email}")
                else:
                    logger.warning(f"⚠️ Party A e-sign authority email not found for user ID: {party_esignature_authority_id}")
            
            # Send email to Party B (Counter-Party) E-Signature Authority
            if counterparty_esignature_authority_id:
                party_b_query = text("""
                    SELECT 
                        u.email,
                        CONCAT(u.first_name, ' ', u.last_name) as full_name
                    FROM users u
                    WHERE u.id = :user_id
                """)
                party_b_info = db.execute(party_b_query, {"user_id": counterparty_esignature_authority_id}).fetchone()
                
                if party_b_info and party_b_info.email:
                    WorkflowEmailService.send_contract_sent_for_signature(
                        db=db,
                        contract_id=contract_id,
                        contract_number=contract.contract_number,
                        contract_title=contract.contract_title,
                        esign_authority_email=party_b_info.email,
                        esign_authority_name=party_b_info.full_name,
                        party_type="Party B (Counter-Party)"
                    )
                    logger.info(f"✉️ E-signature email sent to Party B: {party_b_info.email}")
                else:
                    logger.warning(f"⚠️ Party B e-sign authority email not found for user ID: {counterparty_esignature_authority_id}")
            
            logger.info("✅ E-signature notification emails sent to both authorities")
            
        except Exception as email_error:
            logger.error(f"❌ Error sending e-signature emails: {str(email_error)}")
            # Don't fail the operation if email fails
        
        return {
            "success": True,
            "message": "Contract sent for e-signature",
            "status": "pending_signature"
        }
        
    except Exception as e:
        db.rollback()
        logger.error(f"Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/activity/{contract_id}")
async def get_contract_activity_logs(
    contract_id: int,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get simple activity logs for a contract - who did what and when
    """
    try:
        logger.info(f"📊 Fetching activity logs for contract {contract_id}")
        
        # Get activity logs with user details
        query = text("""
            SELECT 
                al.id,
                al.action_type,
                al.action_details,
                al.created_at,
                al.ip_address,
                u.id as user_id,
                CONCAT(u.first_name, ' ', u.last_name) as user_name,
                u.email as user_email
            FROM audit_logs al
            LEFT JOIN users u ON al.user_id = u.id
            WHERE al.contract_id = :contract_id
            ORDER BY al.created_at DESC
            LIMIT :limit
        """)
        
        result = db.execute(query, {
            "contract_id": contract_id,
            "limit": limit
        })
        
        activities = []
        for row in result:
            # Parse action_details if it's a JSON string
            action_details = {}
            if row.action_details:
                try:
                    if isinstance(row.action_details, str):
                        action_details = json.loads(row.action_details)
                    else:
                        action_details = row.action_details
                except:
                    action_details = {"raw": str(row.action_details)}
            
            # Format action description
            action_desc = get_action_description(row.action_type, action_details)
            
            activities.append({
                "id": row.id,
                "action_type": row.action_type,
                "action_description": action_desc,
                "user_name": row.user_name or "System",
                "user_email": row.user_email or "system@calim360.com",
                "timestamp": row.created_at.isoformat() + 'Z' if row.created_at else None,
                "ip_address": row.ip_address,
                "details": action_details
            })
        
        # Get statistics
        stats_query = text("""
            SELECT 
                COUNT(*) as total_activities,
                MAX(created_at) as last_activity
            FROM audit_logs
            WHERE contract_id = :contract_id
        """)
        
        stats_result = db.execute(stats_query, {"contract_id": contract_id}).fetchone()
        
        return {
            "success": True,
            "contract_id": contract_id,
            "activities": activities,
            "statistics": {
                "total_activities": stats_result.total_activities or 0,
                "last_activity_time": stats_result.last_activity.isoformat() + 'Z' if stats_result.last_activity else None  # 🔥 ADD 'Z'
            }
        }
        
    except Exception as e:
        logger.error(f"❌ Error fetching activity logs: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error fetching activity logs: {str(e)}"
        )


def get_action_description(action_type: str, details: dict) -> str:
    """
    Convert action type to human-readable description
    """
    descriptions = {
        "contract_created": "Created the contract",
        "contract_updated": "Updated contract details",
        "contract_deleted": "Deleted the contract",
        "contract_signed": "Signed the contract",
        "contract_submitted": "Submitted for review",
        "contract_approved": "Approved the contract",
        "contract_rejected": "Rejected the contract",
        "obligation_created": "Created new obligation",
        "obligation_updated": "Updated obligation",
        "obligation_completed": "Completed obligation",
        "document_uploaded": "Uploaded document",
        "document_downloaded": "Downloaded document",
        "comment_added": "Added comment",
        "workflow_started": "Started workflow",
        "workflow_completed": "Completed workflow step",
        "clause_added": "Added contract clause",
        "clause_updated": "Updated contract clause",
        "version_created": "Created new version",
        "signature_requested": "Requested signature",
        "negotiation_started": "Started negotiation",
        "ai_generation": "Generated content using AI",
        "blockchain_storage": "Stored on blockchain"
    }
    
    # Get description or default
    description = descriptions.get(action_type, action_type.replace('_', ' ').title())
    
    # Add specific details if available
    if details.get("entity_type"):
        description += f" ({details['entity_type']})"
    
    return description


@router.post("/{contract_id}/recover-tampered-version")
async def recover_tampered_version(
    contract_id: int,
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete latest tampered version"""
    
    # Check authorization
    contract_check = db.execute(text("""
        SELECT created_by FROM contracts WHERE id = :cid
    """), {"cid": contract_id}).fetchone()
    
    if not contract_check:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    # Get user attributes safely
    user_id = getattr(current_user, 'id', None) or getattr(current_user, 'user_id', None)
    user_type = getattr(current_user, 'user_type', None) or getattr(current_user, 'type', None)
    
    is_internal = user_type == "internal"
    is_initiator = user_id == contract_check.created_by
    
    if not (is_internal or is_initiator):
        raise HTTPException(status_code=403, detail="Not authorized")
    
    # Get latest version
    current = db.execute(text("""
        SELECT id, version_number
        FROM contract_versions
        WHERE contract_id = :cid
        ORDER BY version_number DESC
        LIMIT 1
    """), {"cid": contract_id}).fetchone()
    
    if not current:
        raise HTTPException(status_code=404, detail="No versions found")
    
    # Check if there's a previous version
    previous = db.execute(text("""
        SELECT id
        FROM contract_versions
        WHERE contract_id = :cid AND version_number < :vnum
        ORDER BY version_number DESC
        LIMIT 1
    """), {"cid": contract_id, "vnum": current.version_number}).fetchone()
    
    if not previous:
        raise HTTPException(status_code=400, detail="No previous version available to recover")
    
    try:
        # Simply delete the latest version
        db.execute(text("DELETE FROM contract_versions WHERE id = :vid"), 
                   {"vid": current.id})
        
        db.commit()
        
        return {
            "success": True, 
            "message": "Latest version deleted successfully. Please refresh the page."
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Recovery failed: {str(e)}")