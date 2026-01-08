# File: app/services/document_generator.py
"""
Document Generator Service
Generates Word documents for correspondence with HTML formatting support
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import logging
import re
from html.parser import HTMLParser

logger = logging.getLogger(__name__)


class HTMLToDocxParser(HTMLParser):
    """Parser to convert HTML to Word document with formatting"""
    
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.list_stack = []
        self.bold_stack = []
        self.italic_stack = []
        self.underline_stack = []
        self.heading_level = None
        self.in_list_item = False
        self.preserve_newlines = False
        
    def handle_starttag(self, tag, attrs):
        """Handle HTML opening tags"""
        tag = tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = int(tag[1])
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.style = f'Heading {self.heading_level}'
            
        elif tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.space_after = Pt(10)
            self.current_paragraph.paragraph_format.line_spacing = 1.15
            
        elif tag == 'br':
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run('\n')
                
        elif tag in ['ul', 'ol']:
            self.list_stack.append(tag)
            
        elif tag == 'li':
            self.in_list_item = True
            self.current_paragraph = self.doc.add_paragraph()
            
            # Add bullet or number based on list type
            if self.list_stack and self.list_stack[-1] == 'ul':
                self.current_paragraph.style = 'List Bullet'
            else:
                self.current_paragraph.style = 'List Number'
                
        elif tag == 'strong' or tag == 'b':
            self.bold_stack.append(True)
            
        elif tag == 'em' or tag == 'i':
            self.italic_stack.append(True)
            
        elif tag == 'u':
            self.underline_stack.append(True)
            
        elif tag == 'hr':
            # Add horizontal line
            p = self.doc.add_paragraph()
            p_border = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), 'auto')
            p_border.append(bottom_border)
            p._element.get_or_add_pPr().append(p_border)
            
        elif tag == 'blockquote':
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.left_indent = Inches(0.5)
            self.current_paragraph.paragraph_format.space_before = Pt(6)
            self.current_paragraph.paragraph_format.space_after = Pt(6)
            
        elif tag == 'code':
            if not self.current_paragraph:
                self.current_paragraph = self.doc.add_paragraph()
                
        elif tag == 'pre':
            self.current_paragraph = self.doc.add_paragraph()
            self.preserve_newlines = True
            
    def handle_endtag(self, tag):
        """Handle HTML closing tags"""
        tag = tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            self.heading_level = None
            self.current_paragraph = None
            
        elif tag == 'p':
            self.current_paragraph = None
            
        elif tag in ['ul', 'ol']:
            if self.list_stack:
                self.list_stack.pop()
                
        elif tag == 'li':
            self.in_list_item = False
            self.current_paragraph = None
            
        elif tag == 'strong' or tag == 'b':
            if self.bold_stack:
                self.bold_stack.pop()
                
        elif tag == 'em' or tag == 'i':
            if self.italic_stack:
                self.italic_stack.pop()
                
        elif tag == 'u':
            if self.underline_stack:
                self.underline_stack.pop()
                
        elif tag == 'blockquote':
            self.current_paragraph = None
            
        elif tag == 'pre':
            self.preserve_newlines = False
            self.current_paragraph = None
            
    def handle_data(self, data):
        """Handle text data"""
        if not data.strip() and not self.preserve_newlines:
            return
            
        # Create paragraph if needed
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.space_after = Pt(10)
            self.current_paragraph.paragraph_format.line_spacing = 1.15
        
        # Add run with formatting
        run = self.current_paragraph.add_run(data)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        
        # Apply formatting from stacks
        if self.bold_stack:
            run.bold = True
        if self.italic_stack:
            run.italic = True
        if self.underline_stack:
            run.underline = True
            
        # Apply heading formatting
        if self.heading_level:
            run.bold = True
            if self.heading_level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(26, 54, 93)
            elif self.heading_level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(39, 98, 203)
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(51, 51, 51)
        
        self.current_run = run


class DocumentGenerator:
    """Service for generating Word documents"""
    
    @staticmethod
    def generate_correspondence_docx(
        content: str,
        subject: str = None,
        sender_name: str = None,
        recipient_name: str = None,
        reference: str = None
    ) -> BytesIO:
        """
        Generate a professional Word document for correspondence
        
        Args:
            content: Main content/body text (can be HTML or plain text)
            subject: Subject line
            sender_name: Name of sender
            recipient_name: Name of recipient
            reference: Reference number
            
        Returns:
            BytesIO object containing the Word document
        """
        
        try:
            logger.info("📝 Generating Word document with formatting support")
            
            # Create document
            doc = Document()
            
            # Set document margins (in inches)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add company header
            header_para = doc.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_run = header_para.add_run('CALIM 360 - Smart Contract Lifecycle Management')
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.color.rgb = RGBColor(39, 98, 203)  # #2762cb
            
            doc.add_paragraph()  # Empty line
            
            # Add date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f'Date: {datetime.now().strftime("%d %B %Y")}')
            date_run.font.size = Pt(11)
            
            # Add reference if provided
            if reference:
                ref_para = doc.add_paragraph()
                ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ref_run = ref_para.add_run(f'Ref: {reference}')
                ref_run.font.size = Pt(11)
            
            doc.add_paragraph()  # Empty line
            
            # Add recipient if provided
            if recipient_name:
                recipient_para = doc.add_paragraph(f'To: {recipient_name}')
                recipient_para.paragraph_format.space_after = Pt(6)
            
            # Add subject if provided
            if subject:
                subject_para = doc.add_paragraph()
                subject_run = subject_para.add_run(f'Subject: {subject}')
                subject_run.bold = True
                subject_run.font.size = Pt(12)
                subject_para.paragraph_format.space_after = Pt(12)
            
            doc.add_paragraph()  # Empty line
            
            # Parse and add main content with HTML formatting
            content = content.strip()
            
            # Check if content is HTML
            if DocumentGenerator._is_html(content):
                logger.info("📄 Parsing HTML content with formatting")
                DocumentGenerator._parse_html_to_docx(content, doc)
            else:
                logger.info("📄 Processing plain text content")
                # Handle plain text with paragraph breaks
                paragraphs = content.split('\n\n')
                
                for para_text in paragraphs:
                    if para_text.strip():
                        # Handle single newlines within paragraphs
                        lines = para_text.split('\n')
                        para = doc.add_paragraph()
                        
                        for i, line in enumerate(lines):
                            if line.strip():
                                if i > 0:
                                    para.add_run('\n')
                                run = para.add_run(line.strip())
                                run.font.size = Pt(11)
                                run.font.name = 'Calibri'
                        
                        para.paragraph_format.space_after = Pt(10)
                        para.paragraph_format.line_spacing = 1.15
            
            doc.add_paragraph()  # Empty line
            
            # Add sender signature if provided
            if sender_name:
                doc.add_paragraph()
                signature_para = doc.add_paragraph('Yours faithfully,')
                signature_para.paragraph_format.space_after = Pt(30)
                
                sender_para = doc.add_paragraph(sender_name)
                sender_run = sender_para.runs[0]
                sender_run.bold = True
            
            # Add footer
            doc.add_paragraph()
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(
                '_______________________________________________\n'
                'Generated by CALIM 360'
            )
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            logger.info(" Word document generated successfully with formatting")
            return docx_buffer
            
        except Exception as e:
            logger.error(f"❌ Error generating Word document: {str(e)}")
            raise Exception(f"Failed to generate Word document: {str(e)}")
    
    @staticmethod
    def _is_html(content: str) -> bool:
        """Check if content contains HTML tags"""
        html_pattern = re.compile(r'<[^>]+>')
        return bool(html_pattern.search(content))
    
    @staticmethod
    def _parse_html_to_docx(html_content: str, doc: Document):
        """Parse HTML content and add to Word document with formatting"""
        
        # Clean up HTML
        html_content = DocumentGenerator._clean_html(html_content)
        
        # Parse HTML using custom parser
        parser = HTMLToDocxParser(doc)
        parser.feed(html_content)
    
    @staticmethod
    def _clean_html(html_content: str) -> str:
        """Clean HTML content for parsing"""
        
        # Remove script and style tags
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert common entities
        html_content = html_content.replace('&nbsp;', ' ')
        html_content = html_content.replace('&amp;', '&')
        html_content = html_content.replace('&lt;', '<')
        html_content = html_content.replace('&gt;', '>')
        html_content = html_content.replace('&quot;', '"')
        html_content = html_content.replace('&#39;', "'")
        
        # Remove inline styles to prevent parsing issues
        html_content = re.sub(r'\s*style="[^"]*"', '', html_content)
        html_content = re.sub(r'\s*class="[^"]*"', '', html_content)
        html_content = re.sub(r'\s*id="[^"]*"', '', html_content)
        
        return html_content
    
    @staticmethod
    def generate_simple_docx(content: str) -> BytesIO:
        """
        Generate a simple Word document with just content
        
        Args:
            content: Text content to include
            
        Returns:
            BytesIO object containing the Word document
        """
        
        try:
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add content
            if DocumentGenerator._is_html(content):
                DocumentGenerator._parse_html_to_docx(content, doc)
            else:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.paragraph_format.space_after = Pt(10)
                        for run in para.runs:
                            run.font.size = Pt(11)
                            run.font.name = 'Calibri'
            
            # Save to buffer
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer
            
        except Exception as e:
            logger.error(f"❌ Error generating simple Word document: {str(e)}")
            raise Exception(f"Failed to generate Word document: {str(e)}")